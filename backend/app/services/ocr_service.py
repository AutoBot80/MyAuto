"""AI reader queue and Add Sales upload extraction: Textract FORMS for detail sheets; Aadhaar uses Textract text only (no UIDAI QR decode)."""

from __future__ import annotations

import json
import logging
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Any

from app.db import get_connection
from app.repositories.ai_reader_queue import AiReaderQueueRepository
from app.services.customer_address_infer import (
    enrich_customer_address_from_freeform,
    normalize_address_freeform,
)

logger = logging.getLogger(__name__)


def _aadhar_last4(aadhar_id: str | None) -> str | None:
    """Return only last 4 digits of Aadhar for compliance; never persist full 12 digits."""
    if not aadhar_id or not str(aadhar_id).strip():
        return None
    digits = "".join(c for c in str(aadhar_id) if c.isdigit())
    return digits[-4:] if len(digits) >= 4 else digits or None


def _safe_subfolder_name(subfolder: str) -> str:
    """Safe directory name under ocr_output (one segment, no path separators)."""
    return re.sub(r"[^\w\-]", "_", (subfolder or "").strip()) or "default"


def _safe_output_basename(subfolder: str, filename: str) -> str:
    """Build a safe filename for output: subfolder_filename.txt (no path chars)."""
    safe_sub = _safe_subfolder_name(subfolder)
    safe_name = Path(filename).stem
    safe_name = re.sub(r"[^\w\-.]", "_", safe_name)
    return f"{safe_sub}_{safe_name}.txt"


def _ocr_subfolder_path(output_dir: Path, subfolder: str) -> Path:
    """Path to per-customer subfolder under ocr_output: ocr_output/mobile_ddmmyyyy/."""
    return output_dir / _safe_subfolder_name(subfolder)


# OCR output JSON filename (in ocr_output subfolder only; input files stay as Details.jpg)
OCR_OUTPUT_JSON_STEM = "OCR_To_be_Used"


def _json_output_path(output_dir: Path, subfolder: str) -> Path:
    """Path for structured JSON output: ocr_output/mobile_ddmmyy/OCR_To_be_Used.json."""
    subfolder_name = _safe_subfolder_name(subfolder)
    return output_dir / subfolder_name / f"{OCR_OUTPUT_JSON_STEM}.json"


# Filename patterns for queue processing.
DETAILS_FILENAME_HINTS = (
    "details",
    "detail_sheet",
    "sales_detail",
    "sales detail",
    "a5 template",
)
AADHAR_FILENAME_CONTAINS = "aadhar"

# 15 Aadhar fields (display order and labels for ocr_output file).
AADHAR_15_FIELDS: list[tuple[str, str]] = [
    ("aadhar_id", "Aadhar ID"),
    ("name", "Name"),
    ("gender", "Gender"),
    ("year_of_birth", "Year of birth"),
    ("date_of_birth", "Date of birth"),
    ("care_of", "Care of"),
    ("house", "House"),
    ("street", "Street"),
    ("location", "Location"),
    ("city", "City"),
    ("post_office", "Post Office"),
    ("district", "District"),
    ("sub_district", "Sub District"),
    ("state", "State"),
    ("pin_code", "Pin Code"),
]

# Customer keys for Aadhaar merges (same shape as legacy UIDAI POA; filled from Textract parsers).
AADHAR_QR_CUSTOMER_KEYS: tuple[str, ...] = (
    "aadhar_id",
    "name",
    "gender",
    "year_of_birth",
    "date_of_birth",
    "care_of",
    "house",
    "street",
    "location",
    "city",
    "post_office",
    "district",
    "sub_district",
    "state",
    "pin_code",
)


def _rebuild_aadhar_address_from_parts(
    customer: dict[str, str],
    *,
    preserve_existing_address: bool = False,
) -> None:
    """
    UIDAI POA often uses city (vtc), district, post office without house/street.
    Mutates `customer` in place: sets or clears `address`.
    If POA parts are missing but `address` was already set (e.g. Details sheet), keep it.
    When ``preserve_existing_address`` is True, do not replace a non-empty existing ``address``
    (Details sheet full line) with a shorter UIDAI-composed line.
    """
    parts = [
        customer.get("care_of"),
        customer.get("house"),
        customer.get("street"),
        customer.get("location"),
        customer.get("city"),
        customer.get("district"),
        customer.get("sub_district"),
        customer.get("post_office"),
        customer.get("state"),
        customer.get("pin_code"),
    ]
    address = ", ".join(p for p in parts if p and str(p).strip())
    if address:
        prev = customer.get("address")
        if preserve_existing_address and prev and str(prev).strip():
            return
        customer["address"] = address
    elif not (customer.get("address") and str(customer["address"]).strip()):
        customer.pop("address", None)


def _merge_qr_customer_into_existing(
    existing: dict[str, str],
    qr_merged: dict[str, str],
) -> dict[str, str]:
    """
    Fill blank customer keys from an Aadhaar-derived fragment (Textract-parsed fields).
    Preserves Details sheet / DB values when set.
    """
    out = dict(existing) if existing else {}
    if not qr_merged:
        return out
    for k in AADHAR_QR_CUSTOMER_KEYS:
        qv = qr_merged.get(k)
        if not qv or not str(qv).strip():
            continue
        ev = out.get(k)
        if ev is None or not str(ev).strip():
            out[k] = str(qv).strip()
    # Fill address from POA parts when customer had no address; keep Details sheet address otherwise
    _rebuild_aadhar_address_from_parts(out, preserve_existing_address=True)
    if qr_merged.get("address") and str(qr_merged["address"]).strip():
        if not (out.get("address") and str(out["address"]).strip()):
            out["address"] = str(qr_merged["address"]).strip()
    return out


def _load_aadhar_scan_bytes(subdir: Path) -> dict[str, Any]:
    """Load ``Aadhar.jpg`` / ``Aadhar_back.jpg`` bytes for the Textract pipeline (no QR decode)."""
    out: dict[str, Any] = {"front_bytes": None, "back_bytes": None}
    ap = subdir / "Aadhar.jpg"
    bp = subdir / "Aadhar_back.jpg"
    try:
        out["front_bytes"] = ap.read_bytes() if ap.is_file() else None
    except OSError:
        out["front_bytes"] = None
    try:
        out["back_bytes"] = bp.read_bytes() if bp.is_file() else None
    except OSError:
        out["back_bytes"] = None
    return out


def _normalize_aadhar_back_address_chunk(chunk: str) -> str:
    chunk = chunk.replace("\r\n", "\n")
    chunk = re.sub(r"[\|`]+", " ", chunk)
    chunk = re.sub(r"\s+", " ", chunk).strip(" ,.-|")
    return chunk


def _parse_aadhar_back_address_from_ocr(ocr_text: str) -> dict[str, str]:
    """
    English (and noisy OCR) address on Aadhaar back. UIDAI layout uses ``Address:``;
    OCR often merges lines — fall back to ``C/O:`` … through PIN / Aadhaar line.
    DOB/Gender are not printed on the back; use **Aadhar.jpg** (front) Textract text.
    """
    out: dict[str, str] = {}
    if not ocr_text or len(ocr_text.strip()) < 8:
        return out
    text = ocr_text.replace("\r\n", "\n")
    raw = ""
    # Do not stop at ``C/O:`` on the next line — UIDAI layout is ``Address:`` then ``C/O:`` on the following line.
    m = re.search(
        r"(?is)\bAddress\s*:\s*([\s\S]+?)(?=\n\s*पता\b|\n\s*VID\s*:|\n\s*Virtual\s+ID\b|\Z)",
        text,
    )
    if m:
        raw = m.group(1).strip()
    if not raw or len(raw) < 12:
        co = re.search(r"(?is)\bC/O\s*:", text)
        if co:
            tail = text[co.start() :]
            stop = re.search(r"\d{4}\s+\d{4}\s+\d{4}", tail)
            chunk = tail[: stop.start()] if stop else tail
            stop_vid = re.search(r"(?i)\bVID\s*:", chunk)
            if stop_vid:
                chunk = chunk[: stop_vid.start()]
            raw = chunk.strip()
    if not raw or len(raw) < 8:
        # Textract often puts "Address" on its own line, then "Near …" / locality on the next lines.
        addr_head = re.search(r"(?is)\bAddress\s*\n+", text)
        if addr_head:
            window = text[addr_head.end() : addr_head.end() + 950]
            block_m = re.search(
                r"(?is)^\s*((?:Near|C/O|C\.?\s*O\.?|S/O|W/O|D/O)\s*:?.+?)(?=\n\s*(?:पता|VID|Virtual|Aadhaar|www\.|help@|\d{4}\s+\d{4}\s+\d{4})\b|\n{3,}|\Z)",
                window,
            )
            if block_m:
                raw = block_m.group(1).strip()
        if (not raw or len(raw) < 8) and re.search(r"(?i)\bNear\b", text):
            near_m = re.search(
                r"(?is)\b(Near\s+.{15,350}?(?:Rajasthan|State|,?\s*[A-Za-z]{4,}\s*-\s*\d{6}|\d{6}))",
                text,
            )
            if near_m:
                raw = near_m.group(1).strip()
    if not raw:
        return out
    raw = _normalize_aadhar_back_address_chunk(raw)
    foot = re.split(
        r"(?i)(www\.uidai|help@uidai|unique\s+identification|virtual\s+id)",
        raw,
        maxsplit=1,
    )
    raw = foot[0].strip(" ,.-") if foot else raw
    if len(raw) < 15:
        return out
    parsed = normalize_address_freeform(raw)
    out["address"] = parsed.get("address") or re.sub(r"\s+", " ", raw).strip()
    if parsed.get("care_of"):
        out["care_of"] = parsed["care_of"]
    if parsed.get("pin_code"):
        out["pin_code"] = parsed["pin_code"]
    if parsed.get("state"):
        out["state"] = parsed["state"]
    if parsed.get("city"):
        out["city"] = parsed["city"]
    if parsed.get("district"):
        out["district"] = parsed["district"]
    return out


def _aadhar_normalize_dob_triplet(day: int, month: int, year: int) -> str | None:
    """Validate and return DD/MM/YYYY for Indian Aadhaar-style dates."""
    if not (1 <= month <= 12 and 1 <= day <= 31):
        return None
    if year < 100:
        year = 2000 + year if year < 50 else 1900 + year
    y_max = datetime.now().year
    if year < 1920 or year > y_max:
        return None
    try:
        datetime(year, month, day)
    except ValueError:
        return None
    return f"{day:02d}/{month:02d}/{year}"


def _normalize_aadhar_gender_token(token: str) -> str | None:
    t = (token or "").strip()
    if not t:
        return None
    u = t.upper()
    if u in ("M", "MALE"):
        return "Male"
    if u in ("F", "FEMALE"):
        return "Female"
    if u in ("T", "TRANSGENDER"):
        return "Transgender"
    if t.lower() in ("male", "female", "transgender"):
        return t[:1].upper() + t[1:].lower()
    return None


def _aadhar_dob_window_not_issue_date(text: str, date_start_idx: int) -> bool:
    """True if the date at ``date_start_idx`` is unlikely to be ``Aadhaar … issued:``."""
    window = text[max(0, date_start_idx - 35) : date_start_idx]
    return not re.search(r"(?i)issued|issue\s*date|date\s+of\s+issue", window)


def _gender_skip_word_then_slash_token(suffix: str) -> str | None:
    """
    UIDAI front layout (OCR): after ``dd/mm/yyyy``, skip the next token, then the next ``/``
    introduces gender (e.g. ``yes/ MALE`` where ``Sex`` OCR'd as ``yes``).
    """
    s = suffix.lstrip()
    if not s:
        return None
    if s.startswith("/"):
        after = s[1:].lstrip()
        m_g = re.match(r"\S+", after)
        return _normalize_aadhar_gender_token(m_g.group(0)) if m_g else None
    m_skip = re.match(r"\S+", s)
    if m_skip:
        s = s[m_skip.end() :]
    s = s.lstrip()
    idx = s.find("/")
    if idx < 0:
        return None
    after = s[idx + 1 :].lstrip()
    m_g = re.match(r"\S+", after)
    if not m_g:
        return None
    return _normalize_aadhar_gender_token(m_g.group(0))


def _extract_gender_using_dob_slash_rule(full_text: str, normalized_dob: str | None) -> str | None:
    """
    Locate DOB (``dd/mm/yyyy``), skip the next word, find the next ``/``, read gender after it.
    Avoids anchoring on ``issued: dd/mm/yyyy`` when possible.
    """
    t = full_text.replace("\r\n", "\n")
    end_pos: int | None = None
    if normalized_dob and re.match(r"^\d{1,2}/\d{1,2}/\d{4}$", normalized_dob.strip()):
        p = normalized_dob.strip().split("/")
        d, mo, yr = int(p[0]), int(p[1]), p[2]
        variants = (
            rf"(?<!\d){d}/{mo}/{yr}(?!\d)",
            rf"(?<!\d){d:02d}/{mo:02d}/{yr}(?!\d)",
            rf"(?<!\d){d}/{mo:02d}/{yr}(?!\d)",
            rf"(?<!\d){d:02d}/{mo}/{yr}(?!\d)",
        )
        for rx in variants:
            for m in re.finditer(rx, t):
                if _aadhar_dob_window_not_issue_date(t, m.start()):
                    end_pos = m.end()
                    break
            if end_pos is not None:
                break
    if end_pos is None:
        for m in re.finditer(
            r"(?<!\d)(\d{1,2})[/.\-](\d{1,2})[/.\-]((?:19|20)\d{2})(?!\d)",
            t,
        ):
            if not _aadhar_dob_window_not_issue_date(t, m.start()):
                continue
            day, month, yr_s = int(m.group(1)), int(m.group(2)), int(m.group(3))
            norm = _aadhar_normalize_dob_triplet(day, month, yr_s)
            if norm:
                end_pos = m.end()
                break
    if end_pos is None:
        return None
    return _gender_skip_word_then_slash_token(t[end_pos:])


def _parse_aadhar_front_textract_fallback(text: str) -> dict[str, str]:
    """
    Pull DOB / gender from Aadhaar **front** full text (AWS Textract or same layout in Raw_OCR).
    Used when printed front text omits these fields.
    """
    out: dict[str, str] = {}
    if not text or len(text.strip()) < 5:
        return out
    t = text.replace("\r\n", "\n")

    dob_patterns = [
        r"(?i)\b(?:DOB|D\.?\s*O\.?\s*B\.?|Date\s+of\s+Birth|Birth\s+Date)\s*[:]?\s*(\d{1,2})[/.\-](\d{1,2})[/.\-](\d{2,4})\b",
        r"(?i)\b(?:जन्म\s*तिथि|जन्मतिथि)\s*[:]?\s*(\d{1,2})[/.\-](\d{1,2})[/.\-](\d{2,4})\b",
        # Mis-OCR: "JoH /DB:01/01/2004" — DOB read as DB; slash before DB is common.
        r"(?i)/\s*D\.?B\.?\s*[:]?\s*(\d{1,2})[/.\-](\d{1,2})[/.\-](\d{2,4})\b",
        # Standalone "DB:" or "D.B:" when followed immediately by a slash-date (avoid matching "db" in words).
        r"(?i)(?<![A-Za-z])D\.?B\.?\s*[:]?\s*(\d{1,2})[/.\-](\d{1,2})[/.\-](\d{2,4})\b",
    ]
    for pat in dob_patterns:
        m = re.search(pat, t)
        if m:
            d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
            norm = _aadhar_normalize_dob_triplet(d, mo, y)
            if norm:
                out["date_of_birth"] = norm
                out["year_of_birth"] = norm.split("/")[2]
                break

    if "date_of_birth" not in out:
        # Prefer dd/mm/yyyy (two slashes) near DOB/DB markers over stray numeric runs.
        dob_marker = re.compile(r"(?i)(?:\b(?:dob|date\s+of\s+birth|birth\s+date|जन्म)|/(?:dob|d\.?b\.?)\s*:?|\bd\.?b\.?\s*:)")
        triplet_re = re.compile(r"(?<!\d)(\d{1,2})/(\d{1,2})/((?:19|20)\d{2})(?!\d)")
        best: tuple[int, tuple[int, int, int]] | None = None
        for m in triplet_re.finditer(t):
            d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
            norm = _aadhar_normalize_dob_triplet(d, mo, y)
            if not norm:
                continue
            start, end = m.start(), m.end()
            window_start = max(0, start - 48)
            window = t[window_start:end]
            if not dob_marker.search(window):
                continue
            slash_bonus = window.count("/")
            dist = start - window_start
            score = slash_bonus * 10 - dist
            if best is None or score > best[0]:
                best = (score, (d, mo, y))
        if best:
            d, mo, y = best[1]
            norm = _aadhar_normalize_dob_triplet(d, mo, y)
            if norm:
                out["date_of_birth"] = norm
                out["year_of_birth"] = str(y)

    if "date_of_birth" not in out:
        for line in t.splitlines():
            line = line.strip()
            if len(line) > 72:
                continue
            if not re.search(r"(?i)(birth|dob|db\s*:|/(?:dob|d\.?b\.?)|जन्म)", line):
                continue
            m = re.search(r"(\d{1,2})[/.\-](\d{1,2})[/.\-]((19|20)\d{2})\b", line)
            if m:
                d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
                norm = _aadhar_normalize_dob_triplet(d, mo, y)
                if norm:
                    out["date_of_birth"] = norm
                    out["year_of_birth"] = str(y)
                    break

    # Primary gender from layout: anchor on DOB, skip next token, next "/" then gender token.
    g_dob = _extract_gender_using_dob_slash_rule(t, out.get("date_of_birth"))
    if g_dob:
        out["gender"] = g_dob

    if "gender" not in out:
        gm = (
            re.search(r"(?i)\bGender\s*[:]?\s*(Male|Female|Transgender|M|F|T)\b", t)
            or re.search(r"(?i)\bGender\s+(Male|Female|Transgender)\b", t)
            or re.search(r"(?i)\b(?:लिंग|Gender)\s*[:]?\s*(Male|Female|Transgender)\b", t)
        )
        if gm:
            g = _normalize_aadhar_gender_token(gm.group(1))
            if g:
                out["gender"] = g
        else:
            for label_pat in (r"(?i)\bGender\b", r"(?i)\bलिंग\b"):
                pos = re.search(label_pat, t)
                if pos:
                    window = t[pos.end() : pos.end() + 40]
                    for word in ("Male", "Female", "Transgender"):
                        if re.search(rf"(?i)\b{word}\b", window):
                            out["gender"] = word
                            break
                    if out.get("gender"):
                        break

    if "gender" not in out:
        # UIDAI front prints "Sex / Male" (or similar). OCR often reads "Sex" as "yes" → "yes/ MALE"
        # with no "Gender" label, so the patterns above miss it.
        slash_g = re.search(
            r"(?i)\b(?:yes|yex|yos|ses|sex)\s*/\s*(Male|Female|Transgender|MALE|FEMALE|M|F|T)\b",
            t,
        ) or re.search(
            r"(?i)\bSex\s*[/:]\s*(Male|Female|Transgender|MALE|FEMALE|M|F|T)\b",
            t,
        )
        if slash_g:
            g = _normalize_aadhar_gender_token(slash_g.group(1))
            if g:
                out["gender"] = g

    return out


def _parse_aadhar_name_from_aadhaar_textract(text: str) -> dict[str, str]:
    """
    Heuristic name line on Aadhaar **front** when Textract text has no clear name field.
    Avoids obvious headers / labels; prefers a Title Case line before DOB.
    """
    out: dict[str, str] = {}
    if not text or len(text.strip()) < 8:
        return out
    lines = [ln.strip() for ln in text.replace("\r\n", "\n").split("\n") if ln.strip()]
    skip = re.compile(
        r"(?i)aadhaar|aadhar|enrol|enrollment|government|india|identification|male|female|transgender|"
        r"address|dob|date\s*of\s*birth|year\s*of\s*birth|vid|virtual|help@|uidai|www\.|pin"
    )
    dob_line = re.compile(r"(?i)\b(dob|date\s*of\s*birth|जन्म|birth)\b")
    gov_line = re.compile(r"(?i)government\s+of\s+india")
    noise = re.compile(r"(?i)\b(famoy|family|service|assess|authority|unique)\b")
    candidates: list[str] = []
    gov_idx = -1
    for i, line in enumerate(lines[:25]):
        if gov_line.search(line):
            gov_idx = i
            break
    if gov_idx >= 0:
        for line in lines[gov_idx + 1 : min(gov_idx + 5, len(lines))]:
            if len(line) < 2 or len(line) > 80:
                continue
            if skip.search(line) or noise.search(line):
                continue
            if re.search(r"\d", line):
                continue
            if re.match(r"^[A-Za-z][A-Za-z\s.'-]{1,70}$", line):
                words = [w for w in line.split() if w.strip()]
                if 1 <= len(words) <= 5:
                    out["name"] = line.strip()
                    return out

    for line in lines[:22]:
        if len(line) < 4 or len(line) > 90:
            continue
        if skip.search(line):
            continue
        if noise.search(line):
            continue
        if dob_line.search(line):
            break
        if re.search(r"\d{4}\s+\d{4}\s+\d{4}", line):
            continue
        if re.match(r"^[\d\s/:\-.]+$", line):
            continue
        if re.match(r"^[A-Za-z][A-Za-z\s.'-]{2,70}$", line):
            words = line.split()
            initials = sum(1 for w in words if w and w[0].isupper())
            if 1 <= len(words) <= 6 and initials >= 1:
                # Keep candidates and prefer the nearest one before DOB.
                candidates.append(line)
    if candidates:
        out["name"] = candidates[-1]
    return out


def _aadhar_identity_ok(customer: dict[str, str]) -> bool:
    name = (customer.get("name") or "").strip()
    if len(name) >= 2:
        return True
    aid = "".join(c for c in str(customer.get("aadhar_id") or "") if c.isdigit())
    return len(aid) >= 4


def _aadhar_geo_ok(customer: dict[str, str]) -> bool:
    pin = re.sub(r"\D", "", str(customer.get("pin_code") or ""))
    if len(pin) >= 6:
        return True
    st = (customer.get("state") or "").strip()
    if len(st) >= 3:
        return True
    addr = (customer.get("address") or "").strip()
    if len(addr) >= 22:
        return True
    if (customer.get("district") or customer.get("city")) and len(pin) >= 6:
        return True
    return False


def _pipeline_merge_aadhar_customer(
    front_bytes: bytes | None,
    back_bytes: bytes | None,
    front_textract: dict | None,
    back_textract: dict | None,
) -> tuple[dict[str, str], list[tuple[str, str]], str | None, dict[str, int]]:
    """
    **AWS Textract only** on front/back (no Tesseract, no UIDAI QR in this pipeline).
    """
    from app.services.textract_service import extract_text_from_bytes

    timings: dict[str, int] = {
        "aadhar_textract_front_ms": 0,
        "aadhar_textract_back_ms": 0,
    }

    def _ftext(d: dict | None) -> str:
        if not d or d.get("error"):
            return ""
        return (d.get("full_text") or "").strip()

    raw_parts: list[tuple[str, str]] = []
    customer: dict[str, str] = {}

    ft = front_textract
    if front_bytes and (not ft or ft.get("error") or not _ftext(ft)):
        t_tx = time.perf_counter()
        try:
            ft = extract_text_from_bytes(front_bytes)
        except Exception:
            ft = {"error": "textract", "full_text": ""}
        timings["aadhar_textract_front_ms"] += int((time.perf_counter() - t_tx) * 1000)
    front_txt = _ftext(ft)
    if front_txt:
        raw_parts.append(("Aadhar.jpg", front_txt))
        customer = _merge_aadhar_textract_fallback_dict(
            customer, _parse_aadhar_front_textract_fallback(front_txt)
        )
        customer = _merge_aadhar_textract_fallback_dict(
            customer, _parse_aadhar_name_from_aadhaar_textract(front_txt)
        )

    bt = back_textract
    if back_bytes and not _aadhar_geo_ok(customer):
        if not bt or bt.get("error") or not _ftext(bt):
            t_tx = time.perf_counter()
            try:
                bt = extract_text_from_bytes(back_bytes)
            except Exception:
                bt = {"error": "textract", "full_text": ""}
            timings["aadhar_textract_back_ms"] += int((time.perf_counter() - t_tx) * 1000)
        back_txt = _ftext(bt)
        if back_txt:
            raw_parts.append(("Aadhar_back.jpg", back_txt))
            customer = _merge_aadhar_textract_fallback_dict(
                customer, _parse_aadhar_back_address_from_ocr(back_txt)
            )

    note: str | None = None
    if not _aadhar_identity_ok(customer):
        note = (
            "Aadhaar automated read did not yield a clear name or ID number. "
            "Enter or verify customer fields manually."
        )

    if customer.get("aadhar_id"):
        customer["aadhar_id"] = _aadhar_last4(customer["aadhar_id"]) or ""
    _default_gender_male_if_unread(customer)
    customer = enrich_customer_address_from_freeform(customer)
    return customer, raw_parts, note, timings


def _merge_granular_upload_process_timings(
    section_timings_ms: dict[str, int],
    *,
    prefetch_job_ms: dict[str, int],
    frag_a: dict[str, Any] | None,
    frag_d: dict[str, Any] | None,
) -> None:
    """Populate per-process durations (prefetch job wall times + sync pipeline/FORMS times)."""
    pt = (frag_a or {}).get("process_timings") or {}
    section_timings_ms["aadhar_textract_front_ms"] = int(prefetch_job_ms.get("aadhar_front", 0)) + int(
        pt.get("aadhar_textract_front_ms", 0)
    )
    section_timings_ms["aadhar_textract_back_ms"] = int(prefetch_job_ms.get("aadhar_back", 0)) + int(
        pt.get("aadhar_textract_back_ms", 0)
    )
    dsync = (
        int((frag_d or {}).get("detail_sheet_textract_sync_ms", 0))
        if frag_d and frag_d.get("ok")
        else 0
    )
    section_timings_ms["detail_sheet_textract_ms"] = int(prefetch_job_ms.get("details_forms", 0)) + dsync


def _compile_details_sheet_fragment(
    input_path: Path,
    textract_forms_prefetch: dict | None,
) -> dict:
    """
    Parse Sales Detail sheet to structured fragments only (no JSON write).
    Raster/PDF: Textract **AnalyzeDocument** FORMS+TABLES — right tool for printed form key-values.
    """
    fmt = _details_input_format(input_path)
    detail_sheet_textract_sync_ms = 0
    if fmt == "docx":
        key_value_pairs, docx_full = _key_value_pairs_from_docx(input_path)
        if not key_value_pairs and not (docx_full or "").strip():
            raise RuntimeError(
                "Could not read the Word document (.docx). Ensure it is a valid .docx Sales Detail Sheet."
            )
        result = {"error": None, "full_text": docx_full, "key_value_pairs": key_value_pairs}
    elif fmt in ("jpeg", "png", "pdf"):
        from app.services.textract_service import extract_forms_from_bytes

        if textract_forms_prefetch is not None:
            result = textract_forms_prefetch
        else:
            t_df = time.perf_counter()
            result = extract_forms_from_bytes(input_path.read_bytes())
            detail_sheet_textract_sync_ms = int((time.perf_counter() - t_df) * 1000)
        if result.get("error"):
            raise RuntimeError(str(result.get("error")))
    else:
        raise RuntimeError(
            f"Unsupported Details file format (detected={fmt!r}). "
            "Use a JPEG/PNG scan, PDF export, or .docx Sales Detail Sheet."
        )

    key_value_pairs = result.get("key_value_pairs") or []
    vehicle = _map_key_value_pairs_to_vehicle(key_value_pairs)
    insurance = _map_key_value_pairs_to_insurance(key_value_pairs)
    details_customer = _map_key_value_pairs_to_details_customer(key_value_pairs)
    details_customer_name = _extract_details_customer_name(key_value_pairs)
    if result.get("full_text"):
        from_vehicle = _parse_vehicle_from_full_text(result["full_text"])
        for k, v in from_vehicle.items():
            if v and not vehicle.get(k):
                vehicle[k] = v
        from_full = _parse_insurance_from_full_text(result["full_text"])
        if from_full.get("customer_name") and not details_customer_name:
            details_customer_name = from_full["customer_name"]
        for k, v in from_full.items():
            if k == "customer_name":
                continue
            if v and not insurance.get(k):
                insurance[k] = v
    if details_customer.get("name") and not details_customer_name:
        details_customer_name = details_customer.get("name")

    return {
        "vehicle": vehicle,
        "insurance": insurance,
        "details_customer": details_customer,
        "details_customer_name": details_customer_name,
        "full_text": result.get("full_text") or "",
        "detail_sheet_textract_sync_ms": detail_sheet_textract_sync_ms,
    }


def _default_gender_male_if_unread(customer: dict[str, str]) -> None:
    """If gender was not extracted (Textract/parsed text), assume Male for DMS/insurance flows."""
    g = customer.get("gender")
    if g is None or not str(g).strip():
        customer["gender"] = "Male"


def _merge_aadhar_textract_fallback_dict(customer: dict[str, str], hints: dict[str, str]) -> dict[str, str]:
    """
    Fill blank customer fields from Textract-derived hints. For ``address``, prefer a longer /
    clearer Textract line over a short or empty prior value (common for back-of-card English).
    """
    out = dict(customer) if customer else {}
    if not hints:
        return out
    for k in AADHAR_QR_CUSTOMER_KEYS:
        if k == "address":
            continue
        hv = hints.get(k)
        if not hv or not str(hv).strip():
            continue
        cur = out.get(k)
        if cur is None or not str(cur).strip():
            out[k] = str(hv).strip()
    ha = (hints.get("address") or "").strip()
    ca = (out.get("address") or "").strip()
    if ha:
        if not ca:
            out["address"] = ha
        elif len(ha) > len(ca) + 12:
            out["address"] = ha
    _rebuild_aadhar_address_from_parts(out, preserve_existing_address=True)
    if ha and (not out.get("address") or not str(out["address"]).strip()):
        out["address"] = ha
    if hints.get("pin_code") and not (out.get("pin_code") and str(out["pin_code"]).strip()):
        out["pin_code"] = str(hints["pin_code"]).strip()
    if hints.get("care_of") and not (out.get("care_of") and str(out["care_of"]).strip()):
        out["care_of"] = str(hints["care_of"]).strip()
    return out


def _parse_raw_ocr_sections(content: str) -> dict[str, str]:
    """Split ``Raw_OCR.txt`` (``--- filename ---`` sections) into filename -> body."""
    sections: dict[str, str] = {}
    current: str | None = None
    buf: list[str] = []
    for line in content.splitlines():
        m = re.match(r"^---\s*(.+?)\s*---\s*$", line.strip())
        if m:
            if current is not None:
                sections[current] = "\n".join(buf).strip()
            current = m.group(1).strip()
            buf = []
        else:
            buf.append(line)
    if current is not None:
        sections[current] = "\n".join(buf).strip()
    return sections


def _apply_aadhar_textract_fallbacks_from_parts(
    ocr_output_dir: Path,
    subfolder: str,
    parts: list[tuple[str, str]],
) -> None:
    """Merge DOB/gender (front) and address (back) from Textract blobs into OCR_To_be_Used.json."""
    json_path = _json_output_path(ocr_output_dir, subfolder)
    if not json_path.is_file():
        return
    try:
        data = json.loads(json_path.read_text(encoding="utf-8"))
    except Exception:
        return
    customer = data.get("customer") or {}
    if not isinstance(customer, dict):
        customer = {}
    for fn, tx in parts:
        if not tx or not str(tx).strip():
            continue
        fl = fn.strip().replace("\\", "/").split("/")[-1].lower()
        if fl == "aadhar.jpg":
            customer = _merge_aadhar_textract_fallback_dict(
                customer, _parse_aadhar_front_textract_fallback(tx)
            )
        elif fl == "aadhar_back.jpg":
            customer = _merge_aadhar_textract_fallback_dict(
                customer, _parse_aadhar_back_address_from_ocr(tx)
            )
    if customer.get("aadhar_id"):
        customer["aadhar_id"] = _aadhar_last4(customer["aadhar_id"]) or ""
    _default_gender_male_if_unread(customer)
    customer = enrich_customer_address_from_freeform(customer)
    data["customer"] = customer
    try:
        json_path.write_text(json.dumps(data, indent=2), encoding="utf-8")
    except Exception:
        pass


def _apply_aadhar_textract_fallbacks_from_raw_ocr_file(
    ocr_output_dir: Path,
    subfolder: str,
) -> None:
    """Same as parts merge, using persisted Raw_OCR.txt (e.g. get_extracted_details)."""
    raw_path = _ocr_subfolder_path(ocr_output_dir, subfolder) / "Raw_OCR.txt"
    if not raw_path.is_file():
        return
    try:
        sections = _parse_raw_ocr_sections(raw_path.read_text(encoding="utf-8", errors="replace"))
    except Exception:
        return
    parts = [(k, v) for k, v in sections.items() if v]
    if parts:
        _apply_aadhar_textract_fallbacks_from_parts(ocr_output_dir, subfolder, parts)


# Map Textract form keys (normalized) to our vehicle detail fields.
_VEHICLE_KEY_ALIASES = {
    # Many dealer PDFs use "Chassis Number" / "Engine Number" (not "Chassis No.")
    "frame_no": [
        "frame no",
        "frame no.",
        "frame number",
        "chassis",
        "chassis no",
        "chassis no.",
        "chassis number",
        "vin",
        "vehicle identification number",
    ],
    "engine_no": ["engine no", "engine no.", "engine number", "engine"],
    "model_colour": ["model & colour", "model and colour", "model/colour", "model", "colour", "color"],
    "key_no": ["key no", "key no.", "key number", "key"],
    "battery_no": ["battery no", "battery no.", "battery number", "battery"],
}

# Map Textract form keys to insurance/nominee fields (from Details sheet).
_INSURANCE_KEY_ALIASES = {
    "profession": [
        "profession",
        "profession:",
        "occupation",
        "customer profession",
        "employment",
        "nature of occupation",
        "job",
        "customer occupation",
    ],
    "financier": [
        "financier",
        "financier name",
        "finance company",
        "bank",
        "financer",
        "financing bank",
        "name of financier",
        "bank / financier",
        "financier / bank",
        "finance bank",
        "hypothecation",
        "loan from",
        "financing institution",
    ],
    "insurer": [
        "insurer name (if needed)",
        "insurer name if needed",
        "insurer",
        "insurer name",
        "name of insurer",
        "insurance company",
        "insurance provider",
        # Do not use bare "company" — matches unrelated labels (e.g. finance company).
    ],
    "marital_status": ["marital status", "customer marital status", "married status"],
    "nominee_gender": ["nominee gender", "gender of nominee", "sex of nominee", "nominee sex"],
    "nominee_name": [
        "nominee name",
        "nominee name:",
        "nominee",
        "name of nominee",
        "name of the nominee",
        "nominee's name",
        "insurance nominee",
        "nominee details",
    ],
    "nominee_age": [
        "nominee age",
        "nominee age:",
        "age of nominee",
        "nominee age (years)",
        "age (years)",
    ],
    "nominee_relationship": [
        "nominee relationship",
        "nominee relationship:",
        "relationship with customer",
        "relation with proposer",
        "relation to insured",
        "nominee relation",
        "relationship",
        "relation",
    ],
}

# Map Textract form keys to customer name on Details sheet.
_DETAILS_CUSTOMER_NAME_ALIASES = [
    "customer name", "customer name:", "buyer name", "buyer's name",
    "buyer name:", "name of customer",
    "full name", "full name:", "customer full name",
]

_DETAILS_CUSTOMER_KEY_ALIASES = {
    "name": ["full name", "customer name", "name of customer", "buyer name", "name"],
    "mobile_number": ["mobile number", "mobile no", "mobile", "contact number", "contact no"],
    "alt_phone_num": ["alternate", "alternate no", "alternate number", "alternate mobile", "alternate mobile number", "landline", "landline number"],
    "aadhar_id": ["aadhaar number", "aadhar number", "aadhaar no", "aadhar no", "aadhaar", "aadhar"],
    "profession": [
        "profession",
        "occupation",
        "customer profession",
        "employment",
        "nature of occupation",
        "customer occupation",
    ],
    "marital_status": ["marital status", "customer marital status"],
    "financier": [
        "financier name",
        "financier",
        "financer name",
        "financer",
        "finance company",
        "bank",
        "financing bank",
        "name of financier",
        "bank / financier",
    ],
    "nominee_name": ["nominee name", "name of nominee", "name of the nominee", "nominee's name"],
    "nominee_age": ["nominee age", "age of nominee", "age (years)"],
    "nominee_gender": ["nominee gender", "gender of nominee", "sex of nominee"],
    "nominee_relationship": [
        "nominee relationship",
        "nominee relation",
        "relationship with customer",
        "relation with proposer",
        "relationship",
        "relation",
    ],
}

# Map Textract form keys for insurance policy document (Insurance.jpg).
_INSURANCE_POLICY_KEY_ALIASES = {
    "insurer": ["insurer", "insurance company", "company", "insurance provider", "name of insurer", "national insurance"],
    "policy_num": ["policy no", "policy no.", "policy number", "policy num", "cert. no", "cert no", "certificate no"],
    "policy_from": ["tp valid from", "valid from", "validity from", "policy from", "from date", "tp valid"],
    "policy_to": ["tp valid to", "valid to", "validity to", "policy to", "to date", "midnight of"],
    "premium": ["gross premium", "premium", "total premium", "gross premium amount", "premium of rs"],
}


def _normalize_key_for_match(key: str) -> str:
    """Lowercase and collapse spaces for key matching."""
    return re.sub(r"\s+", " ", (key or "").lower().strip())


def _details_input_format(path: Path) -> str:
    """
    Detect real file type for the sales detail upload. V2 always saves as Details.jpg
    even when the user uploads .docx or PDF, so we must sniff magic bytes.
    """
    try:
        head = path.read_bytes()[:12]
    except OSError:
        return "unknown"
    if len(head) >= 2 and head[:2] == b"\xff\xd8":
        return "jpeg"
    if len(head) >= 4 and head[:4] == b"\x89PNG":
        return "png"
    if len(head) >= 4 and head[:4] == b"%PDF":
        return "pdf"
    if len(head) >= 2 and head[:2] == b"PK":
        return "docx"
    return "unknown"


def _parallel_textract_prefetch_upload_subfolder(subdir: Path) -> tuple[dict[str, dict], dict[str, int]]:
    """
    Run AWS Textract calls for all upload files concurrently to cut wall time on scans-v2.
    Keys: ``aadhar_front`` / ``aadhar_back`` (DetectDocumentText), ``details_forms``
    (**AnalyzeDocument** FORMS+TABLES for structured sales detail scans/PDFs), ``insurance``,
    ``financing``. ``.docx`` Details is not prefetched (parsed locally).

    Returns ``(results_by_key, job_duration_ms_by_key)`` for logging and ``section_timings_ms``.
    """
    from concurrent.futures import ThreadPoolExecutor

    from app.config import OCR_UPLOAD_TEXTRACT_TIMEOUT_SEC
    from app.services.textract_service import extract_forms_from_bytes, extract_text_from_bytes

    jobs: list[tuple[str, bytes, str]] = []
    ap = subdir / "Aadhar.jpg"
    if ap.is_file():
        try:
            jobs.append(("aadhar_front", ap.read_bytes(), "text"))
        except OSError as e:
            logger.warning("prefetch: could not read Aadhar.jpg: %s", e)
    dp = subdir / "Details.jpg"
    if dp.is_file() and _details_input_format(dp) in ("jpeg", "png", "pdf"):
        try:
            jobs.append(("details_forms", dp.read_bytes(), "forms"))
        except OSError as e:
            logger.warning("prefetch: could not read Details.jpg: %s", e)
    ip = subdir / "Insurance.jpg"
    if ip.is_file():
        try:
            jobs.append(("insurance", ip.read_bytes(), "text"))
        except OSError as e:
            logger.warning("prefetch: could not read Insurance.jpg: %s", e)
    for fname, key in (("Aadhar_back.jpg", "aadhar_back"), ("Financing.jpg", "financing")):
        p = subdir / fname
        if p.is_file():
            try:
                jobs.append((key, p.read_bytes(), "text"))
            except OSError as e:
                logger.warning("prefetch: could not read %s: %s", fname, e)

    out: dict[str, dict] = {}
    job_ms: dict[str, int] = {}
    if not jobs:
        return out, job_ms

    timeout = max(30, OCR_UPLOAD_TEXTRACT_TIMEOUT_SEC)
    max_workers = min(5, len(jobs))

    def _run(job: tuple[str, bytes, str]) -> tuple[str, dict, int]:
        key, blob, mode = job
        t0 = time.perf_counter()
        try:
            if mode == "forms":
                res = extract_forms_from_bytes(blob)
            else:
                res = extract_text_from_bytes(blob)
        except Exception as e:
            res = {
                "error": str(e),
                "full_text": "",
                "key_value_pairs": [],
                "blocks": [],
                "raw_response": None,
            }
        elapsed_ms = int((time.perf_counter() - t0) * 1000)
        return key, res, elapsed_ms

    with ThreadPoolExecutor(max_workers=max_workers) as ex:
        keyed_futures = []
        for job in jobs:
            keyed_futures.append((job[0], ex.submit(_run, job)))
        for key, fut in keyed_futures:
            try:
                k, res, elapsed_ms = fut.result(timeout=timeout)
                out[k] = res
                job_ms[k] = elapsed_ms
            except Exception as e:
                logger.warning("prefetch Textract failed for %s: %s", key, e)
                out[key] = {
                    "error": str(e),
                    "full_text": "",
                    "key_value_pairs": [],
                    "blocks": [],
                    "raw_response": None,
                }
                job_ms[key] = 0
    return out, job_ms


def _key_value_pairs_from_docx(doc_path: Path) -> tuple[list[dict[str, str]], str]:
    """
    Parse Word Sales Detail Sheet (.docx) into Textract-like key/value pairs.
    AWS Textract does not accept .docx bytes; this path is used for native Word uploads.
    """
    try:
        from docx import Document
    except ImportError:
        return [], ""

    doc = Document(str(doc_path))
    pairs: list[dict[str, str]] = []
    text_lines: list[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        text_lines.append(t)
        if ":" in t:
            left, right = t.split(":", 1)
            lk, rv = left.strip(), right.strip()
            if lk and rv and len(lk) <= 160 and len(rv) <= 600:
                pairs.append({"key": lk, "value": rv})
        elif "\t" in t:
            parts = re.split(r"\t+", t, maxsplit=1)
            if len(parts) == 2 and parts[0].strip() and parts[1].strip():
                pairs.append({"key": parts[0].strip(), "value": parts[1].strip()})

    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if not any(cells):
                continue
            n = len(cells)
            if n >= 2:
                for i in range(0, n - 1, 2):
                    key, val = cells[i], cells[i + 1]
                    if key and val and key.casefold() != val.casefold():
                        pairs.append({"key": key, "value": val})

    seen: set[tuple[str, str]] = set()
    deduped: list[dict[str, str]] = []
    for kv in pairs:
        k = (kv.get("key") or "").strip()
        v = (kv.get("value") or "").strip()
        if not k or not v:
            continue
        sig = (k.casefold(), v.casefold())
        if sig in seen:
            continue
        seen.add(sig)
        deduped.append({"key": k, "value": v})

    full_text = "\n".join(text_lines)
    for table in doc.tables:
        for row in table.rows:
            row_txt = " | ".join((c.text or "").strip() for c in row.cells)
            if row_txt.strip():
                full_text += "\n" + row_txt
    return deduped, full_text


def _sanitize_nominee_age(val: str | None) -> str | None:
    """Extract numeric part from nominee age (e.g. '30/m' -> '30'). Returns None if no digits."""
    if not val or not str(val).strip():
        return None
    s = str(val).strip()
    m = re.match(r"^(\d{1,3})", s)
    return m.group(1) if m and 1 <= int(m.group(1)) <= 150 else None


def _normalize_name_for_match(name: str | None) -> str:
    """Normalize name for matching: lowercase, strip, collapse spaces."""
    if not name or not str(name).strip():
        return ""
    return " ".join(str(name).lower().strip().split())


def _names_match(name1: str | None, name2: str | None) -> bool:
    """Return True if names likely refer to same person. Handles OCR variations (spacing, case)."""
    n1 = _normalize_name_for_match(name1)
    n2 = _normalize_name_for_match(name2)
    if not n1 or not n2:
        return True  # Can't compare if one missing; allow
    if n1 == n2:
        return True
    # One name contains the other (e.g. "john doe" in "john doe kumar")
    if n1 in n2 or n2 in n1:
        return True
    # Stricter: no first-word-only fallback; require substantial match
    return False


def _validate_name_match(aadhar_name: str | None, details_name: str | None, insurance_name: str | None) -> str | None:
    """Return error message if names from Aadhar, Details and Insurance do not match. None if OK."""
    names = [
        ("Aadhar", aadhar_name),
        ("Details sheet", details_name),
        ("Insurance", insurance_name),
    ]
    present = [(label, n) for label, n in names if n and str(n).strip()]
    if len(present) < 2:
        return None  # Need at least 2 to compare
    first_label, first_val = present[0]
    for label, val in present[1:]:
        if not _names_match(first_val, val):
            return (
                f"Name mismatch: '{first_label}' has a different name than '{label}'. "
                "Ensure the name on Aadhar front, Details sheet and Insurance document match."
            )
    return None


def _map_key_value_pairs_to_vehicle(pairs: list[dict]) -> dict[str, str]:
    """Map key_value_pairs from Textract to structured vehicle fields (frame_no, engine_no, model_colour, key_no, battery_no)."""
    out: dict[str, str] = {}
    key_norm_to_value: dict[str, str] = {}
    for kv in pairs:
        k = (kv.get("key") or "").strip()
        v = (kv.get("value") or "").strip()
        if not k:
            continue
        key_norm = _normalize_key_for_match(k)
        key_norm_to_value[key_norm] = v
        if ":" in key_norm:
            key_norm_to_value[key_norm.replace(":", "").strip()] = v

    for field, aliases in _VEHICLE_KEY_ALIASES.items():
        if field in out:
            continue
        for alias in aliases:
            anorm = _normalize_key_for_match(alias)
            if anorm in key_norm_to_value:
                out[field] = key_norm_to_value[anorm]
                break
            for kn, v in key_norm_to_value.items():
                if anorm in kn or kn in anorm:
                    out[field] = v
                    break
            if field in out:
                break

    # Combine Model and Colour into model_colour if we have them separately
    model_val = key_norm_to_value.get(_normalize_key_for_match("model"), "").strip() or next(
        (
            v.strip()
            for kn, v in key_norm_to_value.items()
            if "model" in kn and "colour" not in kn and "color" not in kn and v.strip()
        ),
        "",
    )
    colour_val = key_norm_to_value.get(_normalize_key_for_match("colour"), "").strip() or key_norm_to_value.get(
        _normalize_key_for_match("color"), ""
    ).strip()
    if model_val or colour_val:
        combined = ", ".join(filter(None, [model_val, colour_val]))
        if combined:
            out["model_colour"] = combined

    # Guard against section-heading bleed when Key/Battery fields are blank or scratched.
    for fld in ("key_no", "battery_no"):
        val = (out.get(fld) or "").strip()
        if val and re.search(r"(?i)\b(nominee|payment|insurance|details|customer|vehicle)\b", val):
            out.pop(fld, None)

    return out


def _clean_sales_sheet_scalar(value: str) -> str:
    """Strip placeholder underscores / blanks from a field value."""
    s = (value or "").strip()
    if not s:
        return ""
    if re.match(r"^[\s_.,-]+$", s):
        return ""
    return s.strip()


def _parse_vehicle_from_full_text(full_text: str) -> dict[str, str]:
    """
    Fallback when Textract FORMS misses pairs (common on some PDFs): parse LINE layout
    like 'Chassis Number: 59324 Engine Number: 50581'.
    """
    out: dict[str, str] = {}
    if not full_text or not isinstance(full_text, str):
        return out
    text = full_text.replace("\r", "\n")
    lines = [ln.strip() for ln in text.split("\n")]

    def _section_or_noise_line(s: str) -> bool:
        return bool(re.search(r"(?i)\b(nominee|payment|insurance|customer details|vehicle details)\b", s))

    def _best_numeric_token(s: str, *, min_len: int = 3, max_len: int = 12) -> str:
        toks = re.findall(r"\d+", s or "")
        toks = [t for t in toks if min_len <= len(t) <= max_len]
        if not toks:
            return ""
        # Prefer 4-6 length tokens for vehicle partials; else longest.
        pref = [t for t in toks if 4 <= len(t) <= 6]
        return (pref[0] if pref else sorted(toks, key=len, reverse=True)[0]).strip()

    def _extract_field_near_label(
        label_rx: str,
        *,
        min_len: int = 3,
        max_len: int = 12,
        look_ahead: int = 3,
    ) -> str:
        for idx, ln in enumerate(lines):
            if not re.search(label_rx, ln, re.I):
                continue
            # 1) Inline value after colon in same line.
            m_inline = re.search(r":\s*(.+)$", ln)
            if m_inline:
                cand_inline = _best_numeric_token(m_inline.group(1), min_len=min_len, max_len=max_len)
                if cand_inline:
                    return cand_inline
            # 2) Nearest numeric line below label (for scratched/blank field values).
            for j in range(idx + 1, min(len(lines), idx + 1 + look_ahead)):
                nln = lines[j]
                if not nln or _section_or_noise_line(nln):
                    continue
                # Skip another label line.
                if ":" in nln and re.search(r"[A-Za-z]", nln):
                    continue
                cand = _best_numeric_token(nln, min_len=min_len, max_len=max_len)
                if cand:
                    return cand
            # 3) Small backward window (number slightly above label).
            for j in range(max(0, idx - 2), idx):
                pln = lines[j]
                if not pln or _section_or_noise_line(pln):
                    continue
                if ":" in pln and re.search(r"[A-Za-z]", pln):
                    continue
                cand = _best_numeric_token(pln, min_len=min_len, max_len=max_len)
                if cand:
                    return cand
        return ""

    m = re.search(r"(?i)Model\s*:\s*([^\n]+?)(?=\s+Colour\s*:)", text)
    if m:
        mv = _clean_sales_sheet_scalar(m.group(1))
        if mv:
            out["model_colour"] = mv
    m = re.search(r"(?i)Colour\s*:\s*([^\n]+)", text)
    if m:
        cv = _clean_sales_sheet_scalar(m.group(1))
        if cv:
            if out.get("model_colour"):
                out["model_colour"] = f"{out['model_colour']}, {cv}"
            else:
                out["model_colour"] = cv

    m = re.search(r"(?i)Chassis\s+Number\s*:\s*([^\n]+?)(?=\s+Engine\s+Number\s*:|\n|$)", text)
    if m:
        fv = _clean_sales_sheet_scalar(m.group(1))
        if fv:
            out["frame_no"] = fv
    m = re.search(r"(?i)Engine\s+Number\s*:\s*([^\n]+?)(?=\s+Key\s+Number\s*:|\n|$)", text)
    if m:
        ev = _clean_sales_sheet_scalar(m.group(1))
        if ev:
            out["engine_no"] = ev
    m = re.search(r"(?i)Key\s+Number\s*:\s*([^\n]+?)(?=\s+Battery\s+Number\s*:|\n|$)", text)
    if m:
        kv = _clean_sales_sheet_scalar(m.group(1))
        if kv and not re.search(r"(?i)\b(nominee|payment|insurance|details|customer|vehicle)\b", kv):
            out["key_no"] = kv
    m = re.search(r"(?i)Battery\s+Number\s*:\s*([^\n]*)(?=\n|$)", text)
    if m:
        bv = _clean_sales_sheet_scalar(m.group(1))
        if bv and not re.search(r"(?i)\b(nominee|payment|insurance|details|customer|vehicle)\b", bv):
            out["battery_no"] = bv

    # Nearby-number rescue for scratched/blank fields in Details sheet.
    if "frame_no" not in out or not str(out.get("frame_no") or "").strip():
        v = _extract_field_near_label(r"\b(chassis|frame)\s+number\b", min_len=4, max_len=12)
        if v:
            out["frame_no"] = v
    if "engine_no" not in out or not str(out.get("engine_no") or "").strip():
        v = _extract_field_near_label(r"\bengine\s+number\b", min_len=4, max_len=12)
        if v:
            out["engine_no"] = v
    if "key_no" not in out or not str(out.get("key_no") or "").strip():
        v = _extract_field_near_label(r"\bkey\s+number\b", min_len=3, max_len=10)
        if v:
            out["key_no"] = v
    if "battery_no" not in out or not str(out.get("battery_no") or "").strip():
        v = _extract_field_near_label(r"\bbattery\s+number\b", min_len=3, max_len=12)
        if v:
            out["battery_no"] = v

    # Alternate labels
    if "frame_no" not in out:
        m = re.search(r"(?i)(?:Frame|Chassis)\s+No\.?\s*:\s*([^\n]+?)(?=\s+Engine|\n|$)", text)
        if m:
            fv = _clean_sales_sheet_scalar(m.group(1))
            if fv:
                out["frame_no"] = fv
    return out


def _extract_details_customer_name(pairs: list[dict]) -> str | None:
    """Extract customer name from Details sheet key-value pairs."""
    key_lower_to_value: dict[str, str] = {}
    for kv in pairs:
        k = (kv.get("key") or "").strip()
        v = (kv.get("value") or "").strip()
        if not k or not v:
            continue
        key_norm = _normalize_key_for_match(k)
        key_lower_to_value[key_norm] = v
        if ":" in key_norm:
            key_lower_to_value[key_norm.replace(":", "").strip()] = v
    for alias in _DETAILS_CUSTOMER_NAME_ALIASES:
        anorm = _normalize_key_for_match(alias)
        if anorm in key_lower_to_value:
            val = key_lower_to_value[anorm].strip()
            if len(val) >= 2 and len(val) <= 80 and not re.search(r"(?i)\bnominee\b", val):
                return val
        for k, v in key_lower_to_value.items():
            if re.search(r"(?i)\bnominee\b", k):
                continue
            if anorm in k or k in anorm:
                val = v.strip()
                if len(val) >= 2 and len(val) <= 80 and not re.search(r"(?i)\bnominee\b", val):
                    return val
    return None


def _extract_financier_from_payment_line(text: str) -> str | None:
    """Arya-style sales detail: 'Payment: (A) Cash ... (D) Finance: SHRIRAM FIN' (no separate Financier label)."""
    if not text or not isinstance(text, str):
        return None
    low = text.lower()
    if "finance" not in low:
        return None
    m = re.search(r"(?i)(?:\(D\)\s*|D\)\s*)Finance\s*:\s*([^\n]+)", text)
    if not m and "payment" in low and re.search(r"(?i)\(D\)", text):
        m = re.search(r"(?i)Finance\s*:\s*([^\n]+)", text)
    if not m:
        return None
    val = m.group(1).strip()
    for cut in (" I agree", " Insurer", " Insurer Name", " Customer Signature"):
        if cut in val:
            val = val.split(cut)[0].strip()
    val = val.rstrip(".,; ")
    if len(val) < 2 or len(val) > 160:
        return None
    if re.match(r"^(yes|no|na|n/?a|cash|required|optional)\s*$", val, re.I):
        return None
    return val


def _map_key_value_pairs_to_insurance(pairs: list[dict]) -> dict[str, str]:
    """Map key_value_pairs from Textract to insurance/nominee fields."""
    out: dict[str, str] = {}
    key_lower_to_value: dict[str, str] = {}
    for kv in pairs:
        k = (kv.get("key") or "").strip()
        v = (kv.get("value") or "").strip()
        if not k:
            continue
        key_norm = _normalize_key_for_match(k)
        key_lower_to_value[key_norm] = v
        # Also store with colon stripped for "Nominee Name:" style
        if ":" in key_norm:
            key_lower_to_value[key_norm.replace(":", "").strip()] = v

    for field, aliases in _INSURANCE_KEY_ALIASES.items():
        if field in out:
            continue
        for alias in aliases:
            anorm = _normalize_key_for_match(alias)
            if anorm in key_lower_to_value:
                out[field] = key_lower_to_value[anorm]
                break
            # Match if key contains alias (e.g. "nominee name" in "Nominee Name")
            for k, v in key_lower_to_value.items():
                if anorm in k or k in anorm:
                    out[field] = v
                    break
            if field in out:
                break

    if "financier" not in out:
        for vv in key_lower_to_value.values():
            got = _extract_financier_from_payment_line(vv)
            if got:
                out["financier"] = got
                break
    return out


def _map_key_value_pairs_to_details_customer(pairs: list[dict]) -> dict[str, str]:
    """Map Details sheet key-value pairs to customer-side fields from new A5 template."""
    out: dict[str, str] = {}
    key_lower_to_value: dict[str, str] = {}
    for kv in pairs:
        k = (kv.get("key") or "").strip()
        v = (kv.get("value") or "").strip()
        if not k:
            continue
        key_norm = _normalize_key_for_match(k)
        key_lower_to_value[key_norm] = v
        if ":" in key_norm:
            key_lower_to_value[key_norm.replace(":", "").strip()] = v

    for field, aliases in _DETAILS_CUSTOMER_KEY_ALIASES.items():
        if field in out:
            continue
        for alias in aliases:
            anorm = _normalize_key_for_match(alias)
            if anorm in key_lower_to_value and key_lower_to_value[anorm].strip():
                out[field] = key_lower_to_value[anorm].strip()
                break
            for k, v in key_lower_to_value.items():
                if (anorm in k or k in anorm) and v.strip():
                    out[field] = v.strip()
                    break
            if field in out:
                break

    if out.get("mobile_number"):
        digits = "".join(ch for ch in out["mobile_number"] if ch.isdigit())
        out["mobile_number"] = digits[-10:] if digits else ""
    if out.get("alt_phone_num"):
        digits = "".join(ch for ch in out["alt_phone_num"] if ch.isdigit())
        out["alt_phone_num"] = digits[-10:] if digits else ""
    if out.get("aadhar_id"):
        out["aadhar_id"] = _aadhar_last4(out.get("aadhar_id")) or ""
    if out.get("nominee_age"):
        out["nominee_age"] = _sanitize_nominee_age(out["nominee_age"]) or ""
    return out


def _parse_insurance_from_full_text(full_text: str) -> dict[str, str]:
    """Try to extract insurance/customer extras from full text (fallback when not in key-value pairs)."""
    out: dict[str, str] = {}
    if not full_text or not isinstance(full_text, str):
        return out
    text = full_text.strip()
    pay_fin = _extract_financier_from_payment_line(text)
    if pay_fin:
        out["financier"] = pay_fin
    # Patterns: "Label" or "Label:" followed by value on same or next line
    patterns = [
        # Detail-sheet lines e.g. "Insurer Name (if needed): SOMPO" (longer labels first)
        ("insurer name (if needed)", "insurer"),
        ("insurer name if needed", "insurer"),
        ("name of insurer", "insurer"),
        # Avoid bare "insurer name" here — it matches inside "Insurer Name (if needed)" and captures the wrong group.
        ("insurance provider", "insurer"),
        ("customer profession", "profession"),
        ("profession of customer", "profession"),
        ("profession", "profession"),
        ("occupation", "profession"),
        ("nature of occupation", "profession"),
        ("name of financier", "financier"),
        ("financing bank", "financier"),
        ("finance company", "financier"),
        ("financier / bank", "financier"),
        ("bank / financier", "financier"),
        ("financier", "financier"),
        ("marital status", "marital_status"),
        ("nominee gender", "nominee_gender"),
        ("sex of nominee", "nominee_gender"),
        ("name of the nominee", "nominee_name"),
        ("name of nominee", "nominee_name"),
        ("nominee name", "nominee_name"),
        ("nominee's name", "nominee_name"),
        ("nominee age", "nominee_age"),
        ("age of nominee", "nominee_age"),
        ("nominee relationship", "nominee_relationship"),
        ("relationship with customer", "nominee_relationship"),
        ("relation with proposer", "nominee_relationship"),
        ("customer name", "customer_name"),
        ("buyer name", "customer_name"),
        ("buyer's name", "customer_name"),
    ]
    for label, key in patterns:
        if key in out:
            continue
        pat = re.compile(
            rf"(?im){re.escape(label)}\s*[:\t]?\s*\n?\s*([^\n]+)",
        )
        m = pat.search(text)
        if m:
            val = m.group(1).strip()
            if val and len(val) < 200:
                out[key] = val

    # Word / table layouts: label alone on one line, value on the next (e.g. "Profession" then "Farmer")
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    for i, ln in enumerate(lines):
        low = _normalize_key_for_match(ln)
        nxt = lines[i + 1].strip() if i + 1 < len(lines) else ""
        if not nxt or len(nxt) > 180:
            continue
        if "profession" not in out and low in ("profession", "occupation", "customer profession", "employment"):
            out["profession"] = nxt
        if "financier" not in out and low in (
            "financier",
            "financing bank",
            "name of financier",
            "finance company",
            "bank name",
        ):
            out["financier"] = nxt
        if "nominee_name" not in out and ("nominee" in low and "name" in low) and "age" not in low:
            out["nominee_name"] = nxt
        if "nominee_age" not in out and low in ("nominee age", "age of nominee", "nominee age (years)"):
            out["nominee_age"] = _sanitize_nominee_age(nxt) or nxt
        if "nominee_gender" not in out and low in ("nominee gender", "gender of nominee", "sex of nominee"):
            out["nominee_gender"] = nxt
        if "nominee_relationship" not in out and low in (
            "nominee relationship",
            "relationship with customer",
            "relation with proposer",
            "relation to insured",
        ):
            out["nominee_relationship"] = nxt
        same = re.match(r"(?i)^(profession|occupation|customer profession)\s+(.{1,120})$", ln)
        if same and "profession" not in out:
            out["profession"] = same.group(2).strip()
        same_f = re.match(
            r"(?i)^(financier|financing bank|name of financier)\s+(.{1,160})$",
            ln,
        )
        if same_f and "financier" not in out:
            out["financier"] = same_f.group(2).strip()
        if "insurer" not in out and "insurer" in low and "name" in low and "nominee" not in low:
            if nxt and len(nxt) < 120 and not re.match(r"^[:\s_]+$", nxt):
                out["insurer"] = nxt.strip()

    return out


def _map_key_value_pairs_to_insurance_policy(pairs: list[dict]) -> dict[str, str]:
    """Map key_value_pairs from Textract to insurance policy fields (insurer, policy_from, policy_to, premium)."""
    out: dict[str, str] = {}
    key_lower_to_value: dict[str, str] = {}
    for kv in pairs:
        k = (kv.get("key") or "").strip()
        v = (kv.get("value") or "").strip()
        if not k:
            continue
        key_norm = _normalize_key_for_match(k)
        key_lower_to_value[key_norm] = v
        if ":" in key_norm:
            key_lower_to_value[key_norm.replace(":", "").strip()] = v

    for field, aliases in _INSURANCE_POLICY_KEY_ALIASES.items():
        if field in out:
            continue
        for alias in aliases:
            anorm = _normalize_key_for_match(alias)
            if anorm in key_lower_to_value:
                out[field] = key_lower_to_value[anorm]
                break
            for k, v in key_lower_to_value.items():
                if anorm in k or k in anorm:
                    out[field] = v
                    break
            if field in out:
                break
    return out


def _parse_insurance_policy_from_full_text(full_text: str) -> dict[str, str]:
    """Extract policy number, insurer, policy from/to, gross premium from insurance policy full text.
    Document layout:
      - Policy / Cert. No. -> policy number (e.g. 39010231216200202663)
      - Right after -> insurance provider (e.g. NATIONAL INSURANCE COMPANY LIMITED)
      - Policy Period -> two dates: first=from, second=to (e.g. 15-06-2022, 15-06-2023)
      - Gross Premium -> premium amount (e.g. 5291.00)
    """
    out: dict[str, str] = {}
    if not full_text or not isinstance(full_text, str):
        return out
    text = full_text.strip()

    # 1. Policy number: "Policy\nCert. No.\n39010231216200202663" - number after Cert. No.
    if "policy_num" not in out:
        m = re.search(
            r"(?:Policy\s+)?Cert\.?\s*No\.?\s*:?\s*\n\s*(\d{12,})",
            text,
            re.IGNORECASE,
        )
        if m:
            out["policy_num"] = m.group(1).strip()

    # 2. Insurer: right after policy number - "NATIONAL INSURANCE COMPANY LIMITED"
    if "insurer" not in out:
        m = re.search(
            r"(\d{12,})\s*\n\s*([A-Z][A-Za-z\s]+(?:INSURANCE\s+COMPANY\s+LIMITED|INSURANCE\s+COMPANY|INSURANCE))",
            text,
        )
        if m:
            out["insurer"] = m.group(2).strip()
        else:
            # Fallback: known insurer names
            m = re.search(
                r"([A-Z][A-Za-z\s]+INSURANCE[A-Za-z\s]*(?:COMPANY\s+)?(?:LIMITED|LTD\.?)?)",
                text,
            )
            if m:
                out["insurer"] = m.group(1).strip()

    # 3. Policy from & to:
    #    - policy_from: Date in "dd-mm-yyyy To" block before "OD Policy Period" (e.g. 16-06-2021 near UIN block)
    #    - policy_to: First dd-mm-yyyy after "OD Policy Period" (e.g. 15-06-2022)
    if "policy_from" not in out or "policy_to" not in out:
        date_pat = re.compile(r"\b(\d{1,2}-\d{1,2}-\d{4})\b")
        od_match = re.search(r"OD\s+Policy\s+Period\s*:?\s*\n", text, re.IGNORECASE)
        if od_match:
            before_od = text[: od_match.start()]
            after_od = text[od_match.end() :]
            if "policy_from" not in out:
                # Prefer date in "dd-mm-yyyy To" format (policy period start); else first date before OD
                m_from = re.search(r"(\d{1,2}-\d{1,2}-\d{4})\s+To", before_od)
                if m_from:
                    out["policy_from"] = m_from.group(1)
                else:
                    dates_before = date_pat.findall(before_od)
                    if dates_before:
                        out["policy_from"] = dates_before[0]
            if "policy_to" not in out:
                dates_after = date_pat.findall(after_od)
                if dates_after:
                    out["policy_to"] = dates_after[0]  # First date after OD = policy end (e.g. 15-06-2022)
        if "policy_from" not in out or "policy_to" not in out:
            # Fallback: Policy Period with two dates
            for label in [r"Policy\s+Period", r"OD\s+Policy\s+Period"]:
                m = re.search(
                    rf"{label}\s*:?\s*\n\s*(\d{{1,2}}-\d{{1,2}}-\d{{4}})\s*\n\s*(\d{{1,2}}-\d{{1,2}}-\d{{4}})",
                    text,
                    re.IGNORECASE,
                )
                if m:
                    if "policy_from" not in out:
                        out["policy_from"] = m.group(1)
                    if "policy_to" not in out:
                        out["policy_to"] = m.group(2)
                    break

    # 4. Gross Premium: "Gross Premium\n5291.00"
    if "premium" not in out:
        m = re.search(r"Gross\s+Premium\s*:?\s*\n\s*([\d,]+\.?\d*)", text, re.IGNORECASE)
        if m:
            out["premium"] = m.group(1).strip().replace(",", "")
        else:
            m = re.search(r"Premium\s+of\s+Rs\.?\s*:?\s*\n\s*([\d,]+\.?\d*)", text, re.IGNORECASE)
            if m:
                out["premium"] = m.group(1).strip().replace(",", "")

    # 5. Policy holder / insured name (multiple patterns for different document formats)
    if "policy_holder_name" not in out:
        for label in [
            r"Name\s+of\s+(?:the\s+)?(?:insured|proposer|policy\s*holder)",
            r"Policy\s+holder\s*:?",
            r"Registered\s+Owner\s*:?",
            r"Name\s+of\s+insured\s*:?",
            r"Proposer['\u2019]?s?\s+name\s*:?",
            r"Name\s+of\s+proposer\s*:?",
            r"Insured\s+name\s*:?",
            r"(?:Name|Proposer)\s*:?\s*\n",  # "Name:" or "Proposer:" on own line
        ]:
            m = re.search(rf"{label}\s*\n?\s*([A-Za-z][A-Za-z\s\.]{1,79})", text, re.IGNORECASE)
            if m:
                val = m.group(1).strip()
                # Exclude values that look like insurer names or numbers
                if len(val) >= 2 and len(val) <= 80 and not re.match(r"^\d+$", val):
                    if "insurance" not in val.lower() and "company" not in val.lower() and "ltd" not in val.lower():
                        out["policy_holder_name"] = val
                        break

    return out


class OcrService:
    """Process AI reader queue with Textract (forms mode); write results to flat files. Details sheet only."""

    def __init__(
        self,
        uploads_dir: Path | None = None,
        ocr_output_dir: Path | None = None,
    ):
        from app.config import OCR_OUTPUT_DIR, UPLOADS_DIR

        self.uploads_dir = Path(uploads_dir or UPLOADS_DIR).resolve()
        self.ocr_output_dir = Path(ocr_output_dir or OCR_OUTPUT_DIR).resolve()

    def _ensure_ocr_output_dir(self) -> None:
        self.ocr_output_dir.mkdir(parents=True, exist_ok=True)

    def get_output_path(self, subfolder: str, filename: str) -> Path:
        """Path where extracted text is written: ocr_output/mobile_ddmmyy/filename_stem.txt."""
        self._ensure_ocr_output_dir()
        stem = Path(filename).stem
        safe_stem = re.sub(r"[^\w\-.]", "_", stem)
        # Always use a subfolder (mobile_ddmmyy); never write at ocr_output root
        subfolder_name = _safe_subfolder_name(subfolder)
        subfolder_path = self.ocr_output_dir / subfolder_name
        subfolder_path.mkdir(parents=True, exist_ok=True)
        return subfolder_path / f"{safe_stem}.txt"

    def read_extraction_file(self, subfolder: str, filename: str) -> str | None:
        """Read extracted text from the flat file if it exists."""
        path = self.get_output_path(subfolder, filename)
        if not path.exists():
            return None
        return path.read_text(encoding="utf-8", errors="replace")

    def process_next(self) -> dict | None:
        """
        Process the oldest queued item: Details sheet (Textract) or Aadhar (Vision).
        Picks any queued item; branches on filename (details vs aadhar). Returns result dict or None if none queued.
        """
        with get_connection() as conn:
            AiReaderQueueRepository.ensure_table(conn)
            row = AiReaderQueueRepository.get_oldest_queued(conn)  # any queued item
            if not row:
                return None

            qid = row["id"]
            subfolder = row["subfolder"]
            filename = row["filename"]

            AiReaderQueueRepository.update_status(conn, qid, "processing")
            conn.commit()

        input_path = self.uploads_dir / subfolder / filename
        if not input_path.exists():
            with get_connection() as conn:
                AiReaderQueueRepository.update_status(conn, qid, "failed")
                conn.commit()
            return {
                "id": qid,
                "subfolder": subfolder,
                "filename": filename,
                "status": "failed",
                "error": "File not found",
                "extracted_text": None,
                "output_path": None,
                "document_type": None,
                "classification_confidence": None,
            }

        fn_lower = filename.lower()
        if any(hint in fn_lower for hint in DETAILS_FILENAME_HINTS):
            return self._process_details_sheet(qid, subfolder, filename, input_path)
        if AADHAR_FILENAME_CONTAINS in fn_lower:
            return self._process_aadhar(qid, subfolder, filename, input_path)
        # Unknown type: mark done so queue advances
        with get_connection() as conn:
            AiReaderQueueRepository.update_classification(conn, qid, "Other", 0.0)
            AiReaderQueueRepository.update_status(conn, qid, "done")
            conn.commit()
        return {
            "id": qid,
            "subfolder": subfolder,
            "filename": filename,
            "status": "done",
            "error": None,
            "extracted_text": None,
            "output_path": None,
            "document_type": "Other",
            "classification_confidence": 0.0,
        }

    def _write_aadhar_fields_summary_file(self, subfolder: str, customer: dict[str, str]) -> None:
        self._ensure_ocr_output_dir()
        output_path = self.get_output_path(subfolder, "Aadhar.jpg")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        lines = ["Aadhar scan – 15 extracted fields", ""]
        for key, label in AADHAR_15_FIELDS:
            value = customer.get(key)
            if value and str(value).strip():
                lines.append(f"{label}: {value.strip()}")
        if customer.get("address") and str(customer["address"]).strip():
            lines.append(f"Address (constructed): {customer['address'].strip()}")
        output_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

    def _persist_upload_merge(
        self,
        subfolder: str,
        frag_a: dict[str, Any] | None,
        frag_d: dict[str, Any] | None,
    ) -> dict[str, str]:
        """Single JSON write: Aadhaar customer fragment + Details vehicle/insurance (after parallel compile)."""
        customer: dict[str, str] = {}
        vehicle: dict = {}
        insurance_merged: dict = {}
        details_customer_name = None

        if frag_a and frag_a.get("ok", True) and frag_a.get("customer") is not None:
            customer = dict(frag_a["customer"])

        if frag_d and not frag_d.get("error"):
            vehicle = frag_d.get("vehicle") or {}
            insurance = frag_d.get("insurance") or {}
            dc = frag_d.get("details_customer") or {}
            details_customer_name = frag_d.get("details_customer_name")

            for key, val in dc.items():
                if not val:
                    continue
                if key in (
                    "profession",
                    "marital_status",
                    "financier",
                    "nominee_name",
                    "nominee_age",
                    "nominee_gender",
                    "nominee_relationship",
                ):
                    continue
                if not customer.get(key):
                    customer[key] = val

            insurance_merged = {
                **{k: v for k, v in insurance.items() if v},
                **{
                    k: v
                    for k, v in dc.items()
                    if k
                    in (
                        "profession",
                        "marital_status",
                        "financier",
                        "nominee_name",
                        "nominee_age",
                        "nominee_gender",
                        "nominee_relationship",
                    )
                    and v
                },
            }

        customer = enrich_customer_address_from_freeform(customer)

        extraction_error = None
        if not _aadhar_identity_ok(customer):
            extraction_error = (
                (frag_a or {}).get("extraction_note")
                or "Aadhaar automated read did not yield a clear name or ID number. "
                "Enter or verify customer fields manually."
            )

        json_path = _json_output_path(self.ocr_output_dir, subfolder)
        json_path.parent.mkdir(parents=True, exist_ok=True)
        details_json: dict[str, Any] = {
            "vehicle": vehicle,
            "customer": customer,
            "insurance": insurance_merged,
        }
        if details_customer_name:
            details_json["details_customer_name"] = details_customer_name
        if extraction_error:
            details_json["extraction_error"] = extraction_error
        json_path.write_text(json.dumps(details_json, indent=2), encoding="utf-8")
        return customer

    def _upload_fragment_aadhar(
        self,
        scan_bundle: dict[str, Any],
        prefetch: dict[str, dict],
    ) -> dict[str, Any]:
        customer, raw_parts, note, process_timings = _pipeline_merge_aadhar_customer(
            scan_bundle.get("front_bytes"),
            scan_bundle.get("back_bytes"),
            prefetch.get("aadhar_front"),
            prefetch.get("aadhar_back"),
        )
        return {
            "ok": True,
            "customer": customer,
            "raw_parts": raw_parts,
            "extraction_note": note,
            "process_timings": process_timings,
        }

    def _upload_fragment_details(
        self,
        details_path: Path,
        prefetch: dict[str, dict],
    ) -> dict[str, Any]:
        try:
            frag = _compile_details_sheet_fragment(details_path, prefetch.get("details_forms"))
            frag["ok"] = True
            return frag
        except Exception as e:
            return {"ok": False, "error": str(e)}

    def process_uploaded_subfolder(self, subfolder: str) -> dict:
        """
        Run extraction directly on uploaded files (no queue).

        Flow: **AWS Textract prefetch** (when enabled) on Aadhaar and other scans; **Aadhaar assembly**
        uses **Textract text only** on front/back (no UIDAI QR in this path). **Details sheet** is
        compiled in **parallel** with the Aadhaar pipeline, then results are merged once into JSON.
        **Details** raster/PDF: Textract **AnalyzeDocument FORMS** (structured key-values).

        Writes Raw_OCR.txt and returns ``section_timings_ms`` for operator visibility.
        """
        subdir = self.uploads_dir / subfolder
        if not subdir.exists() or not subdir.is_dir():
            return {"error": f"Subfolder not found: {subfolder}", "processed": []}

        processed: list[str] = []
        errors: list[str] = []
        section_timings_ms: dict[str, int] = {}
        t_total = time.perf_counter()

        from concurrent.futures import ThreadPoolExecutor

        from app.config import OCR_UPLOAD_PARALLEL_TEXTRACT

        scan_bundle = _load_aadhar_scan_bytes(subdir)
        prefetch: dict[str, dict] = {}
        prefetch_job_ms: dict[str, int] = {}
        t_pref = time.perf_counter()
        if OCR_UPLOAD_PARALLEL_TEXTRACT:
            prefetch, prefetch_job_ms = _parallel_textract_prefetch_upload_subfolder(subdir)
            section_timings_ms["aws_textract_prefetch_ms"] = int(
                (time.perf_counter() - t_pref) * 1000
            )

        aadhar_path = subdir / "Aadhar.jpg"
        details_path = subdir / "Details.jpg"

        frag_a: dict[str, Any] | None = None
        frag_d: dict[str, Any] | None = None
        t_par = time.perf_counter()
        if aadhar_path.exists() or details_path.exists():
            with ThreadPoolExecutor(max_workers=2) as ex:
                fa = (
                    ex.submit(self._upload_fragment_aadhar, scan_bundle, prefetch)
                    if aadhar_path.exists()
                    else None
                )
                fb = (
                    ex.submit(self._upload_fragment_details, details_path, prefetch)
                    if details_path.exists()
                    else None
                )
                if fa is not None:
                    frag_a = fa.result()
                if fb is not None:
                    frag_d = fb.result()
        section_timings_ms["parallel_aadhar_details_compile_ms"] = int(
            (time.perf_counter() - t_par) * 1000
        )

        _merge_granular_upload_process_timings(
            section_timings_ms,
            prefetch_job_ms=prefetch_job_ms,
            frag_a=frag_a,
            frag_d=frag_d,
        )
        logger.info(
            "ocr_upload_process_timings subfolder=%s aadhar_textract_front_ms=%s "
            "aadhar_textract_back_ms=%s detail_sheet_textract_ms=%s",
            subfolder,
            section_timings_ms.get("aadhar_textract_front_ms", 0),
            section_timings_ms.get("aadhar_textract_back_ms", 0),
            section_timings_ms.get("detail_sheet_textract_ms", 0),
        )

        self._raw_ocr_parts = []
        if frag_a and frag_a.get("raw_parts"):
            self._raw_ocr_parts.extend(frag_a["raw_parts"])
        if frag_d and frag_d.get("ok") and frag_d.get("full_text"):
            self._raw_ocr_parts.append(("Details.jpg", frag_d["full_text"]))

        if aadhar_path.exists():
            if frag_a and frag_a.get("ok"):
                processed.append("Aadhar.jpg")
            else:
                errors.append(f"Aadhar.jpg: {frag_a.get('error') if frag_a else 'failed'}")
        if details_path.exists():
            if frag_d and frag_d.get("ok"):
                processed.append("Details.jpg")
            else:
                errors.append(f"Details.jpg: {frag_d.get('error') if frag_d else 'failed'}")

        t_merge = time.perf_counter()
        merged_customer: dict[str, str] = {}
        if aadhar_path.exists() or details_path.exists():
            merged_customer = self._persist_upload_merge(subfolder, frag_a, frag_d)
            if aadhar_path.exists() and merged_customer:
                self._write_aadhar_fields_summary_file(subfolder, merged_customer)
        section_timings_ms["merge_write_json_ms"] = int((time.perf_counter() - t_merge) * 1000)

        insurance_path = subdir / "Insurance.jpg"
        t_ins = time.perf_counter()
        if insurance_path.exists():
            try:
                self._process_insurance_sheet(
                    subfolder,
                    "Insurance.jpg",
                    insurance_path,
                    textract_prefetch=prefetch.get("insurance"),
                )
                processed.append("Insurance.jpg")
            except Exception as e:
                errors.append(f"Insurance.jpg: {e}")
        section_timings_ms["insurance_ms"] = int((time.perf_counter() - t_ins) * 1000)

        t_extras = time.perf_counter()
        seen_raw = {fn for fn, _ in self._raw_ocr_parts}
        for extra_file, prefetch_key in (
            ("Aadhar_back.jpg", "aadhar_back"),
            ("Financing.jpg", "financing"),
        ):
            if extra_file in seen_raw:
                continue
            extra_path = subdir / extra_file
            if extra_path.exists():
                try:
                    from app.services.textract_service import extract_text_from_bytes

                    if prefetch_key in prefetch:
                        result = prefetch[prefetch_key]
                    else:
                        result = extract_text_from_bytes(extra_path.read_bytes())
                    if not result.get("error") and result.get("full_text"):
                        self._raw_ocr_parts.append((extra_file, result["full_text"]))
                except Exception:
                    pass
        section_timings_ms["extras_raw_ms"] = int((time.perf_counter() - t_extras) * 1000)

        t_raw = time.perf_counter()
        if self._raw_ocr_parts:
            self._ensure_ocr_output_dir()
            subfolder_name = _safe_subfolder_name(subfolder)
            subfolder_path = self.ocr_output_dir / subfolder_name
            subfolder_path.mkdir(parents=True, exist_ok=True)
            raw_lines = []
            for fname, text in self._raw_ocr_parts:
                raw_lines.append(f"--- {fname} ---")
                raw_lines.append(text.strip() if text else "")
                raw_lines.append("")
            (subfolder_path / "Raw_OCR.txt").write_text("\n".join(raw_lines), encoding="utf-8")
            _apply_aadhar_textract_fallbacks_from_parts(
                self.ocr_output_dir, subfolder, list(self._raw_ocr_parts)
            )
        section_timings_ms["raw_ocr_finalize_ms"] = int((time.perf_counter() - t_raw) * 1000)

        self._raw_ocr_parts = None

        section_timings_ms["total_ms"] = int((time.perf_counter() - t_total) * 1000)

        result: dict[str, Any] = {"processed": processed, "section_timings_ms": section_timings_ms}
        if errors:
            result["errors"] = errors
        return result

    def _process_details_sheet(
        self,
        qid: int | None,
        subfolder: str,
        filename: str,
        input_path: Path,
        textract_forms_prefetch: dict | None = None,
    ) -> dict:
        """Run Textract (forms) on Details sheet; write text + JSON; optionally update queue."""
        if qid is not None:
            with get_connection() as conn:
                AiReaderQueueRepository.update_classification(conn, qid, "Details sheet", 1.0)
                conn.commit()
        try:
            fmt = _details_input_format(input_path)
            if fmt == "docx":
                key_value_pairs, docx_full = _key_value_pairs_from_docx(input_path)
                if not key_value_pairs and not (docx_full or "").strip():
                    raise RuntimeError(
                        "Could not read the Word document (.docx). Ensure it is a valid .docx Sales Detail Sheet."
                    )
                result = {
                    "error": None,
                    "full_text": docx_full,
                    "key_value_pairs": key_value_pairs,
                }
            elif fmt in ("jpeg", "png", "pdf"):
                from app.services.textract_service import extract_forms_from_bytes

                if textract_forms_prefetch is not None:
                    result = textract_forms_prefetch
                else:
                    result = extract_forms_from_bytes(input_path.read_bytes())
                if result.get("error"):
                    raise RuntimeError(result["error"])
            else:
                raise RuntimeError(
                    f"Unsupported Details file format (detected={fmt!r}). "
                    "Use a JPEG/PNG scan, PDF export, or .docx Sales Detail Sheet."
                )

            key_value_pairs = result.get("key_value_pairs") or []
            lines = []
            lines.append("Document: Details sheet (Textract Forms)\n")
            for kv in key_value_pairs:
                lines.append(f"{kv.get('key', '')}: {kv.get('value', '')}")
            if result.get("full_text"):
                lines.append("\n--- Full text ---\n")
                lines.append(result["full_text"])
            text = "\n".join(lines)

            if hasattr(self, "_raw_ocr_parts") and self._raw_ocr_parts is not None:
                self._raw_ocr_parts.append((filename, result.get("full_text") or ""))

            self._ensure_ocr_output_dir()
            output_path = self.get_output_path(subfolder, filename)  # used for return value only; no separate .txt file

            vehicle = _map_key_value_pairs_to_vehicle(key_value_pairs)
            insurance = _map_key_value_pairs_to_insurance(key_value_pairs)
            details_customer = _map_key_value_pairs_to_details_customer(key_value_pairs)
            details_customer_name = _extract_details_customer_name(key_value_pairs)
            if result.get("full_text"):
                from_vehicle = _parse_vehicle_from_full_text(result["full_text"])
                for k, v in from_vehicle.items():
                    if v and not vehicle.get(k):
                        vehicle[k] = v
                from_full = _parse_insurance_from_full_text(result["full_text"])
                if from_full.get("customer_name") and not details_customer_name:
                    details_customer_name = from_full["customer_name"]
                for k, v in from_full.items():
                    if k == "customer_name":
                        continue  # Used for details_customer_name, not insurance
                    if v and not insurance.get(k):
                        insurance[k] = v
            json_path = _json_output_path(self.ocr_output_dir, subfolder)
            json_path.parent.mkdir(parents=True, exist_ok=True)
            customer = {}
            existing_insurance: dict = {}
            if json_path.exists():
                try:
                    existing = json.loads(json_path.read_text(encoding="utf-8"))
                    customer = existing.get("customer") or {}
                    if not isinstance(customer, dict):
                        customer = {}
                    # Compliance: never persist full Aadhar
                    if customer.get("aadhar_id"):
                        customer["aadhar_id"] = _aadhar_last4(customer["aadhar_id"]) or ""
                    existing_insurance = existing.get("insurance") or {}
                    if not isinstance(existing_insurance, dict):
                        existing_insurance = {}
                except Exception:
                    pass
            if details_customer.get("name") and not details_customer_name:
                details_customer_name = details_customer.get("name")

            customer_merged = dict(customer)
            for key, val in details_customer.items():
                if not val:
                    continue
                if key in (
                    "profession",
                    "marital_status",
                    "financier",
                    "nominee_name",
                    "nominee_age",
                    "nominee_gender",
                    "nominee_relationship",
                ):
                    continue
                if not customer_merged.get(key):
                    customer_merged[key] = val

            insurance_merged = {
                **existing_insurance,
                **{k: v for k, v in insurance.items() if v},
                **{
                    k: v
                    for k, v in details_customer.items()
                    if k
                    in (
                        "profession",
                        "marital_status",
                        "financier",
                        "nominee_name",
                        "nominee_age",
                        "nominee_gender",
                        "nominee_relationship",
                    )
                    and v
                },
            }
            customer_merged = enrich_customer_address_from_freeform(customer_merged)
            details_json = {"vehicle": vehicle, "customer": customer_merged, "insurance": insurance_merged}
            if details_customer_name:
                details_json["details_customer_name"] = details_customer_name
            json_path.parent.mkdir(parents=True, exist_ok=True)
            json_path.write_text(json.dumps(details_json, indent=2), encoding="utf-8")

            if qid is not None:
                with get_connection() as conn:
                    AiReaderQueueRepository.update_status(conn, qid, "done")
                    conn.commit()

            return {
                "id": qid,
                "subfolder": subfolder,
                "filename": filename,
                "status": "done",
                "error": None,
                "extracted_text": text,
                "output_path": str(output_path),
                "document_type": "Details sheet",
                "classification_confidence": 1.0,
            }
        except Exception as e:
            if qid is not None:
                with get_connection() as conn:
                    AiReaderQueueRepository.update_status(conn, qid, "failed")
                    conn.commit()
                return {
                    "id": qid,
                    "subfolder": subfolder,
                    "filename": filename,
                    "status": "failed",
                    "error": str(e),
                    "extracted_text": None,
                    "output_path": None,
                    "document_type": None,
                    "classification_confidence": None,
                }
            raise

    def _process_aadhar(
        self,
        qid: int | None,
        subfolder: str,
        filename: str,
        input_path: Path,
        front_textract_prefetch: dict | None = None,
    ) -> dict:
        """Run Aadhar extraction: **AWS Textract** on front (and back when geo fields are weak); merge into subfolder JSON."""
        if qid is not None:
            with get_connection() as conn:
                AiReaderQueueRepository.update_classification(conn, qid, "Aadhar", 1.0)
                conn.commit()
        try:
            raw_bytes = input_path.read_bytes()
            back_path = input_path.parent / "Aadhar_back.jpg"
            back_bytes = back_path.read_bytes() if back_path.is_file() else None
            ft = front_textract_prefetch
            if ft is None:
                t_ft = time.perf_counter()
                try:
                    from app.services.textract_service import extract_text_from_bytes

                    ft = extract_text_from_bytes(raw_bytes)
                except Exception:
                    ft = {"error": "textract", "full_text": ""}
                front_textract_sync_ms = int((time.perf_counter() - t_ft) * 1000)
            else:
                front_textract_sync_ms = 0

            piped, raw_parts, note, pipe_timings = _pipeline_merge_aadhar_customer(
                raw_bytes, back_bytes, ft, None
            )
            logger.info(
                "ocr_queue_aadhar_process_timings subfolder=%s "
                "aadhar_textract_front_prefetch_or_sync_ms=%s aadhar_textract_pipeline_front_ms=%s "
                "aadhar_textract_pipeline_back_ms=%s",
                subfolder,
                front_textract_sync_ms,
                pipe_timings.get("aadhar_textract_front_ms", 0),
                pipe_timings.get("aadhar_textract_back_ms", 0),
            )

            if hasattr(self, "_raw_ocr_parts") and self._raw_ocr_parts is not None:
                for fn, txt in raw_parts:
                    self._raw_ocr_parts.append((fn, txt))

            self._ensure_ocr_output_dir()
            json_path = _json_output_path(self.ocr_output_dir, subfolder)
            data = {"vehicle": {}, "customer": {}, "insurance": {}}
            existing_customer: dict[str, str] = {}
            if json_path.exists():
                try:
                    data = json.loads(json_path.read_text(encoding="utf-8"))
                    if not isinstance(data.get("vehicle"), dict):
                        data["vehicle"] = {}
                    if not isinstance(data.get("customer"), dict):
                        data["customer"] = {}
                    if not isinstance(data.get("insurance"), dict):
                        data["insurance"] = {}
                    ec = data.get("customer") or {}
                    if isinstance(ec, dict):
                        existing_customer = {str(k): str(v) for k, v in ec.items() if v is not None}
                except Exception:
                    pass
            customer = _merge_qr_customer_into_existing(existing_customer, piped)
            if customer.get("aadhar_id"):
                customer["aadhar_id"] = _aadhar_last4(customer["aadhar_id"]) or ""
            _default_gender_male_if_unread(customer)
            customer = enrich_customer_address_from_freeform(customer)
            data["customer"] = customer
            if note and not _aadhar_identity_ok(customer):
                data["extraction_error"] = note
            elif "extraction_error" in data and _aadhar_identity_ok(customer):
                data.pop("extraction_error", None)
            json_path.parent.mkdir(parents=True, exist_ok=True)
            json_path.write_text(json.dumps(data, indent=2), encoding="utf-8")

            # Write a file in ocr_output/subfolder listing the 15 extracted Aadhar fields (with display labels)
            output_path = self.get_output_path(subfolder, filename)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            lines = ["Aadhar scan – 15 extracted fields", ""]
            for key, label in AADHAR_15_FIELDS:
                value = customer.get(key)
                if value and str(value).strip():
                    lines.append(f"{label}: {value.strip()}")
            if customer.get("address") and str(customer["address"]).strip():
                lines.append(f"Address (constructed): {customer['address'].strip()}")
            summary = "\n".join(f"{k}: {v}" for k, v in customer.items() if v)
            output_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

            if qid is not None:
                with get_connection() as conn:
                    AiReaderQueueRepository.update_status(conn, qid, "done")
                    conn.commit()

            return {
                "id": qid,
                "subfolder": subfolder,
                "filename": filename,
                "status": "done",
                "error": None,
                "extracted_text": summary,
                "output_path": str(output_path),
                "document_type": "Aadhar",
                "classification_confidence": 1.0,
            }
        except Exception as e:
            if qid is not None:
                with get_connection() as conn:
                    AiReaderQueueRepository.update_status(conn, qid, "failed")
                    conn.commit()
                return {
                    "id": qid,
                    "subfolder": subfolder,
                    "filename": filename,
                    "status": "failed",
                    "error": str(e),
                    "extracted_text": None,
                    "output_path": None,
                    "document_type": None,
                    "classification_confidence": None,
                }
            raise

    def _process_insurance_sheet(
        self,
        subfolder: str,
        filename: str,
        input_path: Path,
        textract_prefetch: dict | None = None,
    ) -> None:
        """Run Textract on Insurance.jpg (TEXT mode); extract insurer, policy from/to, premium via regex on full_text."""
        from app.services.textract_service import extract_text_from_bytes

        if textract_prefetch is not None:
            result = textract_prefetch
        else:
            result = extract_text_from_bytes(input_path.read_bytes())

        if result.get("error"):
            raise RuntimeError(result["error"])

        full_text = result.get("full_text") or ""
        insurance_policy = _parse_insurance_policy_from_full_text(full_text)

        if hasattr(self, "_raw_ocr_parts") and self._raw_ocr_parts is not None:
            self._raw_ocr_parts.append(("Insurance.jpg", full_text))

        # Merge into Details.json insurance section
        json_path = _json_output_path(self.ocr_output_dir, subfolder)
        data: dict = {"vehicle": {}, "customer": {}, "insurance": {}}
        if json_path.exists():
            try:
                data = json.loads(json_path.read_text(encoding="utf-8"))
                if not isinstance(data.get("vehicle"), dict):
                    data["vehicle"] = {}
                if not isinstance(data.get("customer"), dict):
                    data["customer"] = {}
                if not isinstance(data.get("insurance"), dict):
                    data["insurance"] = {}
            except Exception:
                pass
        existing_insurance = data.get("insurance") or {}
        insurance_merged = {**existing_insurance, **{k: v for k, v in insurance_policy.items() if v}}
        data["insurance"] = insurance_merged
        # Compliance: sanitize customer aadhar before persisting
        customer = data.get("customer") or {}
        if customer.get("aadhar_id"):
            customer["aadhar_id"] = _aadhar_last4(customer["aadhar_id"]) or ""
        json_path.parent.mkdir(parents=True, exist_ok=True)
        json_path.write_text(json.dumps(data, indent=2), encoding="utf-8")

    def get_extracted_details(self, subfolder: str) -> dict | None:
        """
        Return structured extracted details (vehicle, customer) for a subfolder.
        Loads from JSON if present. Aadhaar fields are filled from upload-time Textract and
        ``Raw_OCR.txt`` fallbacks (no per-poll UIDAI QR decode).
        """
        self._ensure_ocr_output_dir()
        json_path = _json_output_path(self.ocr_output_dir, subfolder)
        data = {"vehicle": {}, "customer": {}}
        if json_path.exists():
            try:
                data = json.loads(json_path.read_text(encoding="utf-8"))
                if not isinstance(data.get("vehicle"), dict):
                    data["vehicle"] = {}
                if not isinstance(data.get("customer"), dict):
                    data["customer"] = {}
                if not isinstance(data.get("insurance"), dict):
                    data["insurance"] = {}
            except Exception:
                pass
        if "insurance" not in data:
            data["insurance"] = {}

        customer = data.get("customer") or {}
        if not isinstance(customer, dict):
            customer = {}
        aadhar_path = self.uploads_dir / subfolder / "Aadhar.jpg"
        back_p = self.uploads_dir / subfolder / "Aadhar_back.jpg"

        # Textract text in Raw_OCR.txt: fill DOB/gender from front, address from back when parsers missed fields.
        _apply_aadhar_textract_fallbacks_from_raw_ocr_file(self.ocr_output_dir, subfolder)
        if json_path.exists():
            try:
                data = json.loads(json_path.read_text(encoding="utf-8"))
                if not isinstance(data.get("vehicle"), dict):
                    data["vehicle"] = {}
                if not isinstance(data.get("customer"), dict):
                    data["customer"] = {}
                if not isinstance(data.get("insurance"), dict):
                    data["insurance"] = {}
            except Exception:
                pass

        # Compliance: sanitize customer aadhar on return (handles legacy JSON with full Aadhar)
        customer = data.get("customer") or {}
        if customer.get("aadhar_id"):
            customer["aadhar_id"] = _aadhar_last4(customer["aadhar_id"]) or ""
        if aadhar_path.is_file() or back_p.is_file():
            _default_gender_male_if_unread(customer)
            customer = enrich_customer_address_from_freeform(customer)
            data["customer"] = customer
            try:
                json_path.parent.mkdir(parents=True, exist_ok=True)
                json_path.write_text(json.dumps(data, indent=2), encoding="utf-8")
            except Exception:
                pass
        # Sanitize nominee_age: extract digits only (e.g. "30/m" -> "30"); clear if invalid
        ins = data.get("insurance") or {}
        if ins.get("nominee_age"):
            sanitized = _sanitize_nominee_age(str(ins["nominee_age"]))
            ins["nominee_age"] = sanitized
        # Name match validation: Aadhar, Details sheet, Insurance
        aadhar_name = customer.get("name")
        details_name = data.get("details_customer_name")
        insurance_name = (data.get("insurance") or {}).get("policy_holder_name")
        name_err = _validate_name_match(aadhar_name, details_name, insurance_name)
        if name_err:
            data["name_mismatch_error"] = name_err
        return data

    def list_extractions(self, limit: int = 200) -> list[dict]:
        """List queue items (oldest first) with extracted text from flat files."""
        with get_connection() as conn:
            AiReaderQueueRepository.ensure_table(conn)
            rows = AiReaderQueueRepository.list_all(conn, limit=limit)

        rows = list(reversed(rows))
        result = []
        for row in rows:
            r = dict(row)
            for k in ("created_at", "updated_at"):
                if k in r and isinstance(r[k], datetime):
                    r[k] = r[k].isoformat()
            text = self.read_extraction_file(r["subfolder"], r["filename"])
            r["extracted_text"] = text
            r["output_path"] = str(self.get_output_path(r["subfolder"], r["filename"]))
            result.append(r)
        return result


def validate_name_match_for_subfolder(ocr_output_dir: Path, subfolder: str) -> str | None:
    """
    Load OCR JSON for subfolder and validate name match across Aadhar, Details, Insurance.
    Returns error message if mismatch, None if OK.
    """
    json_path = _json_output_path(ocr_output_dir, subfolder)
    if not json_path.exists():
        return None
    try:
        data = json.loads(json_path.read_text(encoding="utf-8"))
    except Exception:
        return None
    customer = data.get("customer") or {}
    aadhar_name = customer.get("name")
    details_name = data.get("details_customer_name")
    insurance_name = (data.get("insurance") or {}).get("policy_holder_name")
    return _validate_name_match(aadhar_name, details_name, insurance_name)
