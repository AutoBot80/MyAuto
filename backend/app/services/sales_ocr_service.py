"""AI reader queue and Add Sales upload extraction: Textract FORMS for detail sheets; Aadhaar uses Textract text only (no UIDAI QR decode)."""

from __future__ import annotations

import json
import logging
import re
import time
from difflib import SequenceMatcher
from datetime import datetime
from pathlib import Path
from typing import Any

from app.db import get_connection
from app.repositories.ai_reader_queue import AiReaderQueueRepository
from app.services.utility_functions import (
    default_profession_if_empty,
    fuzzy_best_option_label,
    normalize_for_fuzzy_match,
    normalize_nominee_relationship_value,
    sanitize_details_sheet_insurer_value,
)
from app.services.customer_address_infer import (
    enrich_customer_address_from_freeform,
    normalize_address_freeform,
)
from app.services.ocr_extraction_log import append_ocr_extraction_log
from app.services.page_classifier import (
    FILENAME_AADHAR_FRONT,
    FILENAME_SALES_DETAIL_SHEET_PDF,
    LEGACY_AADHAR_FRONT_JPG,
    LEGACY_DETAILS_JPG,
)

logger = logging.getLogger(__name__)

# Bulk pre-OCR drops single-page PDFs here for Textract (see ``pre_ocr_service.FOR_OCR_SUBDIR``).
FOR_OCR_SUBDIR = "for_OCR"


def _prefer_for_ocr_input(subdir: Path, pdf_name: str, jpg_name: str, *legacy_root: str) -> Path:
    """Prefer ``for_OCR/<pdf_name>`` when present (bulk pipeline); else first existing sale-folder file."""
    p = subdir / FOR_OCR_SUBDIR / pdf_name
    if p.is_file():
        return p
    for name in (jpg_name,) + legacy_root:
        r = subdir / name
        if r.is_file():
            return r
    return subdir / jpg_name


def _prefer_details_sheet_input(subdir: Path) -> Path:
    """Sales detail: pre-OCR PDF ``Sales_Detail_Sheet.pdf``, ``for_OCR`` PDFs, or legacy ``Details.jpg``."""
    candidates = [
        subdir / FOR_OCR_SUBDIR / FILENAME_SALES_DETAIL_SHEET_PDF,
        subdir / FOR_OCR_SUBDIR / "Details.pdf",
        subdir / FILENAME_SALES_DETAIL_SHEET_PDF,
        subdir / "Details.pdf",
        subdir / LEGACY_DETAILS_JPG,
    ]
    for c in candidates:
        if c.is_file():
            return c
    return subdir / LEGACY_DETAILS_JPG


def _rel_upload_label(path: Path, subdir: Path) -> str:
    """Stable label for Raw_OCR / logs (e.g. ``for_OCR/Aadhar.pdf``)."""
    try:
        return str(path.relative_to(subdir)).replace("\\", "/")
    except ValueError:
        return path.name


def _aadhar_last4(aadhar_id: str | None) -> str | None:
    """Return only last 4 digits of Aadhar for compliance; never persist full 12 digits."""
    if not aadhar_id or not str(aadhar_id).strip():
        return None
    digits = "".join(c for c in str(aadhar_id) if c.isdigit())
    return digits[-4:] if len(digits) >= 4 else digits or None


def _safe_subfolder_name(subfolder: str) -> str:
    """Safe directory name under ocr_output (one segment, no path separators)."""
    return re.sub(r"[^\w\-]", "_", (subfolder or "").strip()) or "default"


def _safe_output_basename(subfolder: str, filename: str) -> str:
    """Build a safe filename for output: subfolder_filename.txt (no path chars)."""
    safe_sub = _safe_subfolder_name(subfolder)
    safe_name = Path(filename).stem
    safe_name = re.sub(r"[^\w\-.]", "_", safe_name)
    return f"{safe_sub}_{safe_name}.txt"


def _ocr_subfolder_path(output_dir: Path, subfolder: str) -> Path:
    """Path to per-customer subfolder under ocr_output: ocr_output/mobile_ddmmyy/."""
    return output_dir / _safe_subfolder_name(subfolder)


# OCR output JSON filename (in ocr_output subfolder only; input files stay as Details.jpg)
OCR_OUTPUT_JSON_STEM = "OCR_To_be_Used"


def _json_output_path(output_dir: Path, subfolder: str) -> Path:
    """Path for structured JSON output: ocr_output/mobile_ddmmyy/OCR_To_be_Used.json."""
    subfolder_name = _safe_subfolder_name(subfolder)
    return output_dir / subfolder_name / f"{OCR_OUTPUT_JSON_STEM}.json"


# Filename patterns for queue processing.
DETAILS_FILENAME_HINTS = (
    "details",
    "detail_sheet",
    "sales_detail",
    "sales detail",
    "a5 template",
)
AADHAR_FILENAME_CONTAINS = "aadhar"

# 15 Aadhar fields (display order and labels for ocr_output file).
AADHAR_15_FIELDS: list[tuple[str, str]] = [
    ("aadhar_id", "Aadhar ID"),
    ("name", "Name"),
    ("gender", "Gender"),
    ("year_of_birth", "Year of birth"),
    ("date_of_birth", "Date of birth"),
    ("care_of", "Care of"),
    ("house", "House"),
    ("street", "Street"),
    ("location", "Location"),
    ("city", "City"),
    ("post_office", "Post Office"),
    ("district", "District"),
    ("sub_district", "Sub District"),
    ("state", "State"),
    ("pin_code", "Pin Code"),
]

# Customer keys for Aadhaar merges (same shape as legacy UIDAI POA; filled from Textract parsers).
AADHAR_QR_CUSTOMER_KEYS: tuple[str, ...] = (
    "aadhar_id",
    "name",
    "gender",
    "year_of_birth",
    "date_of_birth",
    "care_of",
    "house",
    "street",
    "location",
    "city",
    "post_office",
    "district",
    "sub_district",
    "state",
    "pin_code",
)


def _rebuild_aadhar_address_from_parts(
    customer: dict[str, str],
    *,
    preserve_existing_address: bool = False,
) -> None:
    """
    UIDAI POA often uses city (vtc), district, post office without house/street.
    Mutates `customer` in place: sets or clears `address`.
    If POA parts are missing but `address` was already set (e.g. Details sheet), keep it.
    When ``preserve_existing_address`` is True, do not replace a non-empty existing ``address``
    (Details sheet full line) with a shorter UIDAI-composed line.
    """
    parts = [
        customer.get("care_of"),
        customer.get("house"),
        customer.get("street"),
        customer.get("location"),
        customer.get("city"),
        customer.get("district"),
        customer.get("sub_district"),
        customer.get("post_office"),
        customer.get("state"),
        customer.get("pin_code"),
    ]
    address = ", ".join(p for p in parts if p and str(p).strip())
    if address:
        prev = customer.get("address")
        if preserve_existing_address and prev and str(prev).strip():
            return
        customer["address"] = address
    elif not (customer.get("address") and str(customer["address"]).strip()):
        customer.pop("address", None)


def _merge_qr_customer_into_existing(
    existing: dict[str, str],
    qr_merged: dict[str, str],
) -> dict[str, str]:
    """
    Fill blank customer keys from an Aadhaar-derived fragment (Textract-parsed fields).
    Preserves Details sheet / DB values when set.
    """
    out = dict(existing) if existing else {}
    if not qr_merged:
        return out
    for k in AADHAR_QR_CUSTOMER_KEYS:
        qv = qr_merged.get(k)
        if not qv or not str(qv).strip():
            continue
        ev = out.get(k)
        if ev is None or not str(ev).strip():
            out[k] = str(qv).strip()
    # Fill address from POA parts when customer had no address; keep Details sheet address otherwise
    _rebuild_aadhar_address_from_parts(out, preserve_existing_address=True)
    if qr_merged.get("address") and str(qr_merged["address"]).strip():
        if not (out.get("address") and str(out["address"]).strip()):
            out["address"] = str(qr_merged["address"]).strip()
    return out


def _load_aadhar_scan_bytes(subdir: Path) -> dict[str, Any]:
    """
    Load Aadhaar front/back bytes for the Textract pipeline (no QR decode).
    Prefers ``for_OCR/Aadhar.pdf`` (and ``Aadhar_back.pdf``) when present.
    """
    out: dict[str, Any] = {
        "front_bytes": None,
        "back_bytes": None,
        "front_src": FILENAME_AADHAR_FRONT,
        "back_src": "Aadhar_back.jpg",
    }
    ap = _prefer_for_ocr_input(subdir, "Aadhar.pdf", FILENAME_AADHAR_FRONT, LEGACY_AADHAR_FRONT_JPG)
    bp = _prefer_for_ocr_input(subdir, "Aadhar_back.pdf", "Aadhar_back.jpg")
    try:
        out["front_bytes"] = ap.read_bytes() if ap.is_file() else None
        out["front_src"] = _rel_upload_label(ap, subdir) if ap.is_file() else FILENAME_AADHAR_FRONT
    except OSError:
        out["front_bytes"] = None
    try:
        out["back_bytes"] = bp.read_bytes() if bp.is_file() else None
        out["back_src"] = _rel_upload_label(bp, subdir) if bp.is_file() else "Aadhar_back.jpg"
    except OSError:
        out["back_bytes"] = None
    return out


def _normalize_aadhar_back_address_chunk(chunk: str) -> str:
    chunk = chunk.replace("\r\n", "\n")
    chunk = re.sub(r"[\|`]+", " ", chunk)
    chunk = re.sub(r"\s+", " ", chunk).strip(" ,.-|")
    return chunk


def _clean_aadhar_back_cross_column_noise(addr: str) -> str:
    """
    Remove common Aadhaar-back OCR bleed where Hindi-side glyph noise gets merged into
    the English address line as short uppercase/junk comma clauses (e.g. ``49 ERHIGH``, ``HARR``).
    """
    if not addr:
        return addr

    parts = [p.strip(" ,.-") for p in re.split(r"\s*,\s*", addr) if p and p.strip(" ,.-")]
    cleaned: list[str] = []
    seen: set[str] = set()
    seen_pin = False

    for p in parts:
        pl = p.lower()
        # Keep care-of relation clauses intact.
        if re.search(r"(?i)\b(?:c|s|w|d)\s*/\s*o\b", p):
            key = re.sub(r"\s+", " ", pl).strip()
            if key and key not in seen:
                cleaned.append(p)
                seen.add(key)
            continue

        # Drop obvious OCR noise/footer fragments.
        if re.search(r"(?i)www\.|help@|uidai|virtual\s*id|aadhaar", p):
            continue
        # 1-5 digits plus junk word is usually bleed, not PIN (PIN is 6 digits).
        if re.search(r"(?<!\d)\d{1,5}\s+[A-Za-z]{3,}", p) and not re.search(r"(?<!\d)\d{6}(?!\d)", p):
            continue
        # Very short all-uppercase word chunks are commonly Hindi-side OCR artifacts.
        words = re.findall(r"[A-Za-z]+", p)
        if words and all(w.isupper() for w in words) and max(len(w) for w in words) <= 6:
            # Preserve meaningful short locality tokens only when mixed-case exists.
            if not re.search(r"[a-z]", p):
                continue

        # Deduplicate pin if repeated.
        if re.fullmatch(r"\d{6}", p):
            if seen_pin:
                continue
            seen_pin = True

        key = re.sub(r"\s+", " ", pl).strip()
        if key in seen:
            continue
        cleaned.append(p)
        seen.add(key)

    out = ", ".join(cleaned)
    # Collapse accidental duplicated ``..., State, PIN, State, PIN`` tail.
    out = re.sub(
        r"(?i)\b([A-Za-z][A-Za-z ]{2,}),\s*(\d{6})\s*,\s*\1\s*,\s*\2\b",
        r"\1, \2",
        out,
    )
    return re.sub(r"\s+", " ", out).strip(" ,")


_DATE_ONLY_LINE_RE = re.compile(r"^\s*\d{1,2}[/.-]\d{1,2}[/.-]\d{2,4}\s*$")


def _is_aadhar_back_date_only_line(line: str) -> bool:
    """Strip standalone enrollment / issue dates (e.g. ``05/10/2014``) from letter-format address blocks."""
    s = (line or "").strip()
    if not s:
        return False
    return bool(_DATE_ONLY_LINE_RE.match(s))


def _looks_like_aadhar_letter_printed_name_line(line: str) -> bool:
    """Two title-case words, no digits — typical ``First Last`` line under ``To,`` (avoid stripping 3-word localities)."""
    s = (line or "").strip()
    if re.search(r"\d", s):
        return False
    parts = s.split()
    if len(parts) != 2:
        return False
    for p in parts:
        if not re.match(r"^[A-Z][a-zA-Z'.-]+$", p):
            return False
    return True


def _aadhar_back_name_line_matches_hint(line: str, name_hint: str | None) -> bool:
    if not name_hint or len(name_hint.strip()) < 2:
        return False
    a = re.sub(r"\s+", " ", line.strip().lower())
    b = re.sub(r"\s+", " ", name_hint.strip().lower())
    return a == b or (len(b) >= 5 and (a in b or b in a))


def _trim_aadhar_letter_to_block_noise(lines: list[str]) -> list[str]:
    """Drop short OCR junk after ``To,`` (e.g. ``arries 3114``) before the real name/address."""
    out = list(lines)
    while out:
        ln = out[0]
        if re.match(r"(?i)^[a-z]{4,}\s+\d{3,}\s*$", ln) or re.match(r"(?i)^[A-Za-z]+\s+\d{4,}\s*$", ln):
            out.pop(0)
            continue
        if len(ln) <= 4 and re.search(r"\d", ln):
            out.pop(0)
            continue
        break
    return out


def _parse_aadhar_back_letter_format_address(text: str, name_hint: str | None = None) -> str | None:
    """
    UIDAI **letter** printout (no ``Address:`` label): block after ``To,`` is often
    boilerplate → Enrollment → ``To,`` → name → address lines → standalone **date** → state + PIN → **mobile**.

    Used **only as a last resort** from :func:`_parse_aadhar_back_address_from_ocr` after card-style
    ``Address:`` / ``C/O:`` / ``Near`` patterns (and optional retry when trimmed text is still too short).

    We keep lines from the first address-like content through state+PIN, strip date-only lines,
    and drop a printed **name** line when it matches ``name_hint`` (from front) or looks like two title-case words.
    """
    if not re.search(r"(?is)(?:^|\n)\s*To\s*,\s*\n", text):
        return None
    if not re.search(
        r"(?i)Unique\s+Identification\s+Authority|Enrollment\s+No|Government\s+of\s+India",
        text,
    ):
        return None
    m = re.search(r"(?is)(?:^|\n)\s*To\s*,\s*\n([\s\S]+)$", text)
    if not m:
        return None
    block = m.group(1)
    lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
    cut: list[str] = []
    for ln in lines:
        if re.search(r"(?i)www\.|help@uidai|virtual\s+id\b", ln):
            break
        cut.append(ln)
    lines = _trim_aadhar_letter_to_block_noise(cut)
    if len(lines) < 1:
        return None
    if _aadhar_back_name_line_matches_hint(lines[0], name_hint):
        lines = lines[1:]
    elif _looks_like_aadhar_letter_printed_name_line(lines[0]):
        lines = lines[1:]
    lines = [ln for ln in lines if not _is_aadhar_back_date_only_line(ln)]
    if lines:
        last_compact = re.sub(r"\D", "", lines[-1])
        if len(last_compact) == 10 and last_compact[0] in "6789":
            lines = lines[:-1]
    if not lines:
        return None
    raw = "\n".join(lines).strip()
    if len(raw) < 12:
        return None
    return raw


def _parse_aadhar_back_address_from_ocr(ocr_text: str, name_hint: str | None = None) -> dict[str, str]:
    """
    English (and noisy OCR) address on Aadhaar back.

    **Order:** (1) Card-style — ``Address:``, ``C/O:`` …, ``Address`` newline + block, ``Near …``.
    (2) **Last resort only:** UIDAI **letter** printout (``To,`` / enrollment / name / address / date / PIN / mobile)
    via :func:`_parse_aadhar_back_letter_format_address` when (1) did not yield a usable address, or after
    trimming/footer the text is still too short. Pass ``name_hint`` (holder name from front) to drop the
    duplicate name line under ``To,``. DOB/Gender are not printed on the back; use **Aadhar.jpg** (front).
    """
    out: dict[str, str] = {}
    if not ocr_text or len(ocr_text.strip()) < 8:
        return out
    text = ocr_text.replace("\r\n", "\n")
    raw = ""
    # Do not stop at ``C/O:`` on the next line — UIDAI layout is ``Address:`` then ``C/O:`` on the following line.
    m = re.search(
        r"(?is)\bAddress\s*:\s*([\s\S]+?)(?=\n\s*पता\b|\n\s*VID\s*:|\n\s*Virtual\s+ID\b|\Z)",
        text,
    )
    if m:
        raw = m.group(1).strip()
    if not raw or len(raw) < 12:
        so_iter = re.finditer(r"(?is)\b(?:C/O|S/O|W/O|D/O)\s*:", text)
        best_chunk = ""
        for co in so_iter:
            tail = text[co.start() :]
            stop = re.search(r"\d{4}\s+\d{4}\s+\d{4}", tail)
            chunk = tail[: stop.start()] if stop else tail
            stop_vid = re.search(r"(?i)\bVID\s*:", chunk)
            if stop_vid:
                chunk = chunk[: stop_vid.start()]
            chunk = chunk.strip()
            if re.search(r"(?i)\b(?:PO|DIST|District|Tehsil|Post\s*Office)\s*:", chunk):
                best_chunk = chunk
                break
            if not best_chunk or len(chunk) > len(best_chunk):
                best_chunk = chunk
        if best_chunk:
            raw = best_chunk
    if not raw or len(raw) < 8:
        # Textract often puts "Address" on its own line, then "Near …" / locality on the next lines.
        addr_head = re.search(r"(?is)\bAddress\s*\n+", text)
        if addr_head:
            window = text[addr_head.end() : addr_head.end() + 950]
            block_m = re.search(
                r"(?is)^\s*((?:Near|C/O|C\.?\s*O\.?|S/O|W/O|D/O)\s*:?.+?)(?=\n\s*(?:पता|VID|Virtual|Aadhaar|www\.|help@|\d{4}\s+\d{4}\s+\d{4})\b|\n{3,}|\Z)",
                window,
            )
            if block_m:
                raw = block_m.group(1).strip()
        if (not raw or len(raw) < 8) and re.search(r"(?i)\bNear\b", text):
            near_m = re.search(
                r"(?is)\b(Near\s+.{15,350}?(?:Rajasthan|State|,?\s*[A-Za-z]{4,}\s*-\s*\d{6}|\d{6}))",
                text,
            )
            if near_m:
                raw = near_m.group(1).strip()

    used_letter = False
    if not raw or len(raw.strip()) < 12:
        letter = _parse_aadhar_back_letter_format_address(text, name_hint=name_hint)
        if letter:
            raw = letter
            used_letter = True

    if not raw:
        return out
    raw = _normalize_aadhar_back_address_chunk(raw)
    foot = re.split(
        r"(?i)(www\.uidai|help@uidai|unique\s+identification|virtual\s+id)",
        raw,
        maxsplit=1,
    )
    raw = foot[0].strip(" ,.-") if foot else raw
    if len(raw) < 15 and not used_letter:
        letter = _parse_aadhar_back_letter_format_address(text, name_hint=name_hint)
        if letter:
            raw = _normalize_aadhar_back_address_chunk(letter)
            foot2 = re.split(
                r"(?i)(www\.uidai|help@uidai|unique\s+identification|virtual\s+id)",
                raw,
                maxsplit=1,
            )
            raw = foot2[0].strip(" ,.-") if foot2 else raw
    if len(raw) < 15:
        return out
    parsed = normalize_address_freeform(raw)
    out["address"] = _clean_aadhar_back_cross_column_noise(
        parsed.get("address") or re.sub(r"\s+", " ", raw).strip()
    )
    if parsed.get("care_of"):
        out["care_of"] = parsed["care_of"]
    if parsed.get("pin_code"):
        out["pin_code"] = parsed["pin_code"]
    if parsed.get("state"):
        out["state"] = parsed["state"]
    if parsed.get("city"):
        out["city"] = parsed["city"]
    if parsed.get("district"):
        out["district"] = parsed["district"]
    return out


def _aadhar_normalize_dob_triplet(day: int, month: int, year: int) -> str | None:
    """Validate and return DD/MM/YYYY for Indian Aadhaar-style dates."""
    if not (1 <= month <= 12 and 1 <= day <= 31):
        return None
    if year < 100:
        year = 2000 + year if year < 50 else 1900 + year
    y_max = datetime.now().year
    if year < 1920 or year > y_max:
        return None
    try:
        datetime(year, month, day)
    except ValueError:
        return None
    return f"{day:02d}/{month:02d}/{year}"


def _normalize_aadhar_gender_token(token: str) -> str | None:
    t = (token or "").strip()
    if not t:
        return None
    u = t.upper()
    if u in ("M", "MALE"):
        return "Male"
    if u in ("F", "FEMALE"):
        return "Female"
    if u in ("T", "TRANSGENDER"):
        return "Transgender"
    if t.lower() in ("male", "female", "transgender"):
        return t[:1].upper() + t[1:].lower()
    return None


def _aadhar_dob_window_not_issue_date(text: str, date_start_idx: int) -> bool:
    """True if the date at ``date_start_idx`` is unlikely to be ``Aadhaar … issued:``."""
    window = text[max(0, date_start_idx - 35) : date_start_idx]
    return not re.search(r"(?i)issued|issue\s*date|date\s+of\s+issue", window)


def _gender_skip_word_then_slash_token(suffix: str) -> str | None:
    """
    UIDAI front layout (OCR): after ``dd/mm/yyyy``, skip the next token, then the next ``/``
    introduces gender (e.g. ``yes/ MALE`` where ``Sex`` OCR'd as ``yes``).
    """
    s = suffix.lstrip()
    if not s:
        return None
    if s.startswith("/"):
        after = s[1:].lstrip()
        m_g = re.match(r"\S+", after)
        return _normalize_aadhar_gender_token(m_g.group(0)) if m_g else None
    m_skip = re.match(r"\S+", s)
    if m_skip:
        s = s[m_skip.end() :]
    s = s.lstrip()
    idx = s.find("/")
    if idx < 0:
        return None
    after = s[idx + 1 :].lstrip()
    m_g = re.match(r"\S+", after)
    if not m_g:
        return None
    return _normalize_aadhar_gender_token(m_g.group(0))


def _extract_gender_using_dob_slash_rule(full_text: str, normalized_dob: str | None) -> str | None:
    """
    Locate DOB (``dd/mm/yyyy``), skip the next word, find the next ``/``, read gender after it.
    Avoids anchoring on ``issued: dd/mm/yyyy`` when possible.
    """
    t = full_text.replace("\r\n", "\n")
    end_pos: int | None = None
    if normalized_dob and re.match(r"^\d{1,2}/\d{1,2}/\d{4}$", normalized_dob.strip()):
        p = normalized_dob.strip().split("/")
        d, mo, yr = int(p[0]), int(p[1]), p[2]
        variants = (
            rf"(?<!\d){d}/{mo}/{yr}(?!\d)",
            rf"(?<!\d){d:02d}/{mo:02d}/{yr}(?!\d)",
            rf"(?<!\d){d}/{mo:02d}/{yr}(?!\d)",
            rf"(?<!\d){d:02d}/{mo}/{yr}(?!\d)",
        )
        for rx in variants:
            for m in re.finditer(rx, t):
                if _aadhar_dob_window_not_issue_date(t, m.start()):
                    end_pos = m.end()
                    break
            if end_pos is not None:
                break
    if end_pos is None:
        for m in re.finditer(
            r"(?<!\d)(\d{1,2})[/.\-](\d{1,2})[/.\-]((?:19|20)\d{2})(?!\d)",
            t,
        ):
            if not _aadhar_dob_window_not_issue_date(t, m.start()):
                continue
            day, month, yr_s = int(m.group(1)), int(m.group(2)), int(m.group(3))
            norm = _aadhar_normalize_dob_triplet(day, month, yr_s)
            if norm:
                end_pos = m.end()
                break
    if end_pos is None:
        return None
    return _gender_skip_word_then_slash_token(t[end_pos:])


def _parse_aadhar_front_textract_fallback(text: str) -> dict[str, str]:
    """
    Pull DOB / gender from Aadhaar **front** full text (AWS Textract or same layout in Raw_OCR).
    Used when printed front text omits these fields.
    """
    out: dict[str, str] = {}
    if not text or len(text.strip()) < 5:
        return out
    t = text.replace("\r\n", "\n")

    dob_patterns = [
        r"(?i)\b(?:DOB|D\.?\s*O\.?\s*B\.?|Date\s+of\s+Birth|Birth\s+Date)\s*[:]?\s*(\d{1,2})[/.\-](\d{1,2})[/.\-](\d{2,4})\b",
        r"(?i)\b(?:जन्म\s*तिथि|जन्मतिथि)\s*[:]?\s*(\d{1,2})[/.\-](\d{1,2})[/.\-](\d{2,4})\b",
        # Mis-OCR: "JoH /DB:01/01/2004" — DOB read as DB; slash before DB is common.
        r"(?i)/\s*D\.?B\.?\s*[:]?\s*(\d{1,2})[/.\-](\d{1,2})[/.\-](\d{2,4})\b",
        # Standalone "DB:" or "D.B:" when followed immediately by a slash-date (avoid matching "db" in words).
        r"(?i)(?<![A-Za-z])D\.?B\.?\s*[:]?\s*(\d{1,2})[/.\-](\d{1,2})[/.\-](\d{2,4})\b",
    ]
    for pat in dob_patterns:
        m = re.search(pat, t)
        if m:
            d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
            norm = _aadhar_normalize_dob_triplet(d, mo, y)
            if norm:
                out["date_of_birth"] = norm
                out["year_of_birth"] = norm.split("/")[2]
                break

    if "date_of_birth" not in out:
        # Prefer dd/mm/yyyy (two slashes) near DOB/DB markers over stray numeric runs.
        dob_marker = re.compile(r"(?i)(?:\b(?:dob|date\s+of\s+birth|birth\s+date|जन्म)|/(?:dob|d\.?b\.?)\s*:?|\bd\.?b\.?\s*:)")
        triplet_re = re.compile(r"(?<!\d)(\d{1,2})/(\d{1,2})/((?:19|20)\d{2})(?!\d)")
        best: tuple[int, tuple[int, int, int]] | None = None
        for m in triplet_re.finditer(t):
            d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
            norm = _aadhar_normalize_dob_triplet(d, mo, y)
            if not norm:
                continue
            start, end = m.start(), m.end()
            window_start = max(0, start - 48)
            window = t[window_start:end]
            if not dob_marker.search(window):
                continue
            slash_bonus = window.count("/")
            dist = start - window_start
            score = slash_bonus * 10 - dist
            if best is None or score > best[0]:
                best = (score, (d, mo, y))
        if best:
            d, mo, y = best[1]
            norm = _aadhar_normalize_dob_triplet(d, mo, y)
            if norm:
                out["date_of_birth"] = norm
                out["year_of_birth"] = str(y)

    if "date_of_birth" not in out:
        for line in t.splitlines():
            line = line.strip()
            if len(line) > 72:
                continue
            if not re.search(r"(?i)(birth|dob|db\s*:|/(?:dob|d\.?b\.?)|जन्म)", line):
                continue
            m = re.search(r"(\d{1,2})[/.\-](\d{1,2})[/.\-]((19|20)\d{2})\b", line)
            if m:
                d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
                norm = _aadhar_normalize_dob_triplet(d, mo, y)
                if norm:
                    out["date_of_birth"] = norm
                    out["year_of_birth"] = str(y)
                    break

    if "date_of_birth" not in out:
        # Aadhaar letter: only one date is sandwiched between name and MALE/FEMALE.
        # Find the last valid date before the gender token, excluding issue dates.
        gender_pos = re.search(r"(?i)\b(MALE|FEMALE|Transgender)\b", t)
        if gender_pos:
            before_gender = t[: gender_pos.start()]
            _triplet = re.compile(r"(?<!\d)(\d{1,2})[/.\-](\d{1,2})[/.\-]((?:19|20)\d{2})(?!\d)")
            for m in reversed(list(_triplet.finditer(before_gender))):
                if not _aadhar_dob_window_not_issue_date(t, m.start()):
                    continue
                d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
                norm = _aadhar_normalize_dob_triplet(d, mo, y)
                if norm:
                    out["date_of_birth"] = norm
                    out["year_of_birth"] = str(y)
                    break

    # Primary gender from layout: anchor on DOB, skip next token, next "/" then gender token.
    g_dob = _extract_gender_using_dob_slash_rule(t, out.get("date_of_birth"))
    if g_dob:
        out["gender"] = g_dob

    if "gender" not in out:
        gm = (
            re.search(r"(?i)\bGender\s*[:]?\s*(Male|Female|Transgender|M|F|T)\b", t)
            or re.search(r"(?i)\bGender\s+(Male|Female|Transgender)\b", t)
            or re.search(r"(?i)\b(?:लिंग|Gender)\s*[:]?\s*(Male|Female|Transgender)\b", t)
        )
        if gm:
            g = _normalize_aadhar_gender_token(gm.group(1))
            if g:
                out["gender"] = g
        else:
            for label_pat in (r"(?i)\bGender\b", r"(?i)\bलिंग\b"):
                pos = re.search(label_pat, t)
                if pos:
                    window = t[pos.end() : pos.end() + 40]
                    for word in ("Male", "Female", "Transgender"):
                        if re.search(rf"(?i)\b{word}\b", window):
                            out["gender"] = word
                            break
                    if out.get("gender"):
                        break

    if "gender" not in out:
        # UIDAI front prints "पुरुष/ MALE" or "Sex / Male". OCR garbles the Hindi/label
        # into "yes", "you", "yow", "yeu", "puu", etc. before "/ MALE".
        slash_g = re.search(
            r"(?i)\b(?:yes|yex|yos|ses|sex|you|yow|yeu|yoo|puu|pur)\s*/\s*(Male|Female|Transgender|MALE|FEMALE|M|F|T)\b",
            t,
        ) or re.search(
            r"(?i)\bSex\s*[/:]\s*(Male|Female|Transgender|MALE|FEMALE|M|F|T)\b",
            t,
        ) or re.search(
            r"(?i)\b\w{2,6}\s*/\s*(Male|Female|Transgender|MALE|FEMALE)\b",
            t,
        )
        if slash_g:
            g = _normalize_aadhar_gender_token(slash_g.group(1))
            if g:
                out["gender"] = g

    return out


_AADHAR_TEXTRACT_NAME_TITLE_WORD = re.compile(r"^[A-Z][a-z]+(?:['-][A-Z][a-z]+)?$")


def _aadhaar_name_line_score_for_textract(line: str, *, skip: re.Pattern[str], noise: re.Pattern[str]) -> int:
    """
    Score a single OCR line as a likely holder name. Higher is better; -1 = reject.
    Favors two (or more) Latin title-case words (e.g. "Nishant Kumar") over one
    all-lowercase garbage token (e.g. misread "amit") or mixed OCR ("Prints gHR").
    """
    line = line.strip()
    if len(line) < 2 or len(line) > 80:
        return -1
    if skip.search(line) or noise.search(line):
        return -1
    if re.search(r"\d", line):
        return -1
    if not re.match(r"^[A-Za-z][A-Za-z\s.'-]{1,70}$", line):
        return -1
    words = [w for w in line.split() if w.strip()]
    if not words or len(words) > 6:
        return -1
    tc = sum(1 for w in words if _AADHAR_TEXTRACT_NAME_TITLE_WORD.match(w))
    if tc >= 2:
        return 100 + tc
    if tc == 1 and len(words) >= 2:
        return 25
    if tc == 1 and len(words) == 1:
        return 30
    if all(w.islower() for w in words):
        return 2
    return 10


def _pick_best_aadhaar_name_from_textract_lines(
    lines_slice: list[str],
    *,
    skip: re.Pattern[str],
    noise: re.Pattern[str],
    min_score: int = 30,
) -> str | None:
    best_line: str | None = None
    best_sc = -1
    for line in lines_slice:
        sc = _aadhaar_name_line_score_for_textract(line, skip=skip, noise=noise)
        if sc < 0:
            continue
        if sc > best_sc:
            best_sc = sc
            best_line = line.strip()
    if best_line is not None and best_sc >= min_score:
        return best_line
    return None


def _parse_aadhar_name_from_aadhaar_textract(text: str) -> dict[str, str]:
    """
    Heuristic name line on Aadhaar **front** when Textract text has no clear name field.
    Avoids obvious headers / labels; prefers a Title Case line before DOB.
    """
    out: dict[str, str] = {}
    if not text or len(text.strip()) < 8:
        return out
    lines = [ln.strip() for ln in text.replace("\r\n", "\n").split("\n") if ln.strip()]
    skip = re.compile(
        r"(?i)aadhaar|aadhar|enrol|enrollment|government|india|identification|male|female|transgender|"
        r"address|dob|date\s*of\s*birth|year\s*of\s*birth|vid|virtual|help@|uidai|www\.|pin"
    )
    dob_line = re.compile(r"(?i)\b(dob|date\s*of\s*birth|जन्म|birth)\b")
    gov_line = re.compile(r"(?i)government\s+of\s+india")
    noise = re.compile(r"(?i)\b(famoy|family|service|assess|authority|unique)\b")
    candidates: list[str] = []
    issued_re = re.compile(r"(?i)aadhaar\s+no\.?\s*\.?\s*issued")

    gov_idx = next((i for i, line in enumerate(lines[:30]) if gov_line.search(line)), -1)
    issued_idx = next((i for i, line in enumerate(lines[:30]) if issued_re.search(line)), -1)

    # Letter layout: holder name sits between "Government of India" and "Aadhaar no. issued:".
    # Do not take the *first* alphabetic line after Gov (OCR often inserts garbage like "amitt").
    if gov_idx >= 0 and issued_idx > gov_idx:
        picked = _pick_best_aadhaar_name_from_textract_lines(
            lines[gov_idx + 1 : issued_idx],
            skip=skip,
            noise=noise,
            min_score=30,
        )
        if picked:
            out["name"] = picked
            return out

    # Name sometimes appears on the first clean line after "Aadhaar no. issued" (alternate layouts).
    for i, line in enumerate(lines[:25]):
        if issued_re.search(line):
            picked = _pick_best_aadhaar_name_from_textract_lines(
                lines[i + 1 : i + 6],
                skip=skip,
                noise=noise,
                min_score=30,
            )
            if picked:
                out["name"] = picked
                return out
            break

    if gov_idx >= 0:
        picked = _pick_best_aadhaar_name_from_textract_lines(
            lines[gov_idx + 1 : min(gov_idx + 12, len(lines))],
            skip=skip,
            noise=noise,
            min_score=30,
        )
        if picked:
            out["name"] = picked
            return out

    for line in lines[:22]:
        if len(line) < 4 or len(line) > 90:
            continue
        if skip.search(line):
            continue
        if noise.search(line):
            continue
        if dob_line.search(line):
            break
        if re.search(r"\d{4}\s+\d{4}\s+\d{4}", line):
            continue
        if re.match(r"^[\d\s/:\-.]+$", line):
            continue
        if re.match(r"^[A-Za-z][A-Za-z\s.'-]{2,70}$", line):
            words = line.split()
            initials = sum(1 for w in words if w and w[0].isupper())
            if 1 <= len(words) <= 6 and initials >= 1:
                # Keep candidates and prefer the nearest one before DOB.
                candidates.append(line)
    if candidates:
        out["name"] = candidates[-1]
    return out


def _aadhar_identity_ok(customer: dict[str, str]) -> bool:
    name = (customer.get("name") or "").strip()
    if len(name) >= 2:
        return True
    aid = "".join(c for c in str(customer.get("aadhar_id") or "") if c.isdigit())
    return len(aid) >= 4


def _aadhar_geo_ok(customer: dict[str, str]) -> bool:
    pin = re.sub(r"\D", "", str(customer.get("pin_code") or ""))
    if len(pin) >= 6:
        return True
    st = (customer.get("state") or "").strip()
    if len(st) >= 3:
        return True
    addr = (customer.get("address") or "").strip()
    if len(addr) >= 22:
        return True
    if (customer.get("district") or customer.get("city")) and len(pin) >= 6:
        return True
    return False


def _pipeline_merge_aadhar_customer(
    front_bytes: bytes | None,
    back_bytes: bytes | None,
    front_textract: dict | None,
    back_textract: dict | None,
    *,
    front_raw_name: str = FILENAME_AADHAR_FRONT,
    back_raw_name: str = "Aadhar_back.jpg",
) -> tuple[dict[str, str], list[tuple[str, str]], str | None, dict[str, int]]:
    """
    **AWS Textract only** on front/back (no Tesseract, no UIDAI QR in this pipeline).
    """
    from app.services.sales_textract_service import extract_text_from_bytes

    timings: dict[str, int] = {
        "aadhar_textract_front_ms": 0,
        "aadhar_textract_back_ms": 0,
    }

    def _ftext(d: dict | None) -> str:
        if not d or d.get("error"):
            return ""
        return (d.get("full_text") or "").strip()

    raw_parts: list[tuple[str, str]] = []
    customer: dict[str, str] = {}

    # Physical layout (letter / consolidated) is normalized in pre-OCR or
    # :func:`normalize_aadhar_upload_files` before this pipeline runs.

    ft = front_textract
    if front_bytes and (not ft or ft.get("error") or not _ftext(ft)):
        t_tx = time.perf_counter()
        try:
            ft = extract_text_from_bytes(front_bytes)
        except Exception:
            ft = {"error": "textract", "full_text": ""}
        timings["aadhar_textract_front_ms"] += int((time.perf_counter() - t_tx) * 1000)
    front_txt = _ftext(ft)
    if front_txt:
        raw_parts.append((front_raw_name, front_txt))
        customer = _merge_aadhar_textract_fallback_dict(
            customer, _parse_aadhar_front_textract_fallback(front_txt)
        )
        customer = _merge_aadhar_textract_fallback_dict(
            customer, _parse_aadhar_name_from_aadhaar_textract(front_txt)
        )

    bt = back_textract
    if back_bytes and not _aadhar_geo_ok(customer):
        if not bt or bt.get("error") or not _ftext(bt):
            t_tx = time.perf_counter()
            try:
                bt = extract_text_from_bytes(back_bytes)
            except Exception:
                bt = {"error": "textract", "full_text": ""}
            timings["aadhar_textract_back_ms"] += int((time.perf_counter() - t_tx) * 1000)
        back_txt = _ftext(bt)
        if back_txt:
            raw_parts.append((back_raw_name, back_txt))
            customer = _merge_aadhar_textract_fallback_dict(
                customer,
                _parse_aadhar_back_address_from_ocr(
                    back_txt,
                    name_hint=(customer.get("name") or "").strip() or None,
                ),
            )

    note: str | None = None
    if not _aadhar_identity_ok(customer):
        note = (
            "Aadhaar automated read did not yield a clear name or ID number. "
            "Enter or verify customer fields manually."
        )

    if customer.get("aadhar_id"):
        customer["aadhar_id"] = _aadhar_last4(customer["aadhar_id"]) or ""
    _default_gender_male_if_unread(customer)
    customer = enrich_customer_address_from_freeform(customer)
    return customer, raw_parts, note, timings


def _merge_granular_upload_process_timings(
    section_timings_ms: dict[str, int],
    *,
    prefetch_job_ms: dict[str, int],
    frag_a: dict[str, Any] | None,
    frag_d: dict[str, Any] | None,
) -> None:
    """Populate per-process durations (prefetch job wall times + sync pipeline/FORMS times)."""
    pt = (frag_a or {}).get("process_timings") or {}
    section_timings_ms["aadhar_textract_front_ms"] = int(prefetch_job_ms.get("aadhar_front", 0)) + int(
        pt.get("aadhar_textract_front_ms", 0)
    )
    section_timings_ms["aadhar_textract_back_ms"] = int(prefetch_job_ms.get("aadhar_back", 0)) + int(
        pt.get("aadhar_textract_back_ms", 0)
    )
    dsync = (
        int((frag_d or {}).get("detail_sheet_textract_sync_ms", 0))
        if frag_d and frag_d.get("ok")
        else 0
    )
    section_timings_ms["detail_sheet_textract_ms"] = int(prefetch_job_ms.get("details_forms", 0)) + dsync


def _compile_details_sheet_fragment(
    input_path: Path,
    textract_forms_prefetch: dict | None,
) -> dict:
    """
    Parse Sales Detail sheet to structured fragments only (no JSON write).
    Raster/PDF: Textract **AnalyzeDocument** FORMS+TABLES — right tool for printed form key-values.
    """
    fmt = _details_input_format(input_path)
    detail_sheet_textract_sync_ms = 0
    if fmt == "docx":
        key_value_pairs, docx_full = _key_value_pairs_from_docx(input_path)
        if not key_value_pairs and not (docx_full or "").strip():
            raise RuntimeError(
                "Could not read the Word document (.docx). Ensure it is a valid .docx Sales Detail Sheet."
            )
        docx_full = _full_text_from_sales_detail_sheet_heading(docx_full or "")
        result = {"error": None, "full_text": docx_full, "key_value_pairs": key_value_pairs}
    elif fmt in ("jpeg", "png", "pdf"):
        from app.services.sales_textract_service import extract_forms_from_bytes

        if textract_forms_prefetch is not None:
            result = textract_forms_prefetch
        else:
            t_df = time.perf_counter()
            result = extract_forms_from_bytes(input_path.read_bytes())
            detail_sheet_textract_sync_ms = int((time.perf_counter() - t_df) * 1000)
        if result.get("error"):
            raise RuntimeError(str(result.get("error")))
    else:
        raise RuntimeError(
            f"Unsupported Details file format (detected={fmt!r}). "
            "Use a JPEG/PNG scan, PDF export, or .docx Sales Detail Sheet."
        )

    key_value_pairs = result.get("key_value_pairs") or []
    vehicle = _map_key_value_pairs_to_vehicle(key_value_pairs)
    insurance = _map_key_value_pairs_to_insurance(key_value_pairs)
    details_customer = _map_key_value_pairs_to_details_customer(key_value_pairs)
    details_customer_name = _extract_details_customer_name(key_value_pairs)
    if result.get("full_text") or result.get("tables"):
        from_vehicle = _parse_vehicle_from_full_text(result["full_text"] or "")
        for k, v in from_vehicle.items():
            if v and not vehicle.get(k):
                vehicle[k] = v
        extra = _merge_textract_details_fallbacks(
            insurance,
            details_customer,
            full_text=result.get("full_text"),
            tables=result.get("tables"),
        )
        if extra.get("customer_name") and not details_customer_name:
            details_customer_name = extra["customer_name"]
    if details_customer.get("name") and not details_customer_name:
        details_customer_name = details_customer.get("name")

    _sync_nominee_relation_with_gender_across_fragments(insurance, details_customer)

    return {
        "vehicle": vehicle,
        "insurance": insurance,
        "details_customer": details_customer,
        "details_customer_name": details_customer_name,
        "full_text": result.get("full_text") or "",
        "detail_sheet_textract_sync_ms": detail_sheet_textract_sync_ms,
    }


def _default_gender_male_if_unread(customer: dict[str, str]) -> None:
    """If gender was not extracted (Textract/parsed text), assume Male for DMS/insurance flows."""
    g = customer.get("gender")
    if g is None or not str(g).strip():
        customer["gender"] = "Male"


def _merge_aadhar_textract_fallback_dict(customer: dict[str, str], hints: dict[str, str]) -> dict[str, str]:
    """
    Fill blank customer fields from Textract-derived hints. For ``address``, prefer a longer /
    clearer Textract line over a short or empty prior value (common for back-of-card English).
    """
    out = dict(customer) if customer else {}
    if not hints:
        return out
    for k in AADHAR_QR_CUSTOMER_KEYS:
        if k == "address":
            continue
        hv = hints.get(k)
        if not hv or not str(hv).strip():
            continue
        cur = out.get(k)
        if cur is None or not str(cur).strip():
            out[k] = str(hv).strip()
    ha = (hints.get("address") or "").strip()
    ca = (out.get("address") or "").strip()
    if ha:
        if not ca:
            out["address"] = ha
        elif len(ha) > len(ca) + 12:
            out["address"] = ha
    _rebuild_aadhar_address_from_parts(out, preserve_existing_address=True)
    if ha and (not out.get("address") or not str(out["address"]).strip()):
        out["address"] = ha
    if hints.get("pin_code") and not (out.get("pin_code") and str(out["pin_code"]).strip()):
        out["pin_code"] = str(hints["pin_code"]).strip()
    if hints.get("care_of") and not (out.get("care_of") and str(out["care_of"]).strip()):
        out["care_of"] = str(hints["care_of"]).strip()
    return out


def _parse_raw_ocr_sections(content: str) -> dict[str, str]:
    """Split ``Raw_OCR.txt`` (``--- filename ---`` sections) into filename -> body."""
    sections: dict[str, str] = {}
    current: str | None = None
    buf: list[str] = []
    for line in content.splitlines():
        m = re.match(r"^---\s*(.+?)\s*---\s*$", line.strip())
        if m:
            if current is not None:
                sections[current] = "\n".join(buf).strip()
            current = m.group(1).strip()
            buf = []
        else:
            buf.append(line)
    if current is not None:
        sections[current] = "\n".join(buf).strip()
    return sections


def _apply_aadhar_textract_fallbacks_from_parts(
    ocr_output_dir: Path,
    subfolder: str,
    parts: list[tuple[str, str]],
) -> None:
    """Merge DOB/gender (front) and address (back) from Textract blobs into OCR_To_be_Used.json."""
    json_path = _json_output_path(ocr_output_dir, subfolder)
    if not json_path.is_file():
        return
    try:
        data = json.loads(json_path.read_text(encoding="utf-8"))
    except Exception:
        return
    customer = data.get("customer") or {}
    if not isinstance(customer, dict):
        customer = {}

    def _aadhar_section_order(fn: str) -> int:
        fl = fn.strip().replace("\\", "/").split("/")[-1].lower()
        if fl in ("aadhar.jpg", "aadhar.pdf"):
            return 0
        if fl in ("aadhar_back.jpg", "aadhar_back.pdf"):
            return 1
        return 2

    for fn, tx in sorted(parts, key=lambda p: _aadhar_section_order(p[0])):
        if not tx or not str(tx).strip():
            continue
        fl = fn.strip().replace("\\", "/").split("/")[-1].lower()
        if fl in ("aadhar.jpg", "aadhar.pdf"):
            customer = _merge_aadhar_textract_fallback_dict(
                customer, _parse_aadhar_front_textract_fallback(tx)
            )
        elif fl in ("aadhar_back.jpg", "aadhar_back.pdf"):
            hint = (customer.get("name") or "").strip() or None
            customer = _merge_aadhar_textract_fallback_dict(
                customer, _parse_aadhar_back_address_from_ocr(tx, name_hint=hint)
            )
    if customer.get("aadhar_id"):
        customer["aadhar_id"] = _aadhar_last4(customer["aadhar_id"]) or ""
    _default_gender_male_if_unread(customer)
    customer = enrich_customer_address_from_freeform(customer)

    aadhar_parts = [
        (fn, tx)
        for fn, tx in parts
        if "aadhar" in (fn or "").strip().replace("\\", "/").split("/")[-1].lower()
    ]
    blob = _concat_aadhar_scan_ocr_text(aadhar_parts)
    dn = data.get("details_customer_name")
    if dn and blob.strip():
        reconciled = _reconcile_customer_name_aadhar_details(customer.get("name"), str(dn), blob)
        if reconciled:
            customer["name"] = reconciled

    data["customer"] = customer
    try:
        json_path.write_text(json.dumps(data, indent=2), encoding="utf-8")
    except Exception:
        pass


def _apply_aadhar_textract_fallbacks_from_raw_ocr_file(
    ocr_output_dir: Path,
    subfolder: str,
) -> None:
    """Same as parts merge, using persisted Raw_OCR.txt (e.g. get_extracted_details)."""
    raw_path = _ocr_subfolder_path(ocr_output_dir, subfolder) / "Raw_OCR.txt"
    if not raw_path.is_file():
        return
    try:
        sections = _parse_raw_ocr_sections(raw_path.read_text(encoding="utf-8", errors="replace"))
    except Exception:
        return
    parts = [(k, v) for k, v in sections.items() if v]
    if parts:
        _apply_aadhar_textract_fallbacks_from_parts(ocr_output_dir, subfolder, parts)


# Map Textract form keys (normalized) to our vehicle detail fields.
_VEHICLE_KEY_ALIASES = {
    # Many dealer PDFs use "Chassis Number" / "Engine Number" (not "Chassis No.")
    "frame_no": [
        "frame no",
        "frame no.",
        "frame number",
        "chassis",
        "chassis no",
        "chassis no.",
        "chassis number",
        "vin",
        "vehicle identification number",
    ],
    "engine_no": ["engine no", "engine no.", "engine number", "engine"],
    "model_colour": ["model & colour", "model and colour", "model/colour", "model", "colour", "color"],
    "key_no": ["key no", "key no.", "key number", "key"],
    "battery_no": ["battery no", "battery no.", "battery number", "battery"],
}

# Map Textract form keys to insurance/nominee fields (from Details sheet).
_INSURANCE_KEY_ALIASES = {
    "profession": [
        "profession",
        "profession:",
        "occupation",
        "customer profession",
        "employment",
        "nature of occupation",
        "job",
        "customer occupation",
    ],
    "financier": [
        "financier",
        "financier name",
        "finance company",
        "bank",
        "financer",
        "financing bank",
        "name of financier",
        "bank / financier",
        "financier / bank",
        "finance bank",
        "hypothecation",
        "loan from",
        "financing institution",
    ],
    "insurer": [
        "insurer name (if needed)",
        "insurer name if needed",
        "insurer",
        "insurer name",
        "name of insurer",
        "insurance company",
        "insurance provider",
        # Do not use bare "company" — matches unrelated labels (e.g. finance company).
    ],
    "marital_status": [
        "marital status",
        "customer marital status",
        "married status",
        "martial status",
        "customer martial status",
    ],
    "nominee_gender": ["nominee gender", "gender of nominee", "sex of nominee", "nominee sex"],
    "nominee_name": [
        "nominee name",
        "nominee name:",
        "nominee",
        "name of nominee",
        "name of the nominee",
        "nominee's name",
        "insurance nominee",
        "nominee details",
    ],
    "nominee_age": [
        "nominee age",
        "nominee age:",
        "age of nominee",
        "nominee age (years)",
        "age (years)",
    ],
    "nominee_relationship": [
        "nominee relationship",
        "nominee relationship:",
        "relationship with customer",
        "nominee relationship with customer",
        "relation with proposer",
        "relation to insured",
        "nominee relation",
        "relationship of nominee",
        "relationship",
        "relation",
    ],
    "payment_mode": [
        "payment",
        "payment:",
        "payment mode",
        "mode of payment",
        "payment type",
    ],
}

# Map Textract form keys to customer name on Details sheet.
_DETAILS_CUSTOMER_NAME_ALIASES = [
    "customer name", "customer name:", "buyer name", "buyer's name",
    "buyer name:", "name of customer",
    "full name", "full name:", "customer full name",
]

_DETAILS_CUSTOMER_KEY_ALIASES = {
    "name": ["full name", "customer name", "name of customer", "buyer name", "name"],
    "mobile_number": ["mobile number", "mobile no", "mobile", "contact number", "contact no"],
    "alt_phone_num": ["alternate", "alternate no", "alternate number", "alternate mobile", "alternate mobile number", "landline", "landline number"],
    "aadhar_id": ["aadhaar number", "aadhar number", "aadhaar no", "aadhar no", "aadhaar", "aadhar"],
    "profession": [
        "profession",
        "occupation",
        "customer profession",
        "employment",
        "nature of occupation",
        "customer occupation",
    ],
    "marital_status": [
        "marital status",
        "customer marital status",
        "martial status",
        "customer martial status",
    ],
    "financier": [
        "financier name",
        "financier",
        "financer name",
        "financer",
        "finance company",
        "bank",
        "financing bank",
        "name of financier",
        "bank / financier",
    ],
    "nominee_name": ["nominee name", "name of nominee", "name of the nominee", "nominee's name"],
    "nominee_age": ["nominee age", "age of nominee", "age (years)"],
    "nominee_gender": ["nominee gender", "gender of nominee", "sex of nominee"],
    "nominee_relationship": [
        "nominee relationship",
        "nominee relation",
        "relationship with customer",
        "nominee relationship with customer",
        "relation with proposer",
        "relationship of nominee",
        "relationship",
        "relation",
    ],
    "payment_mode": [
        "payment",
        "payment mode",
        "mode of payment",
        "payment type",
    ],
}

# Map Textract form keys for insurance policy document (Insurance.jpg).
_INSURANCE_POLICY_KEY_ALIASES = {
    "insurer": ["insurer", "insurance company", "company", "insurance provider", "name of insurer", "national insurance"],
    "policy_num": ["policy no", "policy no.", "policy number", "policy num", "cert. no", "cert no", "certificate no"],
    "policy_from": ["tp valid from", "valid from", "validity from", "policy from", "from date", "tp valid"],
    "policy_to": ["tp valid to", "valid to", "validity to", "policy to", "to date", "midnight of"],
    "premium": ["gross premium", "premium", "total premium", "gross premium amount", "premium of rs"],
}


def _normalize_key_for_match(key: str) -> str:
    """Lowercase and collapse spaces for key matching."""
    return re.sub(r"\s+", " ", (key or "").lower().strip())


# Printed Details sheets: **handwritten or typed words** for Profession, Marital status, Nominee (fuzzy ≥ 0.5);
# legacy checkbox / tick / **filled-box** rows still supported via ``_normalize_kv_value_for_checkbox_fields`` and
# ``_normalize_filled_box_marks_to_selected_token``.
_WRITTEN_PROFESSION_OPTIONS = ("Private Job", "Employed", "Business", "Farmer")

_FIELD_CHECKBOX_ALIASES: dict[str, list[tuple[str, str]]] = {
    "profession": [
        ("private job", "Private Job"),
        ("private", "Private Job"),
        ("employed", "Employed"),
        ("business", "Business"),
        ("farmer", "Farmer"),
    ],
    "marital_status": [
        ("married", "Married"),
        ("single", "Single"),
        ("unmarried", "Single"),
        ("divorced", "Single"),
        ("widowed", "Single"),
        ("never married", "Single"),
    ],
    "nominee_gender": [
        ("male", "Male"),
        ("female", "Female"),
    ],
    "nominee_relationship": [
        ("father/mother", "Father/Mother"),
        ("son/daughter", "Son/Daughter"),
        ("wife/husband", "Wife/Husband"),
        ("father", "Father"),
        ("mother", "Mother"),
        ("son", "Son"),
        ("daughter", "Daughter"),
        ("wife", "Wife"),
        ("husband", "Husband"),
        ("uncle", "Uncle"),
    ],
    "payment_mode": [
        ("upi/ qr", "UPI/QR"),
        ("upi/qr", "UPI/QR"),
        ("upi", "UPI/QR"),
        ("qr", "UPI/QR"),
        ("cash", "Cash"),
        ("finance", "Finance"),
        ("flipkart", "Flipkart"),
    ],
}


def _match_checkbox_canonical(segment: str, field: str) -> str | None:
    """Map a substring (one checkbox option region) to a canonical label for ``field``."""
    if not segment or not field:
        return None
    seg = re.sub(r"\s+", " ", segment.lower().strip())
    opts = _FIELD_CHECKBOX_ALIASES.get(field)
    if not opts:
        return None
    # Longer keys first so "father/mother" wins over "father"
    ordered = sorted(opts, key=lambda x: len(x[0]), reverse=True)
    for key, canonical in ordered:
        if key in seg:
            return canonical
    return None


def _normalize_filled_box_marks_to_selected_token(s: str) -> str:
    """
    Map **filled checkbox squares** (scanned forms, e.g. Canon) to the same token as a tick ``[✓]``
    so :func:`_extract_tick_before_option_value` and segment splitters behave identically.

    Handles bracketed shapes ``[■]``, ``[█]``, ballot ``[☒]`` / ``[☑]``, and unbracketed ``■ Farmer``-style
    marks before the option label. Empty boxes ``[ ]`` / ``☐`` are left unchanged.
    """
    if not s:
        return s
    u = s
    for sym in (
        "[■]",
        "[█]",
        "[▪]",
        "[◼]",
        "[◾]",
        "[⬛]",
        "[▮]",
        "[●]",
        "[☒]",
        "[☑]",
    ):
        u = u.replace(sym, "[✓]")
    # Unbracketed filled glyphs before an option word (same reading order as ``[✓] Private``)
    u = re.sub(
        r"(?<![\[\w/])([\u25A0\u2588\u25AA\u25FC\u25FE\u2B1B\u25CF\u25AE\u2612])\s+(?=[A-Za-z(])",
        "[✓] ",
        u,
    )
    return u


def _extract_tick_before_option_value(s: str, field: str) -> str | None:
    """
    Dealer forms place the **checkmark before the option** (reading order): ``[✓] Farmer`` / ``✓ Farmer``
    or filled-box ``[■] Farmer`` / ``■ Farmer`` — means **Farmer** is selected — not ``Farmer [✓]`` after another option.

    Match **tick / filled box / bracket mark immediately followed by** the option phrase (longest alias first).
    """
    if not s or not field:
        return None
    u = s.replace("[x]", "[✓]").replace("[X]", "[✓]")
    tick_class = r"[\u2713\u2714\u2611\u2612\u2705☑☒✓✔■▪█◼◾⬛●▮]"
    opts = sorted(
        (_FIELD_CHECKBOX_ALIASES.get(field) or []),
        key=lambda x: len(x[0]),
        reverse=True,
    )
    for key, canonical in opts:
        esc = re.escape(key)
        if re.search(rf"(?i)\[✓\]\s*{esc}", u):
            return canonical
        if re.search(rf"(?i){tick_class}\s*{esc}\b", u):
            return canonical
    return None


def _segments_after_selected_marks(s: str) -> list[str]:
    """Split on ``[✓]`` selected markers; legacy ``[x]`` / ``[X]`` strings are normalized first."""
    if not s:
        return []
    u = s.replace("[x]", "[✓]").replace("[X]", "[✓]")
    out: list[str] = []
    pos = 0
    token = "[✓]"
    while True:
        idx = u.find(token, pos)
        if idx < 0:
            break
        start = idx + len(token)
        rest = u[start:]
        nxt = re.search(r"\[", rest)
        end = nxt.start() if nxt else len(rest)
        chunk = rest[:end].strip()
        if chunk:
            out.append(chunk)
        # Advance past this option + following bracket so a second ``[✓]`` on the same line is found.
        pos = start + end
    return out


def _extract_checkbox_selection_value(raw: str | None, field: str) -> str | None:
    """
    If ``raw`` looks like a Textract checkbox row (``[✓]`` / ``[ ]``, legacy ``[x]``, or legacy ``X``),
    **filled boxes** (``[■]``, ``■``, ``█``, …), unicode tick marks next to labels, or handwritten ticks
    OCR'd as ``✓``/``✔``, return the canonical option for ``field``. Otherwise return ``None`` so callers keep ``raw``.

    **Layout:** Forms use **mark then label** (``[✓] Farmer`` or ``[■] Farmer``). That case is handled first via
    :func:`_extract_tick_before_option_value`; tick-after-label remains as fallback.
    """
    if raw is None or not str(raw).strip():
        return None
    s = str(raw).strip()
    s = s.replace("\u2013", "-").replace("\u2014", "-").replace("\xa0", " ")
    s = _normalize_filled_box_marks_to_selected_token(s)
    tick_first = _extract_tick_before_option_value(s, field)
    if tick_first is not None:
        return tick_first
    # Unicode / printed ticks or filled squares beside option text (FORMS may omit SELECTION_ELEMENT on pencil scans).
    # Prefer the option **immediately next to** a tick or filled box (earliest match in reading order), e.g. ``Farmer ✓ Private``.
    tick_class = r"[\u2713\u2714\u2611\u2612\u2705☑☒✓✔■▪█◼◾⬛●▮]"
    opts_u = sorted(
        (_FIELD_CHECKBOX_ALIASES.get(field) or []),
        key=lambda x: len(x[0]),
        reverse=True,
    )
    best_canon: str | None = None
    best_start: int | None = None
    for key, canonical in opts_u:
        esc = re.escape(key)
        for pat in (rf"(?i){esc}\s*{tick_class}", rf"(?i){tick_class}\s*{esc}"):
            m = re.search(pat, s)
            if not m:
                continue
            st = m.start()
            if best_start is None or st < best_start:
                best_start = st
                best_canon = canonical
    if best_canon is not None:
        return best_canon
    u = s.replace("[x]", "[✓]").replace("[X]", "[✓]")
    if "[✓]" in u or "[ ]" in u:
        segs = _segments_after_selected_marks(s)
        if not segs:
            return None
        primary = segs[0]
        got = _match_checkbox_canonical(primary, field)
        if got:
            return got
        # Short free text after a mark (e.g. OCR typo) — still better than the whole row
        if len(primary.split()) <= 5 and len(primary) <= 80:
            return primary.strip()
        return None
    # Legacy Textract: ``X`` next to an option label (unordered append)
    if re.search(r"(?i)(?<![A-Za-z0-9])X(?![A-Za-z0-9])", s):
        m = re.search(r"(?i)(?<![A-Za-z0-9])X(?![A-Za-z0-9])\s+([A-Za-z][A-Za-z0-9 /]{0,48})", s)
        if m:
            got = _match_checkbox_canonical(m.group(1).strip(), field)
            if got:
                return got
        m = re.search(
            r"(?i)([A-Za-z][A-Za-z0-9 /]{0,48})\s+(?<![A-Za-z0-9])X(?![A-Za-z0-9])",
            s,
        )
        if m:
            got = _match_checkbox_canonical(m.group(1).strip(), field)
            if got:
                return got
    return None


def _normalize_kv_value_for_checkbox_fields(field: str, value: str | None) -> str:
    """Apply checkbox semantics when ``value`` encodes selection marks; else return ``value`` unchanged."""
    if value is None:
        return ""
    v = str(value).strip()
    sel = _extract_checkbox_selection_value(v, field)
    if sel is not None:
        return sel
    return v


def _normalize_nominee_gender_sheet_value(val: str | None) -> str | None:
    """Normalize nominee gender to **Male** / **Female** (substring + fuzzy ≥ 0.5 on those labels)."""
    if val is None or not str(val).strip():
        return None
    raw = re.sub(r"\s+", " ", str(val).strip())
    s = raw.lower()
    if "female" in s:
        return "Female"
    if "male" in s:
        return "Male"
    return fuzzy_best_option_label(raw, ["Male", "Female"], min_score=0.5)


def _normalize_payment_mode_sheet_value(val: str | None) -> str | None:
    """Keep payment mode labels stable (avoid title-casing **UPI**)."""
    if val is None or not str(val).strip():
        return None
    s = str(val).strip()
    got = _extract_checkbox_selection_value(s, "payment_mode")
    if got:
        return got
    got2 = _match_checkbox_canonical(s.lower(), "payment_mode")
    return got2 if got2 else s[:120]


def _refine_nominee_relationship_with_gender(rel: str | None, gender: str | None) -> str | None:
    """
    Checkbox templates use combined labels (**Father/Mother**, **Son/Daughter**, **Wife/Husband**).
    With nominee gender (Male/Female), store the specific relation: e.g. Son/Daughter + Female -> Daughter.
    """
    if rel is None or not str(rel).strip():
        return None
    g = _normalize_nominee_gender_sheet_value(gender) if gender else None
    if g not in ("Male", "Female"):
        return str(rel).strip()
    raw = str(rel).strip()
    r = normalize_nominee_relationship_value(raw) or raw
    r = r.rstrip(".")
    m = re.match(r"(?i)^\s*([A-Za-z]+)\s*/\s*([A-Za-z]+)\s*$", r)
    if not m:
        return r
    a, b = m.group(1).lower(), m.group(2).lower()
    pair = frozenset((a, b))
    if pair == frozenset(("father", "mother")):
        return "Mother" if g == "Female" else "Father"
    if pair == frozenset(("son", "daughter")):
        return "Daughter" if g == "Female" else "Son"
    if pair == frozenset(("wife", "husband")):
        return "Wife" if g == "Female" else "Husband"
    return r


def _sync_nominee_relation_with_gender_across_fragments(
    insurance: dict[str, str],
    details_customer: dict[str, str],
) -> None:
    """
    Relation may land in ``insurance`` and gender in ``details_customer`` (or vice versa).
    Apply combined-label -> specific relation using the best available gender + relation pair.
    """
    rel = (insurance.get("nominee_relationship") or details_customer.get("nominee_relationship") or "").strip()
    gender = (insurance.get("nominee_gender") or details_customer.get("nominee_gender") or "").strip()
    if not rel or not gender:
        return
    refined = _refine_nominee_relationship_with_gender(rel, gender)
    if not refined or refined.strip() == rel.strip():
        return
    insurance["nominee_relationship"] = refined
    if details_customer.get("nominee_relationship"):
        details_customer["nominee_relationship"] = refined


def _details_input_format(path: Path) -> str:
    """
    Detect real file type for the sales detail upload. V2 always saves as Details.jpg
    even when the user uploads .docx or PDF, so we must sniff magic bytes.
    """
    try:
        head = path.read_bytes()[:12]
    except OSError:
        return "unknown"
    if len(head) >= 2 and head[:2] == b"\xff\xd8":
        return "jpeg"
    if len(head) >= 4 and head[:4] == b"\x89PNG":
        return "png"
    if len(head) >= 4 and head[:4] == b"%PDF":
        return "pdf"
    if len(head) >= 2 and head[:2] == b"PK":
        return "docx"
    return "unknown"


def _parallel_textract_prefetch_upload_subfolder(subdir: Path) -> tuple[dict[str, dict], dict[str, int]]:
    """
    Run AWS Textract calls for all upload files concurrently to cut wall time on scans-v2.
    Keys: ``aadhar_front`` / ``aadhar_back`` (DetectDocumentText), ``details_forms``
    (**AnalyzeDocument** FORMS+TABLES for structured sales detail scans/PDFs), ``insurance``,
    ``financing``. ``.docx`` Details is not prefetched (parsed locally).

    Returns ``(results_by_key, job_duration_ms_by_key)`` for logging and ``section_timings_ms``.
    """
    from concurrent.futures import ThreadPoolExecutor

    from app.config import OCR_UPLOAD_TEXTRACT_TIMEOUT_SEC
    from app.services.sales_textract_service import extract_forms_from_bytes, extract_text_from_bytes

    jobs: list[tuple[str, bytes, str]] = []
    ap = _prefer_for_ocr_input(subdir, "Aadhar.pdf", FILENAME_AADHAR_FRONT, LEGACY_AADHAR_FRONT_JPG)
    if ap.is_file():
        try:
            jobs.append(("aadhar_front", ap.read_bytes(), "text"))
        except OSError as e:
            logger.warning("prefetch: could not read %s: %s", ap, e)
    dp = _prefer_details_sheet_input(subdir)
    if dp.is_file() and _details_input_format(dp) in ("jpeg", "png", "pdf"):
        try:
            jobs.append(("details_forms", dp.read_bytes(), "forms"))
        except OSError as e:
            logger.warning("prefetch: could not read %s: %s", dp, e)
    ip = _prefer_for_ocr_input(subdir, "Insurance.pdf", "Insurance.jpg")
    if ip.is_file():
        try:
            jobs.append(("insurance", ip.read_bytes(), "text"))
        except OSError as e:
            logger.warning("prefetch: could not read %s: %s", ip, e)
    for pdf_name, jpg_name, key in (
        ("Aadhar_back.pdf", "Aadhar_back.jpg", "aadhar_back"),
        ("Financing.pdf", "Financing.jpg", "financing"),
    ):
        p = _prefer_for_ocr_input(subdir, pdf_name, jpg_name)
        if p.is_file():
            try:
                jobs.append((key, p.read_bytes(), "text"))
            except OSError as e:
                logger.warning("prefetch: could not read %s: %s", p, e)

    out: dict[str, dict] = {}
    job_ms: dict[str, int] = {}
    if not jobs:
        return out, job_ms

    timeout = max(30, OCR_UPLOAD_TEXTRACT_TIMEOUT_SEC)
    max_workers = min(5, len(jobs))

    def _run(job: tuple[str, bytes, str]) -> tuple[str, dict, int]:
        key, blob, mode = job
        t0 = time.perf_counter()
        try:
            if mode == "forms":
                res = extract_forms_from_bytes(blob)
            else:
                res = extract_text_from_bytes(blob)
        except Exception as e:
            res = {
                "error": str(e),
                "full_text": "",
                "key_value_pairs": [],
                "blocks": [],
                "raw_response": None,
            }
        elapsed_ms = int((time.perf_counter() - t0) * 1000)
        return key, res, elapsed_ms

    with ThreadPoolExecutor(max_workers=max_workers) as ex:
        keyed_futures = []
        for job in jobs:
            keyed_futures.append((job[0], ex.submit(_run, job)))
        for key, fut in keyed_futures:
            try:
                k, res, elapsed_ms = fut.result(timeout=timeout)
                out[k] = res
                job_ms[k] = elapsed_ms
            except Exception as e:
                logger.warning("prefetch Textract failed for %s: %s", key, e)
                out[key] = {
                    "error": str(e),
                    "full_text": "",
                    "key_value_pairs": [],
                    "blocks": [],
                    "raw_response": None,
                }
                job_ms[key] = 0
    return out, job_ms


def _key_value_pairs_from_docx(doc_path: Path) -> tuple[list[dict[str, str]], str]:
    """
    Parse Word Sales Detail Sheet (.docx) into Textract-like key/value pairs.
    AWS Textract does not accept .docx bytes; this path is used for native Word uploads.
    """
    try:
        from docx import Document
    except ImportError:
        return [], ""

    doc = Document(str(doc_path))
    pairs: list[dict[str, str]] = []
    text_lines: list[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        text_lines.append(t)
        if ":" in t:
            left, right = t.split(":", 1)
            lk, rv = left.strip(), right.strip()
            if lk and rv and len(lk) <= 160 and len(rv) <= 600:
                pairs.append({"key": lk, "value": rv})
        elif "\t" in t:
            parts = re.split(r"\t+", t, maxsplit=1)
            if len(parts) == 2 and parts[0].strip() and parts[1].strip():
                pairs.append({"key": parts[0].strip(), "value": parts[1].strip()})

    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if not any(cells):
                continue
            n = len(cells)
            if n >= 2:
                for i in range(0, n - 1, 2):
                    key, val = cells[i], cells[i + 1]
                    if key and val and key.casefold() != val.casefold():
                        pairs.append({"key": key, "value": val})

    seen: set[tuple[str, str]] = set()
    deduped: list[dict[str, str]] = []
    for kv in pairs:
        k = (kv.get("key") or "").strip()
        v = (kv.get("value") or "").strip()
        if not k or not v:
            continue
        sig = (k.casefold(), v.casefold())
        if sig in seen:
            continue
        seen.add(sig)
        deduped.append({"key": k, "value": v})

    full_text = "\n".join(text_lines)
    for table in doc.tables:
        for row in table.rows:
            row_txt = " | ".join((c.text or "").strip() for c in row.cells)
            if row_txt.strip():
                full_text += "\n" + row_txt
    return deduped, full_text


def _sanitize_nominee_age(val: str | None) -> str | None:
    """Extract numeric part from nominee age (e.g. '30/m' -> '30'). Returns None if no digits."""
    if not val or not str(val).strip():
        return None
    s = str(val).strip()
    m = re.match(r"^(\d{1,3})", s)
    return m.group(1) if m and 1 <= int(m.group(1)) <= 150 else None


def _normalize_name_for_match(name: str | None) -> str:
    """Normalize name for matching: lowercase, strip, collapse spaces."""
    if not name or not str(name).strip():
        return ""
    return " ".join(str(name).lower().strip().split())


# Details vs Aadhaar OCR name reconciliation (upload merge + Raw_OCR fallbacks).
NAME_RECONCILE_MIN_RATIO = 0.5


def _details_name_core_for_match(s: str) -> str:
    """Strip trailing ``S/o`` / ``W/o`` / ``D/o`` parentage tails from Details ``Full Name``."""
    t = (s or "").strip()
    if not t:
        return ""
    t = re.sub(r"(?i)\s+[SWD]/[Oo]\s+.+$", "", t).strip()
    t = re.sub(r"(?i)\s+c/o\s+.+$", "", t).strip()
    return t


def _name_fuzzy_ratio(a: str | None, b: str | None) -> float:
    n1 = _normalize_name_for_match(a)
    n2 = _normalize_name_for_match(b)
    if not n1 or not n2:
        return 0.0
    return float(SequenceMatcher(None, n1, n2).ratio())


def _max_fuzzy_vs_details(candidate: str, details_name: str) -> float:
    core = _details_name_core_for_match(details_name)
    return max(
        _name_fuzzy_ratio(candidate, details_name),
        _name_fuzzy_ratio(candidate, core) if core else 0.0,
    )


def _concat_aadhar_scan_ocr_text(raw_parts: list[tuple[str, str]] | None) -> str:
    """Join Textract blobs for Aadhaar front/back only (excludes Details / Insurance sections)."""
    if not raw_parts:
        return ""
    chunks: list[str] = []
    for fn, text in raw_parts:
        fl = (fn or "").strip().replace("\\", "/").split("/")[-1].lower()
        if "aadhar" in fl:
            chunks.append(text or "")
    return "\n".join(chunks)


def _aadhar_blob_line_looks_like_name_line(line: str) -> bool:
    line = line.strip()
    if len(line) < 3 or len(line) > 90:
        return False
    if re.search(r"\d{4}\s+\d{4}\s+\d{4}", line):
        return False
    if re.match(r"^[\d\s/:\-.]+$", line):
        return False
    noise = re.compile(
        r"(?i)aadhaar|aadhar|government|india|identification|male|female|transgender|"
        r"enrol|address|dob|date\s*of\s*birth|vid|virtual|uidai|unique|pin\s*code"
    )
    if noise.search(line):
        return False
    if not re.match(r"^[A-Za-z][A-Za-z\s.'-]{2,88}$", line):
        return False
    words = [w for w in line.split() if w.strip()]
    return 2 <= len(words) <= 8


def _collect_aadhar_name_candidates(blob: str) -> list[str]:
    """Phrases from Aadhaar OCR that might be the holder name (for fuzzy match to Details core name)."""
    if not blob or not str(blob).strip():
        return []
    seen: set[str] = set()
    out: list[str] = []

    def add(s: str) -> None:
        s = " ".join(s.split())
        if len(s) < 3 or len(s) > 90:
            return
        key = s.lower()
        if key in seen:
            return
        seen.add(key)
        out.append(s)

    for ln in blob.replace("\r\n", "\n").split("\n"):
        t = ln.strip()
        if _aadhar_blob_line_looks_like_name_line(t):
            add(t)

    words = re.findall(r"[A-Za-z][A-Za-z.'-]*", blob)[:400]
    max_w = len(words)
    for wlen in range(2, 7):
        for i in range(0, max(0, max_w - wlen + 1)):
            phrase = " ".join(words[i : i + wlen])
            if _aadhar_blob_line_looks_like_name_line(phrase):
                add(phrase)
        if len(out) > 500:
            break
    return out


def _best_aadhar_scan_name_for_details(
    details_name: str,
    aadhar_blob: str,
    *,
    min_ratio: float,
) -> str | None:
    """
    Match Aadhaar OCR phrases against the Details **core** name only (``S/o`` / ``D/o`` / ``W/o``
    tails removed). Does not score against the full Details line including parentage.
    """
    dn = (details_name or "").strip()
    if not dn or not (aadhar_blob or "").strip():
        return None
    query = (_details_name_core_for_match(dn) or dn).strip()
    if not query:
        return None
    ranked: list[tuple[float, int, str]] = []
    for cand in _collect_aadhar_name_candidates(aadhar_blob):
        sc = _name_fuzzy_ratio(cand, query)
        if sc > min_ratio:
            ranked.append((sc, len(cand), cand))
    if not ranked:
        return None
    ranked.sort(key=lambda x: (-x[0], -x[1]))
    return ranked[0][2]


def _reconcile_customer_name_aadhar_details(
    aadhar_name: str | None,
    details_name: str | None,
    aadhar_blob: str,
    *,
    min_ratio: float = NAME_RECONCILE_MIN_RATIO,
) -> str | None:
    """
    1) If parsed Aadhaar ``name`` fuzzy-matches Details ``Full Name`` (above ``min_ratio``), keep it.
    2) Else search Aadhaar OCR for a phrase that fuzzy-matches the Details **core** name (no ``S/o`` /
       ``D/o`` / ``W/o``) above ``min_ratio``; use that substring from the Aadhaar scan.
    3) Else keep the parsed Aadhaar name, or Details name if Aadhaar name was empty.
    """
    dn = (details_name or "").strip()
    an = (aadhar_name or "").strip()
    blob = aadhar_blob or ""
    if not dn:
        return an or None
    if not blob.strip():
        return an or dn or None
    if an and _max_fuzzy_vs_details(an, dn) > min_ratio:
        return an
    best = _best_aadhar_scan_name_for_details(dn, blob, min_ratio=min_ratio)
    if best:
        return best
    if an:
        return an
    return dn or None


def _initcap_words(value: str | None) -> str:
    """Normalize spacing and convert words to InitCap."""
    s = " ".join(str(value or "").strip().split())
    if not s:
        return ""
    return s.title()


def _apply_initcap_on_read(data: dict[str, Any]) -> None:
    """
    Presentation normalization for API reads:
    convert human-readable OCR text fields to InitCap.
    """
    customer = data.get("customer") or {}
    if isinstance(customer, dict):
        if "profession" in customer:
            sp = _sanitize_details_profession_value(customer.get("profession"))
            customer["profession"] = _initcap_words(sp) if sp else default_profession_if_empty("")
        if customer.get("marital_status"):
            ms = _normalize_details_marital_status_value(customer.get("marital_status"))
            if ms:
                customer["marital_status"] = _initcap_words(ms)
            else:
                customer.pop("marital_status", None)
        for k in ("name", "care_of", "city", "district", "sub_district", "post_office", "state"):
            if customer.get(k):
                customer[k] = _initcap_words(customer.get(k))
        data["customer"] = customer

    insurance = data.get("insurance") or {}
    if isinstance(insurance, dict):
        if "profession" in insurance:
            sp = _sanitize_details_profession_value(insurance.get("profession"))
            insurance["profession"] = _initcap_words(sp) if sp else default_profession_if_empty("")
        if insurance.get("marital_status"):
            ms = _normalize_details_marital_status_value(insurance.get("marital_status"))
            if ms:
                insurance["marital_status"] = _initcap_words(ms)
            else:
                insurance.pop("marital_status", None)
        if "insurer" in insurance:
            ins = sanitize_details_sheet_insurer_value(insurance.get("insurer"))
            if ins:
                insurance["insurer"] = _initcap_words(ins)
            else:
                insurance.pop("insurer", None)
        if insurance.get("nominee_relationship"):
            nr = normalize_nominee_relationship_value(insurance.get("nominee_relationship"))
            if nr:
                insurance["nominee_relationship"] = _initcap_words(nr)
            else:
                insurance.pop("nominee_relationship", None)
        if insurance.get("payment_mode"):
            pm = _normalize_payment_mode_sheet_value(insurance.get("payment_mode"))
            if pm:
                insurance["payment_mode"] = pm
            else:
                insurance.pop("payment_mode", None)
        for k in ("nominee_name", "financier", "policy_holder_name"):
            if insurance.get(k):
                insurance[k] = _initcap_words(insurance.get(k))
        data["insurance"] = insurance

    if data.get("details_customer_name"):
        data["details_customer_name"] = _initcap_words(data.get("details_customer_name"))


def _names_match(name1: str | None, name2: str | None) -> bool:
    """Return True if names likely refer to same person. Handles OCR variations (spacing, case)."""
    n1 = _normalize_name_for_match(name1)
    n2 = _normalize_name_for_match(name2)
    if not n1 or not n2:
        return True  # Can't compare if one missing; allow
    if n1 == n2:
        return True
    # One name contains the other (e.g. "john doe" in "john doe kumar")
    if n1 in n2 or n2 in n1:
        return True
    # OCR-tolerant token overlap: allow partial / slightly noisy matches.
    t1 = [t for t in re.sub(r"[^a-z\s]", " ", n1).split() if t]
    t2 = [t for t in re.sub(r"[^a-z\s]", " ", n2).split() if t]
    if not t1 or not t2:
        return True
    s1, s2 = set(t1), set(t2)
    inter = len(s1 & s2)
    if inter >= 1:
        cov = inter / max(1, min(len(s1), len(s2)))
        if cov >= 0.5:
            return True
        # First token often survives OCR better than suffixes.
        if t1[0] == t2[0]:
            return True
    # Minor OCR drift tolerance (e.g. Vishnu vs Wishnu).
    if SequenceMatcher(None, n1, n2).ratio() >= 0.82:
        return True
    if t1 and t2 and SequenceMatcher(None, t1[0], t2[0]).ratio() >= 0.8:
        return True
    return False


def _validate_name_match(
    aadhar_name: str | None,
    details_name: str | None,
    insurance_name: str | None,
    *,
    aadhar_last4: str | None = None,
    details_aadhar_last4: str | None = None,
    insurance_aadhar_last4: str | None = None,
) -> str | None:
    """Return error message if names from Aadhar, Details and Insurance do not match. None if OK."""
    names = [
        ("Aadhar", aadhar_name),
        ("Details sheet", details_name),
        ("Insurance", insurance_name),
    ]
    present = [(label, n) for label, n in names if n and str(n).strip()]
    if len(present) < 2:
        return None  # Need at least 2 to compare
    first_label, first_val = present[0]
    for label, val in present[1:]:
        if not _names_match(first_val, val):
            a4 = _aadhar_last4(aadhar_last4 or "")
            d4 = _aadhar_last4(details_aadhar_last4 or "")
            i4 = _aadhar_last4(insurance_aadhar_last4 or "")
            if a4 and ((d4 and a4 == d4) or (i4 and a4 == i4)):
                return None
            return (
                f"Name mismatch: '{first_label}' has a different name than '{label}'. "
                "Ensure the name on Aadhar front, Details sheet and Insurance document match."
            )
    return None


def _map_key_value_pairs_to_vehicle(pairs: list[dict]) -> dict[str, str]:
    """Map key_value_pairs from Textract to structured vehicle fields (frame_no, engine_no, model_colour, key_no, battery_no)."""
    out: dict[str, str] = {}
    key_norm_to_value: dict[str, str] = {}
    for kv in pairs:
        k = (kv.get("key") or "").strip()
        v = (kv.get("value") or "").strip()
        if not k:
            continue
        key_norm = _normalize_key_for_match(k)
        key_norm_to_value[key_norm] = v
        if ":" in key_norm:
            key_norm_to_value[key_norm.replace(":", "").strip()] = v

    for field, aliases in _VEHICLE_KEY_ALIASES.items():
        if field in out:
            continue
        for alias in aliases:
            anorm = _normalize_key_for_match(alias)
            if anorm in key_norm_to_value:
                out[field] = key_norm_to_value[anorm]
                break
            for kn, v in key_norm_to_value.items():
                if anorm in kn or kn in anorm:
                    out[field] = v
                    break
            if field in out:
                break

    # Combine Model and Colour into model_colour if we have them separately
    model_val = key_norm_to_value.get(_normalize_key_for_match("model"), "").strip() or next(
        (
            v.strip()
            for kn, v in key_norm_to_value.items()
            if "model" in kn and "colour" not in kn and "color" not in kn and v.strip()
        ),
        "",
    )
    colour_val = key_norm_to_value.get(_normalize_key_for_match("colour"), "").strip() or key_norm_to_value.get(
        _normalize_key_for_match("color"), ""
    ).strip()
    if model_val or colour_val:
        combined = ", ".join(filter(None, [model_val, colour_val]))
        if combined:
            out["model_colour"] = combined

    # Guard against section-heading bleed when Key/Battery fields are blank or scratched.
    for fld in ("key_no", "battery_no"):
        val = (out.get(fld) or "").strip()
        if val and re.search(r"(?i)\b(nominee|payment|insurance|details|customer|vehicle)\b", val):
            out.pop(fld, None)

    return out


def _clean_sales_sheet_scalar(value: str) -> str:
    """Strip placeholder underscores / blanks from a field value."""
    s = (value or "").strip()
    if not s:
        return ""
    if re.match(r"^[\s_.,-]+$", s):
        return ""
    return s.strip()


def _canonical_marital_status_from_text(s: str) -> str | None:
    """
    Map OCR / handwritten marital text to **Married** or **Single** (Synonyms for Single: unmarried,
    divorced, widowed, never married). Uses substring rules first, then fuzzy match ≥ 0.5 vs Married/Single.
    """
    t = (s or "").strip()
    t = t.replace("\u2013", "-").replace("\u2014", "-").replace("\u2012", "-")
    t = re.sub(
        r"(?i)^[\s\-–—:._]*(?:marital|martial|marrital)\s*status\s*[:\s]*",
        "",
        t,
    ).strip()
    sl = re.sub(r"\s+", " ", t.lower())
    if not sl:
        return None
    if re.search(r"\bnever\s+married\b", sl):
        return "Single"
    if re.search(r"\b(unmarried|unmaried|un-maried)\b", sl):
        return "Single"
    if re.search(r"\bdivorced\b", sl):
        return "Single"
    if re.search(r"\bwidowed\b", sl) or re.search(r"\bwidow\b", sl):
        return "Single"
    if re.search(r"\bsingle\b", sl):
        return "Single"
    if re.search(r"\bmarried\b", sl):
        return "Married"
    return fuzzy_best_option_label(t, ["Married", "Single"], min_score=0.5)


def _sanitize_details_profession_value(val: str | None) -> str | None:
    """
    Sales detail sheets often place **Profession** and **Marital Status** on one row. When
    Profession is left blank, Textract may merge the rest of the row (e.g. ``Marital Status:
    Unmarried`` / ``- Marital Status: Unmaried``) into the Profession value. Strip that bleed
    so profession is empty unless a real token appears before the marital-status field.
    """
    if val is None or not str(val).strip():
        return None
    s = str(val).strip()
    # Normalize unicode dashes and spaces OCR often mixes with "- Marital Status: …"
    s = s.replace("\u2013", "-").replace("\u2014", "-").replace("\u2012", "-").replace("\xa0", " ")
    # Full-width / compatibility colons (printed forms, OCR)
    s = s.replace("\uff1a", ":").replace("\ufe55", ":")

    # Entire value is only a marital-status label + outcome (incl. OCR "Martial" for "Marital",
    # "Marrital", glued "MaritalStatus:", full-width colon).
    if re.match(
        r"(?i)^[\s\-–—:._]*(?:marital|martial|marrital)\s*status\s*[:：]?\s*(?:unmaried|un-maried|unmarried|single|married|divorced|widowed)?\s*$",
        s,
    ):
        return None

    # Cut before the earliest OCR variant of "marital status" (including glued "maritalstatus").
    _marital_field_pat = (
        r"(?i)\bmarital\s*status\b",
        r"(?i)\bmartial\s*status\b",  # common OCR misread of "Marital"
        r"(?i)\bmarrital\s*status\b",
        r"(?i)marital\s*statu?s\b",
        r"(?i)mari\s*ta?l\s*status\b",
        r"(?i)maritalstatus\b",
        r"(?i)\bmarital\s*stat\b",  # truncated "stat"
    )
    cut_at: int | None = None
    for pat in _marital_field_pat:
        m = re.search(pat, s)
        if m:
            if cut_at is None or m.start() < cut_at:
                cut_at = m.start()
    if cut_at is not None:
        s = s[:cut_at]

    s = s.strip(" \t:._-–—")
    s = _clean_sales_sheet_scalar(s)
    if not s:
        return None

    # Whole value is only marital outcome (incl. OCR "Unmaried")
    if re.match(
        r"(?i)^(unmarried|unmaried|married|single|divorced|widowed)\s*$",
        s,
    ):
        return None

    # Remainder still starts like a marital-status line (no real profession token)
    if re.match(r"(?i)^(marital|martial|marrital|maritalstatus|mari\s*ta?l)\b", s):
        return None

    got = fuzzy_best_option_label(s, list(_WRITTEN_PROFESSION_OPTIONS), min_score=0.5)
    if got:
        return got
    q = normalize_for_fuzzy_match(s)
    ws = [w for w in q.split() if w]
    wset = set(ws)
    if "private" in wset and "job" in wset:
        return "Private Job"
    if ("pvt" in wset or "pvt." in wset) and "job" in wset:
        return "Private Job"
    if q.replace(" ", "") in ("privatejob", "pvtjob"):
        return "Private Job"
    return None


def _normalize_details_marital_status_value(val: str | None) -> str | None:
    """
    Normalize to **Married** or **Single** only. **Single** includes unmarried, divorced, widowed,
    never married (plus OCR **Unmaried**). Fuzzy match ≥ 0.5 vs Married/Single when keywords do not apply.
    """
    if val is None or not str(val).strip():
        return None
    s = str(val).strip()
    s = s.replace("\u2013", "-").replace("\u2014", "-").lstrip("-–—").strip()
    sl = re.sub(r"\s+", " ", s.lower())
    if sl in ("unmaried", "un-maried"):
        return "Single"
    return _canonical_marital_status_from_text(s)


def _parse_vehicle_from_full_text(full_text: str) -> dict[str, str]:
    """
    Fallback when Textract FORMS misses pairs (common on some PDFs): parse LINE layout
    like 'Chassis Number: 59324 Engine Number: 50581'.
    """
    out: dict[str, str] = {}
    if not full_text or not isinstance(full_text, str):
        return out
    text = full_text.replace("\r", "\n")
    lines = [ln.strip() for ln in text.split("\n")]

    def _section_or_noise_line(s: str) -> bool:
        return bool(re.search(r"(?i)\b(nominee|payment|insurance|customer details|vehicle details)\b", s))

    def _numeric_tokens(s: str, *, min_len: int = 3, max_len: int = 12) -> list[str]:
        toks = re.findall(r"\d+", s or "")
        return [t for t in toks if min_len <= len(t) <= max_len]

    def _best_numeric_token(
        s: str,
        *,
        min_len: int = 3,
        max_len: int = 12,
        prefer: str = "first",
    ) -> str:
        toks = _numeric_tokens(s, min_len=min_len, max_len=max_len)
        if not toks:
            return ""
        if prefer == "last":
            return toks[-1].strip()
        if prefer == "single":
            singles = [t for t in toks if len(toks) == 1]
            if singles:
                return singles[0].strip()
        return toks[0].strip()

    def _extract_field_near_label(
        label_rx: str,
        *,
        min_len: int = 3,
        max_len: int = 12,
        look_ahead: int = 3,
        prefer_below_over_inline: bool = False,
        below_token_prefer: str = "first",
        inline_token_prefer: str = "first",
        allow_inline: bool = True,
        allow_backward: bool = True,
    ) -> str:
        for idx, ln in enumerate(lines):
            if not re.search(label_rx, ln, re.I):
                continue
            m_inline = re.search(r":\s*(.+)$", ln)
            cand_inline = ""
            if m_inline and allow_inline:
                cand_inline = _best_numeric_token(
                    m_inline.group(1),
                    min_len=min_len,
                    max_len=max_len,
                    prefer=inline_token_prefer,
                )
            # 1) Prefer nearest numeric line below label (for scratched/overwritten values).
            for j in range(idx + 1, min(len(lines), idx + 1 + look_ahead)):
                nln = lines[j]
                if not nln or _section_or_noise_line(nln):
                    continue
                # Skip another label line.
                if ":" in nln and re.search(r"[A-Za-z]", nln):
                    continue
                cand = _best_numeric_token(
                    nln, min_len=min_len, max_len=max_len, prefer=below_token_prefer
                )
                if cand:
                    return cand
            # 2) Inline value after colon in same line.
            if cand_inline and not prefer_below_over_inline:
                return cand_inline
            # 3) Small backward window (number slightly above label).
            if allow_backward:
                for j in range(max(0, idx - 2), idx):
                    pln = lines[j]
                    if not pln or _section_or_noise_line(pln):
                        continue
                    if ":" in pln and re.search(r"[A-Za-z]", pln):
                        continue
                    cand = _best_numeric_token(
                        pln, min_len=min_len, max_len=max_len, prefer=below_token_prefer
                    )
                    if cand:
                        return cand
            if cand_inline:
                return cand_inline
        return ""

    m = re.search(r"(?i)Model\s*:\s*([^\n]+?)(?=\s+Colour\s*:)", text)
    if m:
        mv = _clean_sales_sheet_scalar(m.group(1))
        if mv:
            out["model_colour"] = mv
    m = re.search(r"(?i)Colour\s*:\s*([^\n]+)", text)
    if m:
        cv = _clean_sales_sheet_scalar(m.group(1))
        if cv:
            if out.get("model_colour"):
                out["model_colour"] = f"{out['model_colour']}, {cv}"
            else:
                out["model_colour"] = cv

    m = re.search(r"(?i)Chassis\s+Number\s*:\s*([^\n]+?)(?=\s+Engine\s+Number\s*:|\n|$)", text)
    if m:
        fv = _clean_sales_sheet_scalar(m.group(1))
        if fv:
            out["frame_no"] = fv
    m = re.search(r"(?i)Engine\s+Number\s*:\s*([^\n]+?)(?=\s+Key\s+Number\s*:|\n|$)", text)
    if m:
        ev = _clean_sales_sheet_scalar(m.group(1))
        if ev:
            out["engine_no"] = ev
    m = re.search(r"(?i)Key\s+Number\s*:\s*([^\n]+?)(?=\s+Battery\s+Number\s*:|\n|$)", text)
    if m:
        kv = _clean_sales_sheet_scalar(m.group(1))
        if kv and not re.search(r"(?i)\b(nominee|payment|insurance|details|customer|vehicle)\b", kv):
            out["key_no"] = kv
    m = re.search(r"(?i)Battery\s+Number\s*:\s*([^\n]*)(?=\n|$)", text)
    if m:
        bv = _clean_sales_sheet_scalar(m.group(1))
        if bv and not re.search(r"(?i)\b(nominee|payment|insurance|details|customer|vehicle)\b", bv):
            out["battery_no"] = bv

    # Nearby-number rescue for scratched/blank fields in Details sheet.
    # Scratch-correction pass (override): prefer corrected numbers written below labels.
    v = _extract_field_near_label(
        r"\b(chassis|frame)\s+number\b",
        min_len=4,
        max_len=12,
        look_ahead=3,
        prefer_below_over_inline=True,
        below_token_prefer="last",
        inline_token_prefer="last",
        allow_inline=True,
        allow_backward=False,
    )
    if v:
        out["frame_no"] = v

    # Engine correction: prefer a single-token corrected line below Engine label (e.g. 76658).
    v = _extract_field_near_label(
        r"\bengine\s+number\b",
        min_len=4,
        max_len=12,
        look_ahead=3,
        prefer_below_over_inline=True,
        below_token_prefer="single",
        inline_token_prefer="last",
        allow_inline=True,
        allow_backward=False,
    )
    if v:
        out["engine_no"] = v
    if "key_no" not in out or not str(out.get("key_no") or "").strip():
        v = _extract_field_near_label(r"\bkey\s+number\b", min_len=3, max_len=10)
        if v:
            out["key_no"] = v
    # Battery should remain blank when blank on sheet; do not rescue from nearby numbers.

    # Alternate labels
    if "frame_no" not in out:
        m = re.search(r"(?i)(?:Frame|Chassis)\s+No\.?\s*:\s*([^\n]+?)(?=\s+Engine|\n|$)", text)
        if m:
            fv = _clean_sales_sheet_scalar(m.group(1))
            if fv:
                out["frame_no"] = fv
    return out


def _extract_details_customer_name(pairs: list[dict]) -> str | None:
    """Extract customer name from Details sheet key-value pairs."""
    key_lower_to_value: dict[str, str] = {}
    for kv in pairs:
        k = (kv.get("key") or "").strip()
        v = (kv.get("value") or "").strip()
        if not k or not v:
            continue
        key_norm = _normalize_key_for_match(k)
        key_lower_to_value[key_norm] = v
        if ":" in key_norm:
            key_lower_to_value[key_norm.replace(":", "").strip()] = v
    for alias in _DETAILS_CUSTOMER_NAME_ALIASES:
        anorm = _normalize_key_for_match(alias)
        if anorm in key_lower_to_value:
            val = key_lower_to_value[anorm].strip()
            if len(val) >= 2 and len(val) <= 80 and not re.search(r"(?i)\bnominee\b", val):
                return val
        for k, v in key_lower_to_value.items():
            if re.search(r"(?i)\bnominee\b", k):
                continue
            if anorm in k or k in anorm:
                val = v.strip()
                if len(val) >= 2 and len(val) <= 80 and not re.search(r"(?i)\bnominee\b", val):
                    return val
    return None


def _extract_financier_from_payment_line(text: str) -> str | None:
    """Arya-style sales detail: 'Payment: (A) Cash ... (D) Finance: SHRIRAM FIN' (no separate Financier label)."""
    if not text or not isinstance(text, str):
        return None
    low = text.lower()
    if "finance" not in low:
        return None
    m = re.search(r"(?i)(?:\(D\)\s*|D\)\s*)Finance\s*:\s*([^\n]+)", text)
    if not m and "payment" in low and re.search(r"(?i)\(D\)", text):
        m = re.search(r"(?i)Finance\s*:\s*([^\n]+)", text)
    if not m:
        return None
    val = m.group(1).strip()
    for cut in (" I agree", " Insurer", " Insurer Name", " Customer Signature"):
        if cut in val:
            val = val.split(cut)[0].strip()
    val = val.rstrip(".,; ")
    if len(val) < 2 or len(val) > 160:
        return None
    if re.match(r"^(yes|no|na|n/?a|cash|required|optional)\s*$", val, re.I):
        return None
    return val


def _map_key_value_pairs_to_insurance(pairs: list[dict]) -> dict[str, str]:
    """Map key_value_pairs from Textract to insurance/nominee fields."""
    out: dict[str, str] = {}
    key_lower_to_value: dict[str, str] = {}
    for kv in pairs:
        k = (kv.get("key") or "").strip()
        v = (kv.get("value") or "").strip()
        if not k:
            continue
        key_norm = _normalize_key_for_match(k)
        key_lower_to_value[key_norm] = v
        # Also store with colon stripped for "Nominee Name:" style
        if ":" in key_norm:
            key_lower_to_value[key_norm.replace(":", "").strip()] = v

    for field, aliases in _INSURANCE_KEY_ALIASES.items():
        if field in out:
            continue
        for alias in aliases:
            anorm = _normalize_key_for_match(alias)
            if anorm in key_lower_to_value:
                out[field] = key_lower_to_value[anorm]
                break
            # Match if key contains alias (e.g. "nominee name" in "Nominee Name")
            for k, v in key_lower_to_value.items():
                if anorm in k or k in anorm:
                    out[field] = v
                    break
            if field in out:
                break

    for _f in ("profession", "marital_status", "nominee_gender", "nominee_relationship", "payment_mode"):
        if _f in out and out[_f]:
            out[_f] = _normalize_kv_value_for_checkbox_fields(_f, out[_f])

    if "financier" not in out:
        for vv in key_lower_to_value.values():
            got = _extract_financier_from_payment_line(vv)
            if got:
                out["financier"] = got
                break
    if "profession" in out:
        sp = _sanitize_details_profession_value(out.get("profession"))
        out["profession"] = sp if sp else default_profession_if_empty("")
    if out.get("marital_status"):
        ms = _normalize_details_marital_status_value(out["marital_status"])
        if ms:
            out["marital_status"] = ms
        else:
            out.pop("marital_status", None)
    if out.get("nominee_gender"):
        ng = _normalize_nominee_gender_sheet_value(out.get("nominee_gender"))
        if ng:
            out["nominee_gender"] = ng
        else:
            out.pop("nominee_gender", None)
    if out.get("payment_mode"):
        pm = _normalize_payment_mode_sheet_value(out.get("payment_mode"))
        if pm:
            out["payment_mode"] = pm
        else:
            out.pop("payment_mode", None)
    if "insurer" in out:
        ins = sanitize_details_sheet_insurer_value(out.get("insurer"))
        if ins:
            out["insurer"] = ins
        else:
            out.pop("insurer", None)
    if out.get("nominee_relationship"):
        out["nominee_relationship"] = normalize_nominee_relationship_value(out["nominee_relationship"])
    rfn = _refine_nominee_relationship_with_gender(out.get("nominee_relationship"), out.get("nominee_gender"))
    if rfn:
        out["nominee_relationship"] = rfn
    return out


def _map_key_value_pairs_to_details_customer(pairs: list[dict]) -> dict[str, str]:
    """Map Details sheet key-value pairs to customer-side fields from new A5 template."""
    out: dict[str, str] = {}
    key_lower_to_value: dict[str, str] = {}
    for kv in pairs:
        k = (kv.get("key") or "").strip()
        v = (kv.get("value") or "").strip()
        if not k:
            continue
        key_norm = _normalize_key_for_match(k)
        key_lower_to_value[key_norm] = v
        if ":" in key_norm:
            key_lower_to_value[key_norm.replace(":", "").strip()] = v

    for field, aliases in _DETAILS_CUSTOMER_KEY_ALIASES.items():
        if field in out:
            continue
        for alias in aliases:
            anorm = _normalize_key_for_match(alias)
            if anorm in key_lower_to_value and key_lower_to_value[anorm].strip():
                out[field] = key_lower_to_value[anorm].strip()
                break
            for k, v in key_lower_to_value.items():
                if (anorm in k or k in anorm) and v.strip():
                    out[field] = v.strip()
                    break
            if field in out:
                break

    for _f in ("profession", "marital_status", "nominee_gender", "nominee_relationship", "payment_mode"):
        if _f in out and out[_f]:
            out[_f] = _normalize_kv_value_for_checkbox_fields(_f, out[_f])

    if out.get("mobile_number"):
        digits = "".join(ch for ch in out["mobile_number"] if ch.isdigit())
        out["mobile_number"] = digits[-10:] if digits else ""
    if out.get("alt_phone_num"):
        digits = "".join(ch for ch in out["alt_phone_num"] if ch.isdigit())
        out["alt_phone_num"] = digits[-10:] if digits else ""
    if out.get("aadhar_id"):
        out["aadhar_id"] = _aadhar_last4(out.get("aadhar_id")) or ""
    if out.get("nominee_age"):
        out["nominee_age"] = _sanitize_nominee_age(out["nominee_age"]) or ""
    if "profession" in out:
        sp = _sanitize_details_profession_value(out.get("profession"))
        out["profession"] = sp if sp else default_profession_if_empty("")
    if out.get("marital_status"):
        ms = _normalize_details_marital_status_value(out["marital_status"])
        if ms:
            out["marital_status"] = ms
        else:
            out.pop("marital_status", None)
    if out.get("nominee_gender"):
        ng = _normalize_nominee_gender_sheet_value(out.get("nominee_gender"))
        if ng:
            out["nominee_gender"] = ng
        else:
            out.pop("nominee_gender", None)
    if out.get("payment_mode"):
        pm = _normalize_payment_mode_sheet_value(out.get("payment_mode"))
        if pm:
            out["payment_mode"] = pm
        else:
            out.pop("payment_mode", None)
    if out.get("nominee_relationship"):
        out["nominee_relationship"] = normalize_nominee_relationship_value(out["nominee_relationship"])
    rfn = _refine_nominee_relationship_with_gender(out.get("nominee_relationship"), out.get("nominee_gender"))
    if rfn:
        out["nominee_relationship"] = rfn
    return out


def _full_text_from_sales_detail_sheet_heading(text: str) -> str:
    """Keep text from the first line containing ``Sales Detail Sheet`` downward (docx / extra preambles)."""
    if not text or not isinstance(text, str):
        return text if isinstance(text, str) else ""
    lines = text.splitlines()
    for i, ln in enumerate(lines):
        if re.search(r"(?i)sales\s*detail\s*sheet", ln):
            return "\n".join(lines[i:]).strip()
    return text.strip()


# Textract FORMS may omit checkbox ticks on scanned PDFs; LINE ``full_text`` + TABLE cells still carry marks.
_CHECKBOX_MERGE_FIELDS = frozenset(
    {"profession", "marital_status", "nominee_gender", "nominee_relationship", "payment_mode"}
)


def _checkbox_field_fully_resolved(field: str, val: str | None) -> bool:
    """True when ``val`` already parsed to a usable checkbox / normalized value for ``field``."""
    if val is None or not str(val).strip():
        return False
    v = str(val).strip()
    if _extract_checkbox_selection_value(v, field) is not None:
        return True
    if field == "profession":
        return bool(_sanitize_details_profession_value(v))
    if field == "marital_status":
        return bool(_normalize_details_marital_status_value(v))
    if field == "nominee_gender":
        return bool(_normalize_nominee_gender_sheet_value(v))
    if field == "nominee_relationship":
        return bool(normalize_nominee_relationship_value(v))
    if field == "payment_mode":
        return bool(_normalize_payment_mode_sheet_value(v))
    return bool(v)


def _parse_sales_detail_checkbox_regions(full_text: str) -> dict[str, str]:
    """
    Scan a window after each printed label — checkbox ticks often appear in LINE text without FORMS pairs.
    """
    out: dict[str, str] = {}
    if not full_text or not str(full_text).strip():
        return out
    text = _full_text_from_sales_detail_sheet_heading(full_text)
    text = text.replace("\r", "\n")
    # (regex after heading, field)
    regions: list[tuple[str, str]] = [
        (r"(?i)profession(?:\s+of\s+customer)?", "profession"),
        (r"(?i)(?:customer\s+)?marital\s*status", "marital_status"),
        (r"(?i)nominee\s*relationship(?:\s*with\s*customer)?", "nominee_relationship"),
        (r"(?i)relationship\s*with\s*customer", "nominee_relationship"),
        (r"(?i)nominee\s*gender|gender\s*of\s*nominee|sex\s*of\s*nominee", "nominee_gender"),
    ]
    for pat, field in regions:
        if field in out:
            continue
        for m in re.finditer(pat, text):
            region = text[m.start() : m.start() + 700]
            cb = _extract_checkbox_selection_value(region, field)
            if cb:
                out[field] = cb
                break
            tail = text[m.end() : m.end() + 400]
            cb = _extract_checkbox_selection_value(tail, field)
            if cb:
                out[field] = cb
                break
    return out


def _parse_sales_detail_checkbox_from_tables(tables: list[list[list[str]]]) -> dict[str, str]:
    """Use TABLE rows as a single string — checkbox rows are often emitted as tables, not KEY_VALUE_SET."""
    out: dict[str, str] = {}
    for table in tables or []:
        for row in table:
            cells = [str(c or "").strip() for c in row if c and str(c).strip()]
            if not cells:
                continue
            joined = " ".join(cells)
            if len(joined) < 6:
                continue
            low = joined.lower()
            if "profession" in low and "profession" not in out:
                cb = _extract_checkbox_selection_value(joined, "profession")
                if cb:
                    out["profession"] = cb
            if "marital" in low and "status" in low and "marital_status" not in out:
                cb = _extract_checkbox_selection_value(joined, "marital_status")
                if cb:
                    out["marital_status"] = cb
            if (
                ("nominee" in low and "relationship" in low)
                or ("relationship" in low and "customer" in low)
            ) and "nominee_relationship" not in out:
                cb = _extract_checkbox_selection_value(joined, "nominee_relationship")
                if cb:
                    out["nominee_relationship"] = cb
            if "nominee" in low and "gender" in low and "nominee_gender" not in out:
                cb = _extract_checkbox_selection_value(joined, "nominee_gender")
                if cb:
                    out["nominee_gender"] = cb
    return out


def _normalize_scanned_checkbox_candidate(field: str, cand: str) -> str:
    if not cand or not str(cand).strip():
        return ""
    v = _normalize_kv_value_for_checkbox_fields(field, cand.strip())
    if field == "profession":
        sp = _sanitize_details_profession_value(v)
        return sp if sp else ""
    if field == "marital_status":
        ms = _normalize_details_marital_status_value(v)
        return ms if ms else ""
    if field == "nominee_gender":
        ng = _normalize_nominee_gender_sheet_value(v)
        return ng if ng else ""
    if field == "nominee_relationship":
        nr = normalize_nominee_relationship_value(v) or v
        return nr.strip()
    if field == "payment_mode":
        pm = _normalize_payment_mode_sheet_value(v)
        return pm if pm else ""
    return v.strip()


def _apply_sales_detail_checkbox_scan(
    insurance: dict[str, str],
    details_customer: dict[str, str],
    *,
    full_text: str,
    tables: list | None,
) -> None:
    """Prefer region/table parses when FORMS values are missing or did not resolve checkboxes."""
    scan: dict[str, str] = {}
    if full_text and str(full_text).strip():
        scan.update(_parse_sales_detail_checkbox_regions(full_text))
    if tables:
        for k, v in _parse_sales_detail_checkbox_from_tables(tables).items():
            if v and (k not in scan or not scan[k]):
                scan[k] = v
    for field, cand in scan.items():
        if not cand:
            continue
        prev = insurance.get(field) or details_customer.get(field)
        if prev and _checkbox_field_fully_resolved(field, prev):
            continue
        norm = _normalize_scanned_checkbox_candidate(field, cand)
        if not norm:
            continue
        insurance[field] = norm
        details_customer[field] = norm


def _merge_textract_details_fallbacks(
    insurance: dict[str, str],
    details_customer: dict[str, str],
    *,
    full_text: str | None,
    tables: list | None = None,
) -> dict[str, str]:
    """
    Merge LINE/table fallbacks into ``insurance`` (and checkbox fields into ``details_customer``).

    Returns extra keys from the full-text parser that are not stored on ``insurance`` (e.g. ``customer_name``).
    """
    extra: dict[str, str] = {}
    ft = (full_text or "").strip()
    from_full: dict[str, str] = _parse_insurance_from_full_text(ft) if ft else {}
    for k, v in from_full.items():
        if k == "customer_name":
            if v:
                extra["customer_name"] = v.strip()
            continue
        if not v:
            continue
        if k in _CHECKBOX_MERGE_FIELDS:
            prev = insurance.get(k)
            if prev and _checkbox_field_fully_resolved(k, prev):
                continue
        elif insurance.get(k):
            continue
        insurance[k] = v.strip()
    _apply_sales_detail_checkbox_scan(
        insurance,
        details_customer,
        full_text=ft,
        tables=tables,
    )
    return extra


def _parse_insurance_from_full_text(full_text: str) -> dict[str, str]:
    """Try to extract insurance/customer extras from full text (fallback when not in key-value pairs)."""
    out: dict[str, str] = {}
    if not full_text or not isinstance(full_text, str):
        return out
    text = _full_text_from_sales_detail_sheet_heading(full_text.strip())
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    pay_fin = _extract_financier_from_payment_line(text)
    if pay_fin:
        out["financier"] = pay_fin
    # Patterns: "Label" or "Label:" followed by value on same or next line
    patterns = [
        # Detail-sheet lines e.g. "Insurer Name (if needed): SOMPO" (longer labels first)
        ("insurer name (if needed)", "insurer"),
        ("insurer name if needed", "insurer"),
        ("name of insurer", "insurer"),
        # Avoid bare "insurer name" here — it matches inside "Insurer Name (if needed)" and captures the wrong group.
        ("insurance provider", "insurer"),
        ("customer profession", "profession"),
        ("profession of customer", "profession"),
        ("profession", "profession"),
        ("occupation", "profession"),
        ("nature of occupation", "profession"),
        ("name of financier", "financier"),
        ("financing bank", "financier"),
        ("finance company", "financier"),
        ("financier / bank", "financier"),
        ("bank / financier", "financier"),
        ("financier", "financier"),
        ("mode of payment", "payment_mode"),
        ("payment mode", "payment_mode"),
        ("marital status", "marital_status"),
        ("nominee gender", "nominee_gender"),
        ("sex of nominee", "nominee_gender"),
        ("name of the nominee", "nominee_name"),
        ("name of nominee", "nominee_name"),
        ("nominee name", "nominee_name"),
        ("nominee's name", "nominee_name"),
        ("nominee age", "nominee_age"),
        ("age of nominee", "nominee_age"),
        ("nominee relationship with customer", "nominee_relationship"),
        ("nominee relationship", "nominee_relationship"),
        ("relationship with customer", "nominee_relationship"),
        ("relationship of nominee", "nominee_relationship"),
        ("relation with proposer", "nominee_relationship"),
        ("customer name", "customer_name"),
        ("buyer name", "customer_name"),
        ("buyer's name", "customer_name"),
    ]
    for label, key in patterns:
        if key in out:
            continue
        pat = re.compile(
            rf"(?im){re.escape(label)}\s*[:\t]?\s*\n?\s*([^\n]+)",
        )
        m = pat.search(text)
        if m:
            val = m.group(1).strip()
            if key == "profession":
                cb = _extract_checkbox_selection_value(val, "profession")
                if cb:
                    val = cb
                val = _sanitize_details_profession_value(val)
                if not val:
                    continue
            elif key == "marital_status":
                cb = _extract_checkbox_selection_value(val, "marital_status")
                if cb:
                    val = cb
                val = _normalize_details_marital_status_value(val)
                if not val:
                    continue
            elif key == "nominee_gender":
                cb = _extract_checkbox_selection_value(val, "nominee_gender")
                if cb:
                    val = cb
                val = _normalize_nominee_gender_sheet_value(val)
                if not val:
                    continue
            elif key == "payment_mode":
                val = _normalize_payment_mode_sheet_value(val)
                if not val:
                    continue
            elif key == "insurer":
                val = sanitize_details_sheet_insurer_value(val)
                if not val:
                    continue
            elif key == "nominee_relationship":
                cb = _extract_checkbox_selection_value(val, "nominee_relationship")
                if cb:
                    val = cb
                val = normalize_nominee_relationship_value(val)
                if not val:
                    continue
            elif not val or len(val) >= 200:
                continue
            out[key] = val

    # Word / table layouts: label alone on one line, value on the next (e.g. "Profession" then "Farmer")
    for i, ln in enumerate(lines):
        low = _normalize_key_for_match(ln)
        nxt = lines[i + 1].strip() if i + 1 < len(lines) else ""
        if not nxt or len(nxt) > 180:
            continue
        if "profession" not in out and low in ("profession", "occupation", "customer profession", "employment"):
            cand = _extract_checkbox_selection_value(nxt, "profession") or _sanitize_details_profession_value(nxt)
            if cand:
                out["profession"] = cand
        if "financier" not in out and low in (
            "financier",
            "financing bank",
            "name of financier",
            "finance company",
            "bank name",
        ):
            out["financier"] = nxt
        if "nominee_name" not in out and ("nominee" in low and "name" in low) and "age" not in low:
            out["nominee_name"] = nxt
        if "nominee_age" not in out and low in ("nominee age", "age of nominee", "nominee age (years)"):
            out["nominee_age"] = _sanitize_nominee_age(nxt) or nxt
        if "nominee_gender" not in out and low in ("nominee gender", "gender of nominee", "sex of nominee"):
            ng = _extract_checkbox_selection_value(nxt, "nominee_gender") or _normalize_nominee_gender_sheet_value(nxt)
            if ng:
                out["nominee_gender"] = ng
        if "nominee_relationship" not in out and low in (
            "nominee relationship",
            "relationship with customer",
            "relation with proposer",
            "relation to insured",
        ):
            nr = _extract_checkbox_selection_value(nxt, "nominee_relationship") or normalize_nominee_relationship_value(
                nxt
            )
            if nr:
                out["nominee_relationship"] = nr
        same = re.match(r"(?i)^(profession|occupation|customer profession)\s+(.{1,120})$", ln)
        if same and "profession" not in out:
            rest = same.group(2).strip()
            cand = _extract_checkbox_selection_value(rest, "profession") or _sanitize_details_profession_value(rest)
            if cand:
                out["profession"] = cand
        same_f = re.match(
            r"(?i)^(financier|financing bank|name of financier)\s+(.{1,160})$",
            ln,
        )
        if same_f and "financier" not in out:
            out["financier"] = same_f.group(2).strip()
        if "insurer" not in out and "insurer" in low and "name" in low and "nominee" not in low:
            if nxt and len(nxt) < 120 and not re.match(r"^[:\s_]+$", nxt):
                cand = sanitize_details_sheet_insurer_value(nxt.strip())
                if cand:
                    out["insurer"] = cand

    # Single LINE blocks with label + checkbox row (new template): "Profession: [✓] Private [ ] Job ..."
    _inline_checkbox = [
        (re.compile(r"(?i)^\s*Profession\s*:\s*(.+)$"), "profession"),
        (re.compile(r"(?i)^\s*Marital\s+Status\s*:\s*(.+)$"), "marital_status"),
        (re.compile(r"(?i)^\s*Nominee\s+Gender\s*:\s*(.+)$"), "nominee_gender"),
        (re.compile(r"(?i)^\s*Nominee\s+Relationship\s*:\s*(.+)$"), "nominee_relationship"),
        (re.compile(r"(?i)^\s*Relation\s*:\s*(.+)$"), "nominee_relationship"),
        (re.compile(r"(?i)^\s*Payment\s*:\s*(.+)$"), "payment_mode"),
    ]
    for ln in lines:
        for rx, field in _inline_checkbox:
            key = field
            if key in out:
                continue
            m = rx.match(ln)
            if not m:
                continue
            rest = m.group(1).strip()
            if field == "profession":
                cand = _extract_checkbox_selection_value(rest, field) or _sanitize_details_profession_value(rest)
                if cand:
                    out[key] = cand
            elif field == "marital_status":
                cand = _extract_checkbox_selection_value(rest, field) or _normalize_details_marital_status_value(rest)
                if cand:
                    out[key] = cand
            elif field == "nominee_gender":
                cand = _extract_checkbox_selection_value(rest, field) or _normalize_nominee_gender_sheet_value(rest)
                if cand:
                    out[key] = cand
            elif field == "nominee_relationship":
                cand = _extract_checkbox_selection_value(rest, field) or normalize_nominee_relationship_value(rest)
                if cand:
                    out[key] = cand
            elif field == "payment_mode":
                cand = _normalize_payment_mode_sheet_value(rest)
                if cand:
                    out[key] = cand
            break

    rfn = _refine_nominee_relationship_with_gender(out.get("nominee_relationship"), out.get("nominee_gender"))
    if rfn:
        out["nominee_relationship"] = rfn
    return out


def _map_key_value_pairs_to_insurance_policy(pairs: list[dict]) -> dict[str, str]:
    """Map key_value_pairs from Textract to insurance policy fields (insurer, policy_from, policy_to, premium)."""
    out: dict[str, str] = {}
    key_lower_to_value: dict[str, str] = {}
    for kv in pairs:
        k = (kv.get("key") or "").strip()
        v = (kv.get("value") or "").strip()
        if not k:
            continue
        key_norm = _normalize_key_for_match(k)
        key_lower_to_value[key_norm] = v
        if ":" in key_norm:
            key_lower_to_value[key_norm.replace(":", "").strip()] = v

    for field, aliases in _INSURANCE_POLICY_KEY_ALIASES.items():
        if field in out:
            continue
        for alias in aliases:
            anorm = _normalize_key_for_match(alias)
            if anorm in key_lower_to_value:
                out[field] = key_lower_to_value[anorm]
                break
            for k, v in key_lower_to_value.items():
                if anorm in k or k in anorm:
                    out[field] = v
                    break
            if field in out:
                break
    return out


def _parse_insurance_policy_from_full_text(full_text: str) -> dict[str, str]:
    """Extract policy number, insurer, policy from/to, gross premium from insurance policy full text.
    Document layout:
      - Policy / Cert. No. -> policy number (e.g. 39010231216200202663)
      - Right after -> insurance provider (e.g. NATIONAL INSURANCE COMPANY LIMITED)
      - Policy Period -> two dates: first=from, second=to (e.g. 15-06-2022, 15-06-2023)
      - Gross Premium -> premium amount (e.g. 5291.00)
    """
    out: dict[str, str] = {}
    if not full_text or not isinstance(full_text, str):
        return out
    text = full_text.strip()

    # 1. Policy number: "Policy\nCert. No.\n39010231216200202663" - number after Cert. No.
    if "policy_num" not in out:
        m = re.search(
            r"(?:Policy\s+)?Cert\.?\s*No\.?\s*:?\s*\n\s*(\d{12,})",
            text,
            re.IGNORECASE,
        )
        if m:
            out["policy_num"] = m.group(1).strip()

    # 2. Insurer: right after policy number - "NATIONAL INSURANCE COMPANY LIMITED"
    if "insurer" not in out:
        m = re.search(
            r"(\d{12,})\s*\n\s*([A-Z][A-Za-z\s]+(?:INSURANCE\s+COMPANY\s+LIMITED|INSURANCE\s+COMPANY|INSURANCE))",
            text,
        )
        if m:
            out["insurer"] = m.group(2).strip()
        else:
            # Fallback: known insurer names
            m = re.search(
                r"([A-Z][A-Za-z\s]+INSURANCE[A-Za-z\s]*(?:COMPANY\s+)?(?:LIMITED|LTD\.?)?)",
                text,
            )
            if m:
                out["insurer"] = m.group(1).strip()

    # 3. Policy from & to:
    #    - policy_from: Date in "dd-mm-yyyy To" block before "OD Policy Period" (e.g. 16-06-2021 near UIN block)
    #    - policy_to: First dd-mm-yyyy after "OD Policy Period" (e.g. 15-06-2022)
    if "policy_from" not in out or "policy_to" not in out:
        date_pat = re.compile(r"\b(\d{1,2}-\d{1,2}-\d{4})\b")
        od_match = re.search(r"OD\s+Policy\s+Period\s*:?\s*\n", text, re.IGNORECASE)
        if od_match:
            before_od = text[: od_match.start()]
            after_od = text[od_match.end() :]
            if "policy_from" not in out:
                # Prefer date in "dd-mm-yyyy To" format (policy period start); else first date before OD
                m_from = re.search(r"(\d{1,2}-\d{1,2}-\d{4})\s+To", before_od)
                if m_from:
                    out["policy_from"] = m_from.group(1)
                else:
                    dates_before = date_pat.findall(before_od)
                    if dates_before:
                        out["policy_from"] = dates_before[0]
            if "policy_to" not in out:
                dates_after = date_pat.findall(after_od)
                if dates_after:
                    out["policy_to"] = dates_after[0]  # First date after OD = policy end (e.g. 15-06-2022)
        if "policy_from" not in out or "policy_to" not in out:
            # Fallback: Policy Period with two dates
            for label in [r"Policy\s+Period", r"OD\s+Policy\s+Period"]:
                m = re.search(
                    rf"{label}\s*:?\s*\n\s*(\d{{1,2}}-\d{{1,2}}-\d{{4}})\s*\n\s*(\d{{1,2}}-\d{{1,2}}-\d{{4}})",
                    text,
                    re.IGNORECASE,
                )
                if m:
                    if "policy_from" not in out:
                        out["policy_from"] = m.group(1)
                    if "policy_to" not in out:
                        out["policy_to"] = m.group(2)
                    break

    # 4. Gross Premium: "Gross Premium\n5291.00"
    if "premium" not in out:
        m = re.search(r"Gross\s+Premium\s*:?\s*\n\s*([\d,]+\.?\d*)", text, re.IGNORECASE)
        if m:
            out["premium"] = m.group(1).strip().replace(",", "")
        else:
            m = re.search(r"Premium\s+of\s+Rs\.?\s*:?\s*\n\s*([\d,]+\.?\d*)", text, re.IGNORECASE)
            if m:
                out["premium"] = m.group(1).strip().replace(",", "")

    # 5. Policy holder / insured name (multiple patterns for different document formats)
    if "policy_holder_name" not in out:
        for label in [
            r"Name\s+of\s+(?:the\s+)?(?:insured|proposer|policy\s*holder)",
            r"Policy\s+holder\s*:?",
            r"Registered\s+Owner\s*:?",
            r"Name\s+of\s+insured\s*:?",
            r"Proposer['\u2019]?s?\s+name\s*:?",
            r"Name\s+of\s+proposer\s*:?",
            r"Insured\s+name\s*:?",
            r"(?:Name|Proposer)\s*:?\s*\n",  # "Name:" or "Proposer:" on own line
        ]:
            m = re.search(rf"{label}\s*\n?\s*([A-Za-z][A-Za-z\s\.]{1,79})", text, re.IGNORECASE)
            if m:
                val = m.group(1).strip()
                # Exclude values that look like insurer names or numbers
                if len(val) >= 2 and len(val) <= 80 and not re.match(r"^\d+$", val):
                    if "insurance" not in val.lower() and "company" not in val.lower() and "ltd" not in val.lower():
                        out["policy_holder_name"] = val
                        break

    return out


class OcrService:
    """Process AI reader queue with Textract (forms mode); write results to flat files. Details sheet only."""

    def __init__(
        self,
        uploads_dir: Path | None = None,
        ocr_output_dir: Path | None = None,
    ):
        from app.config import OCR_OUTPUT_DIR, UPLOADS_DIR

        self.uploads_dir = Path(uploads_dir or UPLOADS_DIR).resolve()
        self.ocr_output_dir = Path(ocr_output_dir or OCR_OUTPUT_DIR).resolve()

    def _ensure_ocr_output_dir(self) -> None:
        self.ocr_output_dir.mkdir(parents=True, exist_ok=True)

    def get_output_path(self, subfolder: str, filename: str) -> Path:
        """Path where extracted text is written: ocr_output/mobile_ddmmyy/filename_stem.txt."""
        self._ensure_ocr_output_dir()
        stem = Path(filename).stem
        safe_stem = re.sub(r"[^\w\-.]", "_", stem)
        # Always use a subfolder (mobile_ddmmyy); never write at ocr_output root
        subfolder_name = _safe_subfolder_name(subfolder)
        subfolder_path = self.ocr_output_dir / subfolder_name
        subfolder_path.mkdir(parents=True, exist_ok=True)
        return subfolder_path / f"{safe_stem}.txt"

    def read_extraction_file(self, subfolder: str, filename: str) -> str | None:
        """Read extracted text from the flat file if it exists."""
        path = self.get_output_path(subfolder, filename)
        if not path.exists():
            return None
        return path.read_text(encoding="utf-8", errors="replace")

    def process_next(self) -> dict | None:
        """
        Process the oldest queued item: Details sheet (Textract) or Aadhar (Vision).
        Picks any queued item; branches on filename (details vs aadhar). Returns result dict or None if none queued.
        """
        with get_connection() as conn:
            AiReaderQueueRepository.ensure_table(conn)
            row = AiReaderQueueRepository.get_oldest_queued(conn)  # any queued item
            if not row:
                return None

            qid = row["id"]
            subfolder = row["subfolder"]
            filename = row["filename"]

            AiReaderQueueRepository.update_status(conn, qid, "processing")
            conn.commit()

        input_path = self.uploads_dir / subfolder / filename
        if not input_path.exists():
            with get_connection() as conn:
                AiReaderQueueRepository.update_status(conn, qid, "failed")
                conn.commit()
            return {
                "id": qid,
                "subfolder": subfolder,
                "filename": filename,
                "status": "failed",
                "error": "File not found",
                "extracted_text": None,
                "output_path": None,
                "document_type": None,
                "classification_confidence": None,
            }

        fn_lower = filename.lower()
        if any(hint in fn_lower for hint in DETAILS_FILENAME_HINTS):
            return self._process_details_sheet(qid, subfolder, filename, input_path)
        if AADHAR_FILENAME_CONTAINS in fn_lower:
            return self._process_aadhar(qid, subfolder, filename, input_path)
        # Unknown type: mark done so queue advances
        with get_connection() as conn:
            AiReaderQueueRepository.update_classification(conn, qid, "Other", 0.0)
            AiReaderQueueRepository.update_status(conn, qid, "done")
            conn.commit()
        return {
            "id": qid,
            "subfolder": subfolder,
            "filename": filename,
            "status": "done",
            "error": None,
            "extracted_text": None,
            "output_path": None,
            "document_type": "Other",
            "classification_confidence": 0.0,
        }

    def _write_aadhar_fields_summary_file(self, subfolder: str, customer: dict[str, str]) -> None:
        self._ensure_ocr_output_dir()
        output_path = self.get_output_path(subfolder, "Aadhar_front_fields.txt")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        lines = ["Aadhar scan – 15 extracted fields", ""]
        for key, label in AADHAR_15_FIELDS:
            value = customer.get(key)
            if value and str(value).strip():
                lines.append(f"{label}: {value.strip()}")
        if customer.get("address") and str(customer["address"]).strip():
            lines.append(f"Address (constructed): {customer['address'].strip()}")
        output_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

    def _persist_upload_merge(
        self,
        subfolder: str,
        frag_a: dict[str, Any] | None,
        frag_d: dict[str, Any] | None,
    ) -> dict[str, str]:
        """Single JSON write: Aadhaar customer fragment + Details vehicle/insurance (after parallel compile)."""
        customer: dict[str, str] = {}
        vehicle: dict = {}
        insurance_merged: dict = {}
        details_customer_name = None

        if frag_a and frag_a.get("ok", True) and frag_a.get("customer") is not None:
            customer = dict(frag_a["customer"])

        if frag_d and not frag_d.get("error"):
            vehicle = frag_d.get("vehicle") or {}
            insurance = frag_d.get("insurance") or {}
            dc = frag_d.get("details_customer") or {}
            details_customer_name = frag_d.get("details_customer_name")

            for key, val in dc.items():
                if not val:
                    continue
                if key in (
                    "profession",
                    "marital_status",
                    "financier",
                    "nominee_name",
                    "nominee_age",
                    "nominee_gender",
                    "nominee_relationship",
                    "payment_mode",
                ):
                    continue
                if not customer.get(key):
                    customer[key] = val

            insurance_merged = {
                **{k: v for k, v in insurance.items() if v},
                **{
                    k: v
                    for k, v in dc.items()
                    if k
                    in (
                        "profession",
                        "marital_status",
                        "financier",
                        "nominee_name",
                        "nominee_age",
                        "nominee_gender",
                        "nominee_relationship",
                        "payment_mode",
                    )
                    and v
                },
            }
            if "profession" in insurance_merged:
                sp_im = _sanitize_details_profession_value(insurance_merged.get("profession"))
                insurance_merged["profession"] = sp_im if sp_im else default_profession_if_empty("")

            if details_customer_name and frag_a and frag_a.get("raw_parts"):
                blob = _concat_aadhar_scan_ocr_text(frag_a["raw_parts"])
                if blob.strip():
                    reconciled = _reconcile_customer_name_aadhar_details(
                        customer.get("name"),
                        details_customer_name,
                        blob,
                    )
                    if reconciled:
                        customer["name"] = reconciled

        customer = enrich_customer_address_from_freeform(customer)

        extraction_error = None
        if not _aadhar_identity_ok(customer):
            extraction_error = (
                (frag_a or {}).get("extraction_note")
                or "Aadhaar automated read did not yield a clear name or ID number. "
                "Enter or verify customer fields manually."
            )

        json_path = _json_output_path(self.ocr_output_dir, subfolder)
        json_path.parent.mkdir(parents=True, exist_ok=True)
        details_json: dict[str, Any] = {
            "vehicle": vehicle,
            "customer": customer,
            "insurance": insurance_merged,
        }
        if details_customer_name:
            details_json["details_customer_name"] = details_customer_name
        if extraction_error:
            details_json["extraction_error"] = extraction_error
        json_path.write_text(json.dumps(details_json, indent=2), encoding="utf-8")
        return customer

    def _upload_fragment_aadhar(
        self,
        scan_bundle: dict[str, Any],
        prefetch: dict[str, dict],
    ) -> dict[str, Any]:
        customer, raw_parts, note, process_timings = _pipeline_merge_aadhar_customer(
            scan_bundle.get("front_bytes"),
            scan_bundle.get("back_bytes"),
            prefetch.get("aadhar_front"),
            prefetch.get("aadhar_back"),
            front_raw_name=str(scan_bundle.get("front_src") or FILENAME_AADHAR_FRONT),
            back_raw_name=str(scan_bundle.get("back_src") or "Aadhar_back.jpg"),
        )
        return {
            "ok": True,
            "customer": customer,
            "raw_parts": raw_parts,
            "extraction_note": note,
            "process_timings": process_timings,
        }

    def _upload_fragment_details(
        self,
        details_path: Path,
        prefetch: dict[str, dict],
    ) -> dict[str, Any]:
        try:
            frag = _compile_details_sheet_fragment(details_path, prefetch.get("details_forms"))
            frag["ok"] = True
            return frag
        except Exception as e:
            return {"ok": False, "error": str(e)}

    def process_uploaded_subfolder(self, subfolder: str) -> dict:
        """
        Run extraction directly on uploaded files (no queue).

        Flow: **AWS Textract prefetch** (when enabled) on Aadhaar and other scans; **Aadhaar assembly**
        uses **Textract text only** on front/back (no UIDAI QR in this path). **Details sheet** is
        compiled in **parallel** with the Aadhaar pipeline, then results are merged once into JSON.
        **Details** raster/PDF: Textract **AnalyzeDocument FORMS** (structured key-values).

        Writes Raw_OCR.txt and returns ``section_timings_ms`` for operator visibility.
        """
        subdir = self.uploads_dir / subfolder
        if not subdir.exists() or not subdir.is_dir():
            return {"error": f"Subfolder not found: {subfolder}", "processed": []}

        processed: list[str] = []
        errors: list[str] = []
        section_timings_ms: dict[str, int] = {}
        t_total = time.perf_counter()
        append_ocr_extraction_log(
            self.ocr_output_dir,
            subfolder,
            "ocr",
            "started (Textract prefetch, parallel compile, merge, Raw_OCR)",
        )

        from concurrent.futures import ThreadPoolExecutor

        from app.config import OCR_UPLOAD_PARALLEL_TEXTRACT

        scan_bundle = _load_aadhar_scan_bytes(subdir)
        prefetch: dict[str, dict] = {}
        prefetch_job_ms: dict[str, int] = {}
        t_pref = time.perf_counter()
        if OCR_UPLOAD_PARALLEL_TEXTRACT:
            prefetch, prefetch_job_ms = _parallel_textract_prefetch_upload_subfolder(subdir)
            section_timings_ms["aws_textract_prefetch_ms"] = int(
                (time.perf_counter() - t_pref) * 1000
            )

        aadhar_path = _prefer_for_ocr_input(subdir, "Aadhar.pdf", FILENAME_AADHAR_FRONT, LEGACY_AADHAR_FRONT_JPG)
        details_path = _prefer_details_sheet_input(subdir)

        frag_a: dict[str, Any] | None = None
        frag_d: dict[str, Any] | None = None
        t_par = time.perf_counter()
        if aadhar_path.exists() or details_path.exists():
            with ThreadPoolExecutor(max_workers=2) as ex:
                fa = (
                    ex.submit(self._upload_fragment_aadhar, scan_bundle, prefetch)
                    if aadhar_path.exists()
                    else None
                )
                fb = (
                    ex.submit(self._upload_fragment_details, details_path, prefetch)
                    if details_path.exists()
                    else None
                )
                if fa is not None:
                    frag_a = fa.result()
                if fb is not None:
                    frag_d = fb.result()
        section_timings_ms["parallel_aadhar_details_compile_ms"] = int(
            (time.perf_counter() - t_par) * 1000
        )

        _merge_granular_upload_process_timings(
            section_timings_ms,
            prefetch_job_ms=prefetch_job_ms,
            frag_a=frag_a,
            frag_d=frag_d,
        )
        logger.info(
            "ocr_upload_process_timings subfolder=%s aadhar_textract_front_ms=%s "
            "aadhar_textract_back_ms=%s detail_sheet_textract_ms=%s",
            subfolder,
            section_timings_ms.get("aadhar_textract_front_ms", 0),
            section_timings_ms.get("aadhar_textract_back_ms", 0),
            section_timings_ms.get("detail_sheet_textract_ms", 0),
        )

        self._raw_ocr_parts = []
        if frag_a and frag_a.get("raw_parts"):
            self._raw_ocr_parts.extend(frag_a["raw_parts"])
        if frag_d and frag_d.get("ok") and frag_d.get("full_text"):
            self._raw_ocr_parts.append((_rel_upload_label(details_path, subdir), frag_d["full_text"]))

        if aadhar_path.exists():
            if frag_a and frag_a.get("ok"):
                processed.append(_rel_upload_label(aadhar_path, subdir))
            else:
                errors.append(
                    f"{_rel_upload_label(aadhar_path, subdir)}: {frag_a.get('error') if frag_a else 'failed'}"
                )
        if details_path.exists():
            if frag_d and frag_d.get("ok"):
                processed.append(_rel_upload_label(details_path, subdir))
            else:
                errors.append(
                    f"{_rel_upload_label(details_path, subdir)}: {frag_d.get('error') if frag_d else 'failed'}"
                )

        t_merge = time.perf_counter()
        merged_customer: dict[str, str] = {}
        if aadhar_path.exists() or details_path.exists():
            merged_customer = self._persist_upload_merge(subfolder, frag_a, frag_d)
            if aadhar_path.exists() and merged_customer:
                self._write_aadhar_fields_summary_file(subfolder, merged_customer)
        section_timings_ms["merge_write_json_ms"] = int((time.perf_counter() - t_merge) * 1000)

        insurance_path = _prefer_for_ocr_input(subdir, "Insurance.pdf", "Insurance.jpg")
        t_ins = time.perf_counter()
        if insurance_path.exists():
            try:
                ins_label = _rel_upload_label(insurance_path, subdir)
                self._process_insurance_sheet(
                    subfolder,
                    ins_label,
                    insurance_path,
                    textract_prefetch=prefetch.get("insurance"),
                )
                processed.append(ins_label)
            except Exception as e:
                errors.append(f"{_rel_upload_label(insurance_path, subdir)}: {e}")
        section_timings_ms["insurance_ms"] = int((time.perf_counter() - t_ins) * 1000)

        t_extras = time.perf_counter()
        seen_raw = {fn for fn, _ in self._raw_ocr_parts}
        for pdf_name, jpg_name, prefetch_key in (
            ("Aadhar_back.pdf", "Aadhar_back.jpg", "aadhar_back"),
            ("Financing.pdf", "Financing.jpg", "financing"),
        ):
            extra_path = _prefer_for_ocr_input(subdir, pdf_name, jpg_name)
            extra_label = _rel_upload_label(extra_path, subdir)
            if extra_label in seen_raw:
                continue
            if extra_path.exists():
                try:
                    from app.services.sales_textract_service import extract_text_from_bytes

                    if prefetch_key in prefetch:
                        result = prefetch[prefetch_key]
                    else:
                        result = extract_text_from_bytes(extra_path.read_bytes())
                    if not result.get("error") and result.get("full_text"):
                        self._raw_ocr_parts.append((extra_label, result["full_text"]))
                except Exception:
                    pass
        section_timings_ms["extras_raw_ms"] = int((time.perf_counter() - t_extras) * 1000)

        t_raw = time.perf_counter()
        if self._raw_ocr_parts:
            self._ensure_ocr_output_dir()
            subfolder_name = _safe_subfolder_name(subfolder)
            subfolder_path = self.ocr_output_dir / subfolder_name
            subfolder_path.mkdir(parents=True, exist_ok=True)
            raw_lines = []
            for fname, text in self._raw_ocr_parts:
                raw_lines.append(f"--- {fname} ---")
                raw_lines.append(text.strip() if text else "")
                raw_lines.append("")
            (subfolder_path / "Raw_OCR.txt").write_text("\n".join(raw_lines), encoding="utf-8")
            _apply_aadhar_textract_fallbacks_from_parts(
                self.ocr_output_dir, subfolder, list(self._raw_ocr_parts)
            )
        section_timings_ms["raw_ocr_finalize_ms"] = int((time.perf_counter() - t_raw) * 1000)

        self._raw_ocr_parts = None

        t_post = time.perf_counter()
        post_ocr_result: dict[str, Any] = {"ok": True, "skipped": True}
        try:
            from app.services.post_ocr_service import run_post_ocr

            post_ocr_result = run_post_ocr(self.uploads_dir, subfolder)
        except Exception as e:
            logger.exception("post_ocr failed subfolder=%s", subfolder)
            post_ocr_result = {"ok": False, "error": str(e)}
        section_timings_ms["post_ocr_ms"] = int((time.perf_counter() - t_post) * 1000)

        if post_ocr_result.get("ok") and "total_bytes_before" in post_ocr_result:
            append_ocr_extraction_log(
                self.ocr_output_dir,
                subfolder,
                "post",
                (
                    f"ok bytes_before={post_ocr_result.get('total_bytes_before')} "
                    f"bytes_after={post_ocr_result.get('total_bytes_after')} "
                    f"max_file_bytes={post_ocr_result.get('max_file_bytes')} "
                    f"still_over={len(post_ocr_result.get('files_still_over_limit') or [])} "
                    f"actions={len(post_ocr_result.get('actions') or [])}"
                ),
            )
        else:
            append_ocr_extraction_log(
                self.ocr_output_dir,
                subfolder,
                "post",
                f"failed error={post_ocr_result.get('error', 'unknown')!r}",
            )

        section_timings_ms["total_ms"] = int((time.perf_counter() - t_total) * 1000)

        result: dict[str, Any] = {
            "processed": processed,
            "section_timings_ms": section_timings_ms,
            "post_ocr": post_ocr_result,
        }
        if errors:
            result["errors"] = errors
        append_ocr_extraction_log(
            self.ocr_output_dir,
            subfolder,
            "ocr",
            (
                f"completed total_ms={section_timings_ms.get('total_ms')} "
                f"textract_prefetch_ms={section_timings_ms.get('aws_textract_prefetch_ms')} "
                f"parallel_compile_ms={section_timings_ms.get('parallel_aadhar_details_compile_ms')} "
                f"merge_write_ms={section_timings_ms.get('merge_write_json_ms')} "
                f"insurance_ms={section_timings_ms.get('insurance_ms')} "
                f"raw_finalize_ms={section_timings_ms.get('raw_ocr_finalize_ms')} "
                f"post_ocr_ms={section_timings_ms.get('post_ocr_ms')} "
                f"processed={len(processed)} errors={len(errors)}"
            ),
        )
        return result

    def _process_details_sheet(
        self,
        qid: int | None,
        subfolder: str,
        filename: str,
        input_path: Path,
        textract_forms_prefetch: dict | None = None,
    ) -> dict:
        """Run Textract (forms) on Details sheet; write text + JSON; optionally update queue."""
        if qid is not None:
            with get_connection() as conn:
                AiReaderQueueRepository.update_classification(conn, qid, "Details sheet", 1.0)
                conn.commit()
        try:
            fmt = _details_input_format(input_path)
            if fmt == "docx":
                key_value_pairs, docx_full = _key_value_pairs_from_docx(input_path)
                if not key_value_pairs and not (docx_full or "").strip():
                    raise RuntimeError(
                        "Could not read the Word document (.docx). Ensure it is a valid .docx Sales Detail Sheet."
                    )
                docx_full = _full_text_from_sales_detail_sheet_heading(docx_full or "")
                result = {
                    "error": None,
                    "full_text": docx_full,
                    "key_value_pairs": key_value_pairs,
                }
            elif fmt in ("jpeg", "png", "pdf"):
                from app.services.sales_textract_service import extract_forms_from_bytes

                if textract_forms_prefetch is not None:
                    result = textract_forms_prefetch
                else:
                    result = extract_forms_from_bytes(input_path.read_bytes())
                if result.get("error"):
                    raise RuntimeError(result["error"])
            else:
                raise RuntimeError(
                    f"Unsupported Details file format (detected={fmt!r}). "
                    "Use a JPEG/PNG scan, PDF export, or .docx Sales Detail Sheet."
                )

            key_value_pairs = result.get("key_value_pairs") or []
            lines = []
            lines.append("Document: Details sheet (Textract Forms)\n")
            for kv in key_value_pairs:
                lines.append(f"{kv.get('key', '')}: {kv.get('value', '')}")
            if result.get("full_text"):
                lines.append("\n--- Full text ---\n")
                lines.append(result["full_text"])
            text = "\n".join(lines)

            if hasattr(self, "_raw_ocr_parts") and self._raw_ocr_parts is not None:
                self._raw_ocr_parts.append((filename, result.get("full_text") or ""))

            self._ensure_ocr_output_dir()
            output_path = self.get_output_path(subfolder, filename)  # used for return value only; no separate .txt file

            vehicle = _map_key_value_pairs_to_vehicle(key_value_pairs)
            insurance = _map_key_value_pairs_to_insurance(key_value_pairs)
            details_customer = _map_key_value_pairs_to_details_customer(key_value_pairs)
            details_customer_name = _extract_details_customer_name(key_value_pairs)
            if result.get("full_text") or result.get("tables"):
                from_vehicle = _parse_vehicle_from_full_text(result["full_text"] or "")
                for k, v in from_vehicle.items():
                    if v and not vehicle.get(k):
                        vehicle[k] = v
                extra = _merge_textract_details_fallbacks(
                    insurance,
                    details_customer,
                    full_text=result.get("full_text"),
                    tables=result.get("tables"),
                )
                if extra.get("customer_name") and not details_customer_name:
                    details_customer_name = extra["customer_name"]
            json_path = _json_output_path(self.ocr_output_dir, subfolder)
            json_path.parent.mkdir(parents=True, exist_ok=True)
            customer = {}
            existing_insurance: dict = {}
            if json_path.exists():
                try:
                    existing = json.loads(json_path.read_text(encoding="utf-8"))
                    customer = existing.get("customer") or {}
                    if not isinstance(customer, dict):
                        customer = {}
                    # Compliance: never persist full Aadhar
                    if customer.get("aadhar_id"):
                        customer["aadhar_id"] = _aadhar_last4(customer["aadhar_id"]) or ""
                    existing_insurance = existing.get("insurance") or {}
                    if not isinstance(existing_insurance, dict):
                        existing_insurance = {}
                except Exception:
                    pass
            if details_customer.get("name") and not details_customer_name:
                details_customer_name = details_customer.get("name")

            customer_merged = dict(customer)
            for key, val in details_customer.items():
                if not val:
                    continue
                if key in (
                    "profession",
                    "marital_status",
                    "financier",
                    "nominee_name",
                    "nominee_age",
                    "nominee_gender",
                    "nominee_relationship",
                    "payment_mode",
                ):
                    continue
                if not customer_merged.get(key):
                    customer_merged[key] = val

            insurance_merged = {
                **existing_insurance,
                **{k: v for k, v in insurance.items() if v},
                **{
                    k: v
                    for k, v in details_customer.items()
                    if k
                    in (
                        "profession",
                        "marital_status",
                        "financier",
                        "nominee_name",
                        "nominee_age",
                        "nominee_gender",
                        "nominee_relationship",
                        "payment_mode",
                    )
                    and v
                },
            }
            if "profession" in insurance_merged:
                sp_im = _sanitize_details_profession_value(insurance_merged.get("profession"))
                insurance_merged["profession"] = sp_im if sp_im else default_profession_if_empty("")
            customer_merged = enrich_customer_address_from_freeform(customer_merged)
            details_json = {"vehicle": vehicle, "customer": customer_merged, "insurance": insurance_merged}
            if details_customer_name:
                details_json["details_customer_name"] = details_customer_name
            json_path.parent.mkdir(parents=True, exist_ok=True)
            json_path.write_text(json.dumps(details_json, indent=2), encoding="utf-8")

            if qid is not None:
                with get_connection() as conn:
                    AiReaderQueueRepository.update_status(conn, qid, "done")
                    conn.commit()

            return {
                "id": qid,
                "subfolder": subfolder,
                "filename": filename,
                "status": "done",
                "error": None,
                "extracted_text": text,
                "output_path": str(output_path),
                "document_type": "Details sheet",
                "classification_confidence": 1.0,
            }
        except Exception as e:
            if qid is not None:
                with get_connection() as conn:
                    AiReaderQueueRepository.update_status(conn, qid, "failed")
                    conn.commit()
                return {
                    "id": qid,
                    "subfolder": subfolder,
                    "filename": filename,
                    "status": "failed",
                    "error": str(e),
                    "extracted_text": None,
                    "output_path": None,
                    "document_type": None,
                    "classification_confidence": None,
                }
            raise

    def _process_aadhar(
        self,
        qid: int | None,
        subfolder: str,
        filename: str,
        input_path: Path,
        front_textract_prefetch: dict | None = None,
    ) -> dict:
        """Run Aadhar extraction: **AWS Textract** on front (and back when geo fields are weak); merge into subfolder JSON."""
        if qid is not None:
            with get_connection() as conn:
                AiReaderQueueRepository.update_classification(conn, qid, "Aadhar", 1.0)
                conn.commit()
        try:
            raw_bytes = input_path.read_bytes()
            back_path = input_path.parent / "Aadhar_back.jpg"
            back_bytes = back_path.read_bytes() if back_path.is_file() else None
            ft = front_textract_prefetch
            if ft is None:
                t_ft = time.perf_counter()
                try:
                    from app.services.sales_textract_service import extract_text_from_bytes

                    ft = extract_text_from_bytes(raw_bytes)
                except Exception:
                    ft = {"error": "textract", "full_text": ""}
                front_textract_sync_ms = int((time.perf_counter() - t_ft) * 1000)
            else:
                front_textract_sync_ms = 0

            piped, raw_parts, note, pipe_timings = _pipeline_merge_aadhar_customer(
                raw_bytes, back_bytes, ft, None
            )
            logger.info(
                "ocr_queue_aadhar_process_timings subfolder=%s "
                "aadhar_textract_front_prefetch_or_sync_ms=%s aadhar_textract_pipeline_front_ms=%s "
                "aadhar_textract_pipeline_back_ms=%s",
                subfolder,
                front_textract_sync_ms,
                pipe_timings.get("aadhar_textract_front_ms", 0),
                pipe_timings.get("aadhar_textract_back_ms", 0),
            )

            if hasattr(self, "_raw_ocr_parts") and self._raw_ocr_parts is not None:
                for fn, txt in raw_parts:
                    self._raw_ocr_parts.append((fn, txt))

            self._ensure_ocr_output_dir()
            json_path = _json_output_path(self.ocr_output_dir, subfolder)
            data = {"vehicle": {}, "customer": {}, "insurance": {}}
            existing_customer: dict[str, str] = {}
            if json_path.exists():
                try:
                    data = json.loads(json_path.read_text(encoding="utf-8"))
                    if not isinstance(data.get("vehicle"), dict):
                        data["vehicle"] = {}
                    if not isinstance(data.get("customer"), dict):
                        data["customer"] = {}
                    if not isinstance(data.get("insurance"), dict):
                        data["insurance"] = {}
                    ec = data.get("customer") or {}
                    if isinstance(ec, dict):
                        existing_customer = {str(k): str(v) for k, v in ec.items() if v is not None}
                except Exception:
                    pass
            customer = _merge_qr_customer_into_existing(existing_customer, piped)
            if customer.get("aadhar_id"):
                customer["aadhar_id"] = _aadhar_last4(customer["aadhar_id"]) or ""
            _default_gender_male_if_unread(customer)
            customer = enrich_customer_address_from_freeform(customer)
            data["customer"] = customer
            if note and not _aadhar_identity_ok(customer):
                data["extraction_error"] = note
            elif "extraction_error" in data and _aadhar_identity_ok(customer):
                data.pop("extraction_error", None)
            json_path.parent.mkdir(parents=True, exist_ok=True)
            json_path.write_text(json.dumps(data, indent=2), encoding="utf-8")

            # Write a file in ocr_output/subfolder listing the 15 extracted Aadhar fields (with display labels)
            output_path = self.get_output_path(subfolder, filename)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            lines = ["Aadhar scan – 15 extracted fields", ""]
            for key, label in AADHAR_15_FIELDS:
                value = customer.get(key)
                if value and str(value).strip():
                    lines.append(f"{label}: {value.strip()}")
            if customer.get("address") and str(customer["address"]).strip():
                lines.append(f"Address (constructed): {customer['address'].strip()}")
            summary = "\n".join(f"{k}: {v}" for k, v in customer.items() if v)
            output_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

            if qid is not None:
                with get_connection() as conn:
                    AiReaderQueueRepository.update_status(conn, qid, "done")
                    conn.commit()

            return {
                "id": qid,
                "subfolder": subfolder,
                "filename": filename,
                "status": "done",
                "error": None,
                "extracted_text": summary,
                "output_path": str(output_path),
                "document_type": "Aadhar",
                "classification_confidence": 1.0,
            }
        except Exception as e:
            if qid is not None:
                with get_connection() as conn:
                    AiReaderQueueRepository.update_status(conn, qid, "failed")
                    conn.commit()
                return {
                    "id": qid,
                    "subfolder": subfolder,
                    "filename": filename,
                    "status": "failed",
                    "error": str(e),
                    "extracted_text": None,
                    "output_path": None,
                    "document_type": None,
                    "classification_confidence": None,
                }
            raise

    def _process_insurance_sheet(
        self,
        subfolder: str,
        filename: str,
        input_path: Path,
        textract_prefetch: dict | None = None,
    ) -> None:
        """Run Textract on Insurance.jpg (TEXT mode); extract insurer, policy from/to, premium via regex on full_text."""
        from app.services.sales_textract_service import extract_text_from_bytes

        if textract_prefetch is not None:
            result = textract_prefetch
        else:
            result = extract_text_from_bytes(input_path.read_bytes())

        if result.get("error"):
            raise RuntimeError(result["error"])

        full_text = result.get("full_text") or ""
        insurance_policy = _parse_insurance_policy_from_full_text(full_text)

        if hasattr(self, "_raw_ocr_parts") and self._raw_ocr_parts is not None:
            self._raw_ocr_parts.append((filename, full_text))

        # Merge into Details.json insurance section
        json_path = _json_output_path(self.ocr_output_dir, subfolder)
        data: dict = {"vehicle": {}, "customer": {}, "insurance": {}}
        if json_path.exists():
            try:
                data = json.loads(json_path.read_text(encoding="utf-8"))
                if not isinstance(data.get("vehicle"), dict):
                    data["vehicle"] = {}
                if not isinstance(data.get("customer"), dict):
                    data["customer"] = {}
                if not isinstance(data.get("insurance"), dict):
                    data["insurance"] = {}
            except Exception:
                pass
        existing_insurance = data.get("insurance") or {}
        insurance_merged = {**existing_insurance, **{k: v for k, v in insurance_policy.items() if v}}
        data["insurance"] = insurance_merged
        # Compliance: sanitize customer aadhar before persisting
        customer = data.get("customer") or {}
        if customer.get("aadhar_id"):
            customer["aadhar_id"] = _aadhar_last4(customer["aadhar_id"]) or ""
        json_path.parent.mkdir(parents=True, exist_ok=True)
        json_path.write_text(json.dumps(data, indent=2), encoding="utf-8")

    def get_extracted_details(self, subfolder: str) -> dict | None:
        """
        Return structured extracted details (vehicle, customer) for a subfolder.
        Loads from JSON if present. Aadhaar fields are filled from upload-time Textract and
        ``Raw_OCR.txt`` fallbacks (no per-poll UIDAI QR decode).
        """
        self._ensure_ocr_output_dir()
        json_path = _json_output_path(self.ocr_output_dir, subfolder)
        data = {"vehicle": {}, "customer": {}}
        if json_path.exists():
            try:
                data = json.loads(json_path.read_text(encoding="utf-8"))
                if not isinstance(data.get("vehicle"), dict):
                    data["vehicle"] = {}
                if not isinstance(data.get("customer"), dict):
                    data["customer"] = {}
                if not isinstance(data.get("insurance"), dict):
                    data["insurance"] = {}
            except Exception:
                pass
        if "insurance" not in data:
            data["insurance"] = {}

        customer = data.get("customer") or {}
        if not isinstance(customer, dict):
            customer = {}
        subdir_gc = self.uploads_dir / subfolder
        aadhar_path = _prefer_for_ocr_input(subdir_gc, "Aadhar.pdf", FILENAME_AADHAR_FRONT, LEGACY_AADHAR_FRONT_JPG)
        back_p = _prefer_for_ocr_input(subdir_gc, "Aadhar_back.pdf", "Aadhar_back.jpg")

        # Textract text in Raw_OCR.txt: fill DOB/gender from front, address from back when parsers missed fields.
        _apply_aadhar_textract_fallbacks_from_raw_ocr_file(self.ocr_output_dir, subfolder)
        if json_path.exists():
            try:
                data = json.loads(json_path.read_text(encoding="utf-8"))
                if not isinstance(data.get("vehicle"), dict):
                    data["vehicle"] = {}
                if not isinstance(data.get("customer"), dict):
                    data["customer"] = {}
                if not isinstance(data.get("insurance"), dict):
                    data["insurance"] = {}
            except Exception:
                pass

        # Compliance: sanitize customer aadhar on return (handles legacy JSON with full Aadhar)
        customer = data.get("customer") or {}
        if customer.get("aadhar_id"):
            customer["aadhar_id"] = _aadhar_last4(customer["aadhar_id"]) or ""
        if aadhar_path.is_file() or back_p.is_file():
            _default_gender_male_if_unread(customer)
            customer = enrich_customer_address_from_freeform(customer)
            data["customer"] = customer
            try:
                json_path.parent.mkdir(parents=True, exist_ok=True)
                json_path.write_text(json.dumps(data, indent=2), encoding="utf-8")
            except Exception:
                pass
        # Sanitize nominee_age: extract digits only (e.g. "30/m" -> "30"); clear if invalid
        ins = data.get("insurance") or {}
        if ins.get("nominee_age"):
            sanitized = _sanitize_nominee_age(str(ins["nominee_age"]))
            ins["nominee_age"] = sanitized
        # Name match validation: Aadhar, Details sheet, Insurance
        aadhar_name = customer.get("name")
        details_name = data.get("details_customer_name")
        insurance = data.get("insurance") or {}
        insurance_name = insurance.get("policy_holder_name")
        name_err = _validate_name_match(
            aadhar_name,
            details_name,
            insurance_name,
            aadhar_last4=customer.get("aadhar_id"),
            details_aadhar_last4=(
                ((data.get("details_customer") or {}).get("aadhar_id") if isinstance(data.get("details_customer"), dict) else None)
                or customer.get("aadhar_id")
            ),
            insurance_aadhar_last4=insurance.get("aadhar_id") or insurance.get("aadhaar_no") or insurance.get("aadhar_no"),
        )
        if name_err:
            data["name_mismatch_error"] = name_err
        _apply_initcap_on_read(data)
        try:
            json_path.parent.mkdir(parents=True, exist_ok=True)
            json_path.write_text(json.dumps(data, indent=2), encoding="utf-8")
        except Exception:
            pass
        return data

    def list_extractions(self, limit: int = 200) -> list[dict]:
        """List queue items (oldest first) with extracted text from flat files."""
        with get_connection() as conn:
            AiReaderQueueRepository.ensure_table(conn)
            rows = AiReaderQueueRepository.list_all(conn, limit=limit)

        rows = list(reversed(rows))
        result = []
        for row in rows:
            r = dict(row)
            for k in ("created_at", "updated_at"):
                if k in r and isinstance(r[k], datetime):
                    r[k] = r[k].isoformat()
            text = self.read_extraction_file(r["subfolder"], r["filename"])
            r["extracted_text"] = text
            r["output_path"] = str(self.get_output_path(r["subfolder"], r["filename"]))
            result.append(r)
        return result


def validate_name_match_for_subfolder(ocr_output_dir: Path, subfolder: str) -> str | None:
    """
    Load OCR JSON for subfolder and validate name match across Aadhar, Details, Insurance.
    Returns error message if mismatch, None if OK.
    """
    json_path = _json_output_path(ocr_output_dir, subfolder)
    if not json_path.exists():
        return None
    try:
        data = json.loads(json_path.read_text(encoding="utf-8"))
    except Exception:
        return None
    customer = data.get("customer") or {}
    aadhar_name = customer.get("name")
    details_name = data.get("details_customer_name")
    insurance = data.get("insurance") or {}
    insurance_name = insurance.get("policy_holder_name")
    return _validate_name_match(
        aadhar_name,
        details_name,
        insurance_name,
        aadhar_last4=customer.get("aadhar_id"),
        details_aadhar_last4=(
            ((data.get("details_customer") or {}).get("aadhar_id") if isinstance(data.get("details_customer"), dict) else None)
            or customer.get("aadhar_id")
        ),
        insurance_aadhar_last4=insurance.get("aadhar_id") or insurance.get("aadhaar_no") or insurance.get("aadhar_no"),
    )
