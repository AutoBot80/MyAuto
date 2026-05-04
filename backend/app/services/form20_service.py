"""
Generate Form 20 (all pages) and Gate Pass from templates.
Form 20: Prefer templates/word/FORM 20 Template.docx (Word) -> fill placeholders -> convert to PDF.
         Else PDF overlay or HTML fallback.
Gate Pass: templates/word/Gate Pass Template.docx -> fill placeholders -> convert to PDF.
Saves Form 20.pdf and Gate Pass.pdf to Uploaded scans/subfolder.
"""
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from app.config import (
    UPLOADS_DIR,
    FORM20_TEMPLATE_SINGLE,
    FORM20_TEMPLATE_FRONT,
    FORM20_TEMPLATE_BACK,
    FORM20_TEMPLATE_DOCX,
    GATE_PASS_TEMPLATE_DOCX,
)

logger = logging.getLogger(__name__)

# ``care_of`` may already include ``S/o`` / ``W/o`` / ``D/o`` / ``C/o`` (Aadhaar OCR).
_CARE_OF_REL_PREFIX = re.compile(r"(?i)^(C|S|W|D)/o\s+")

# Backend dir for HTML templates
_BACKEND_DIR = Path(__file__).resolve().parent.parent.parent
TEMPLATES_DIR = _BACKEND_DIR / "templates"

# Field positions (x, y) in points for Form 20 front page. Adjust if template layout differs.
# Origin: top-left. y increases downward. Based on typical Form 20 layout.
FORM20_FIELD_POSITIONS = {
    "field_1_name_care_of": (120, 135),
    "field_3_address": (120, 175),
    "field_10_dealer_name": (120, 295),
    "field_14_body_type": (120, 365),
    "field_15_vehicle_type": (120, 390),
    "field_16_oem_name": (120, 415),
    "field_17_year_of_mfg": (120, 440),
    "field_18_num_cylinders": (120, 465),
    "field_19_horse_power": (120, 490),
    "field_21_model": (120, 535),
    "field_22_chassis_no": (120, 560),
    "field_24_seating_capacity": (120, 605),
    "field_25_fuel_type": (120, 630),
    "field_28_colour": (120, 675),
}


def _get_vehicle_from_db(vehicle_id: int) -> dict[str, Any]:
    """Fetch vehicle from vehicle_master by vehicle_id."""
    from app.db import get_connection

    conn = get_connection()
    try:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT vehicle_id, chassis, engine, key_num, model, colour,
                       year_of_mfg, cubic_capacity, body_type, seating_capacity,
                       oem_name, vehicle_type, num_cylinders, length_mm, fuel_type
                FROM vehicle_master WHERE vehicle_id = %s
                """,
                (vehicle_id,),
            )
            row = cur.fetchone()
            if not row:
                return {}
            return dict(row)
    except Exception:
        try:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    SELECT vehicle_id, chassis, engine, key_num, model, colour,
                           year_of_mfg, cubic_capacity, body_type, seating_capacity
                    FROM vehicle_master WHERE vehicle_id = %s
                    """,
                    (vehicle_id,),
                )
                row = cur.fetchone()
                return dict(row) if row else {}
        except Exception:
            return {}
    finally:
        conn.close()


def _get_dealer_from_db(dealer_id: int) -> dict[str, Any]:
    """Fetch dealer from dealer_ref by dealer_id. Includes oem_name from oem_ref."""
    from app.db import get_connection

    conn = get_connection()
    try:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT d.dealer_id, d.dealer_name, d.address, d.pin, d.city, d.state,
                       o.oem_name
                FROM dealer_ref d
                LEFT JOIN oem_ref o ON o.oem_id = d.oem_id
                WHERE d.dealer_id = %s
                """,
                (dealer_id,),
            )
            row = cur.fetchone()
            return dict(row) if row else {}
    except Exception:
        return {}
    finally:
        conn.close()


def _build_form20_data(
    customer: dict[str, Any],
    vehicle: dict[str, Any],
    vehicle_id: int | None = None,
    dealer_id: int | None = None,
) -> dict[str, str]:
    """Build Form 20 field values from customer and vehicle. Merge with DB vehicle if vehicle_id given."""
    data: dict[str, str] = {}

    # Merge vehicle from DB if we have vehicle_id
    db_vehicle: dict[str, Any] = {}
    if vehicle_id:
        db_vehicle = _get_vehicle_from_db(vehicle_id)

    v = {**vehicle, **db_vehicle}
    c = customer or {}

    def _str(val: Any) -> str:
        if val is None:
            return ""
        s = str(val).strip()
        return s if s else ""

    # Field 0: City (separate, for templates that use it)
    data["field_0_city"] = _str(c.get("city"))

    # Field 1: Name (separate) and Field 2: Care of (separate)
    name = _str(c.get("name"))
    care_of = _str(c.get("care_of"))
    data["field_1_name"] = name
    data["field_2_care_of"] = care_of
    if name and care_of:
        if _CARE_OF_REL_PREFIX.match(care_of):
            data["field_1_name_care_of"] = f"{name}, {care_of}"
        else:
            data["field_1_name_care_of"] = f"{name} S/O {care_of}"
    else:
        data["field_1_name_care_of"] = name or care_of or ""

    # Field 3: Address
    addr_parts = [
        _str(c.get("address")),
        _str(c.get("city")),
        _str(c.get("state")),
        _str(c.get("pin_code") or c.get("pin")),
    ]
    data["field_3_address"] = ", ".join(p for p in addr_parts if p)

    # Field 10: Dealer name and address (from dealer_ref)
    dealer: dict[str, Any] = {}
    if dealer_id:
        dealer = _get_dealer_from_db(dealer_id)
    dealer_name = _str(dealer.get("dealer_name"))
    dealer_addr_parts = [
        _str(dealer.get("address")),
        _str(dealer.get("city")),
        _str(dealer.get("state")),
        _str(dealer.get("pin") or dealer.get("pin_code")),
    ]
    dealer_address = ", ".join(p for p in dealer_addr_parts if p)
    data["field_10_dealer_name"] = dealer_name
    if dealer_name and dealer_address:
        data["field_10_dealer_name_address"] = f"{dealer_name}, {dealer_address}"
    else:
        data["field_10_dealer_name_address"] = dealer_name or dealer_address or ""

    # Field 14: Body type
    data["field_14_body_type"] = _str(v.get("body_type"))

    # Field 15: Vehicle type
    data["field_15_vehicle_type"] = _str(v.get("vehicle_type"))

    # Field 16: OEM / Make (vehicle oem_name, else dealer's oem_ref.oem_name, else make/maker from DMS)
    data["field_16_oem_name"] = (
        _str(v.get("oem_name"))
        or _str(dealer.get("oem_name"))
        or _str(v.get("make") or v.get("maker"))
    )

    # Field 17: Year of mfg
    data["field_17_year_of_mfg"] = _str(v.get("year_of_mfg"))

    # Field 18: No. of cylinders
    data["field_18_num_cylinders"] = _str(v.get("num_cylinders"))

    # Field 19: Horse power (not stored on vehicle_master; leave blank)
    data["field_19_horse_power"] = ""

    # Field 20: Cubic capacity (cc)
    data["field_20_cubic_capacity"] = _str(v.get("cubic_capacity"))

    # Field 21: Model
    data["field_21_model"] = _str(v.get("model") or v.get("oem_name"))

    # Field 22: Chassis no.
    data["field_22_chassis_no"] = _str(v.get("chassis") or v.get("frame_num") or v.get("frame_no"))

    # Field 24: Seating capacity
    data["field_24_seating_capacity"] = _str(v.get("seating_capacity"))

    # Field 25: Fuel type
    data["field_25_fuel_type"] = _str(v.get("fuel_type")) or "Petrol"

    # Field 28: Colour
    data["field_28_colour"] = _str(v.get("colour") or v.get("color"))

    # Footer: Dealer Saathi© Vehicle ID <vehicle_id> (left, 10pt)
    vid = vehicle_id if vehicle_id else v.get("vehicle_id")
    _em_dash = "\u2014"
    data["footer_text"] = f"Dealer Saathi\u00a9 Vehicle ID {vid if vid else _em_dash}"

    return data


def _build_gate_pass_data(
    customer: dict[str, Any],
    vehicle: dict[str, Any],
    vehicle_id: int | None = None,
    dealer_id: int | None = None,
) -> dict[str, str]:
    """Build Gate Pass field values from customer and vehicle."""
    data: dict[str, str] = {}
    db_vehicle = _get_vehicle_from_db(vehicle_id) if vehicle_id else {}
    dealer = _get_dealer_from_db(dealer_id) if dealer_id else {}
    v = {**vehicle, **db_vehicle}
    c = customer or {}

    def _str(val: Any) -> str:
        if val is None:
            return ""
        return str(val).strip() or ""

    # field_1_oem_name: OEM name (Welcome to X family)
    oem = _str(v.get("oem_name")) or _str(dealer.get("oem_name")) or _str(v.get("make") or v.get("maker"))
    data["field_1_oem_name"] = oem

    # field_0_today_date: Delivery date (today)
    data["field_0_today_date"] = datetime.now().strftime("%d/%m/%Y")

    # field_2_customer_name
    data["field_2_customer_name"] = _str(c.get("name"))

    # field_3_aadhar_id: Last 4 digits only (compliance - DB stores last 4; extraction may have full)
    raw = c.get("aadhar_id") or c.get("aadhar")
    digits = "".join(ch for ch in str(raw or "") if ch.isdigit())
    data["field_3_aadhar_id"] = digits[-4:] if len(digits) >= 4 else digits

    # field_4_model, field_5_color, field_6_key_num, field_7_chassis_num
    data["field_4_model"] = _str(v.get("model") or v.get("oem_name"))
    data["field_5_color"] = _str(v.get("colour") or v.get("color"))
    data["field_6_key_num"] = _str(v.get("key_num") or v.get("key_no"))
    data["field_7_chassis_num"] = _str(v.get("chassis") or v.get("frame_num") or v.get("frame_no"))

    return data


def _fill_docx_template(docx_path: Path, data: dict[str, str], out_path: Path) -> None:
    """Fill {{placeholder}} in Word document and save."""
    from docx import Document

    doc = Document(docx_path)
    replacements = {k: (v or "\u2014") for k, v in data.items()}

    for paragraph in doc.paragraphs:
        _simple_replace_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _simple_replace_paragraph(paragraph, replacements)
    for section in doc.sections:
        for p in section.header.paragraphs:
            _simple_replace_paragraph(p, replacements)
        for p in section.footer.paragraphs:
            _simple_replace_paragraph(p, replacements)

    doc.save(out_path)


def _simple_replace_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """Simple replace when python-docx-replace not available. Uses normal font."""
    full_text = paragraph.text
    if "{{" not in full_text:
        return
    new_text = full_text
    for key, val in replacements.items():
        new_text = new_text.replace("{{" + key + "}}", val or "\u2014")
    if new_text != full_text:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)


def _docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """Convert docx to PDF. Tries docx2pdf (Word on Windows), then LibreOffice."""
    docx_path = Path(docx_path).resolve()
    pdf_path = Path(pdf_path).resolve()

    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        return
    except ImportError:
        pass
    except Exception as e:
        logger.warning("form20: docx2pdf failed: %s", e)

    # Fallback: LibreOffice (outputs input.pdf in outdir)
    import shutil
    import subprocess

    libreoffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not libreoffice:
        # Windows: check common install paths when not in PATH
        for base in (Path("C:/Program Files/LibreOffice/program"), Path("C:/Program Files (x86)/LibreOffice/program")):
            exe = base / "soffice.exe"
            if exe.exists():
                libreoffice = str(exe)
                break
    if libreoffice:
        outdir = pdf_path.parent
        subprocess.run(
            [libreoffice, "--headless", "--convert-to", "pdf", "--outdir", str(outdir), str(docx_path)],
            check=True,
            capture_output=True,
            timeout=60,
        )
        # LibreOffice creates docx_stem.pdf in outdir
        converted = outdir / (docx_path.stem + ".pdf")
        if converted.exists():
            if converted.resolve() != pdf_path.resolve():
                pdf_path.unlink(missing_ok=True)  # Windows: remove existing before rename
                converted.rename(pdf_path)
        return

    raise RuntimeError(
        "Cannot convert docx to PDF. Install docx2pdf (pip install docx2pdf) with Word, or install LibreOffice."
    )


def _split_pdf_pages(pdf_path: Path, out_paths: list[Path]) -> None:
    """Split PDF into separate files (page 0 -> out_paths[0], page 1 -> out_paths[1], etc.)."""
    import fitz

    doc = fitz.open(pdf_path)
    try:
        for i, out_path in enumerate(out_paths):
            page_idx = min(i, doc.page_count - 1)  # If single page, use it for both
            new_doc = fitz.open()
            new_doc.insert_pdf(doc, from_page=page_idx, to_page=page_idx)
            new_doc.save(str(out_path))
            new_doc.close()
    finally:
        doc.close()


def _generate_form20_via_docx(
    docx_template: Path,
    subfolder_path: Path,
    data: dict[str, str],
) -> list[str]:
    """Generate Form 20 from Word template: fill placeholders, convert to PDF (all pages)."""
    import shutil
    import tempfile

    # Copy template to temp file first (avoids Permission denied when original is open in Word/OneDrive)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        shutil.copy2(docx_template, tmp_path)
        work_docx = tmp_path
    except PermissionError:
        tmp_path.unlink(missing_ok=True)
        raise PermissionError(
            f"Cannot read {docx_template}. Close the file if it's open in Word, and ensure OneDrive isn't locking it."
        ) from None

    form20_out = subfolder_path / "Form 20.pdf"
    form20_out.unlink(missing_ok=True)  # Remove existing so rename/overwrite works on Windows

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        filled_docx = Path(tmp.name)
    try:
        _fill_docx_template(work_docx, data, filled_docx)
        _docx_to_pdf(filled_docx, form20_out)
    except RuntimeError as e:
        logger.warning("form20: docx-to-PDF conversion failed (%s), falling back to HTML", e)
        return _generate_form20_via_html(subfolder_path, data)
    finally:
        work_docx.unlink(missing_ok=True)
        filled_docx.unlink(missing_ok=True)

    logger.info("form20: saved %s (from Word template, all pages)", form20_out)
    return ["Form 20.pdf"]


def _generate_gate_pass_via_docx(
    docx_template: Path,
    subfolder_path: Path,
    data: dict[str, str],
) -> list[str]:
    """Generate Gate Pass from Word template: fill placeholders, convert to PDF."""
    import shutil
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        shutil.copy2(docx_template, tmp_path)
        work_docx = tmp_path
    except PermissionError:
        tmp_path.unlink(missing_ok=True)
        logger.warning("gate_pass: Cannot read %s (file may be open)", docx_template)
        return []

    gate_pass_out = subfolder_path / "Gate Pass.pdf"
    gate_pass_out.unlink(missing_ok=True)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        filled_docx = Path(tmp.name)
    try:
        _fill_docx_template(work_docx, data, filled_docx)
        _docx_to_pdf(filled_docx, gate_pass_out)
    except RuntimeError as e:
        logger.warning("gate_pass: docx-to-PDF conversion failed: %s", e)
        return []
    finally:
        work_docx.unlink(missing_ok=True)
        filled_docx.unlink(missing_ok=True)

    logger.info("gate_pass: saved %s", gate_pass_out)
    return ["Gate Pass.pdf"]


def generate_gate_pass_pdf_only(
    subfolder: str,
    customer: dict[str, Any],
    vehicle: dict[str, Any],
    vehicle_id: int | None = None,
    dealer_id: int | None = None,
    uploads_dir: Path | None = None,
) -> Path:
    """
    Generate only ``Gate Pass.pdf`` from ``GATE_PASS_TEMPLATE_DOCX``
    (default: ``<project>/templates/word/Gate Pass Template.docx``).
    Word placeholders are filled via :func:`_build_gate_pass_data` (``{{field_*}}`` keys).
    """
    import re

    uploads_path = Path(uploads_dir or UPLOADS_DIR).resolve()
    safe_sub = re.sub(r"[^\w\-]", "_", (subfolder or "").strip()) or "default"
    subfolder_path = uploads_path / safe_sub
    subfolder_path.mkdir(parents=True, exist_ok=True)
    gate_pass_template = Path(GATE_PASS_TEMPLATE_DOCX).resolve()
    if not gate_pass_template.is_file():
        raise FileNotFoundError(
            f"Gate Pass template not found: {gate_pass_template}. "
            "Set GATE_PASS_TEMPLATE_DOCX or add templates/word/Gate Pass Template.docx"
        )
    gate_pass_data = _build_gate_pass_data(customer, vehicle, vehicle_id, dealer_id)
    saved = _generate_gate_pass_via_docx(gate_pass_template, subfolder_path, gate_pass_data)
    if not saved:
        raise RuntimeError(
            "Gate Pass PDF was not produced (docx fill or PDF conversion failed — see logs)."
        )
    out = subfolder_path / "Gate Pass.pdf"
    if not out.is_file():
        raise RuntimeError("Gate Pass.pdf missing after generation.")
    return out


def _overlay_on_page(
    page,
    data: dict[str, str],
    field_positions: dict[str, tuple[float, float]],
    fontsize: float = 10,
    footer_y: float = 820,
) -> None:
    """Overlay form data and footer on a single PDF page."""
    for key, (x, y) in field_positions.items():
        val = data.get(key)
        if val and key != "footer_text":
            text = str(val)[:80] if len(str(val)) > 80 else str(val)
            page.insert_text((x, y), text, fontsize=fontsize, fontname="helv", color=(0, 0, 0), overlay=True)
    footer = data.get("footer_text", "")
    if footer:
        page.insert_text((42, footer_y), footer, fontsize=8, fontname="helv", color=(0, 0, 0), overlay=True)


def _overlay_all_pages(
    template_path: Path,
    out_path: Path,
    data: dict[str, str],
    field_positions: dict[str, tuple[float, float]],
    fontsize: float = 10,
    footer_y: float = 820,
) -> None:
    """Overlay form data on page 0, footer on all pages. Saves combined PDF."""
    import fitz

    if not template_path.exists():
        raise FileNotFoundError(f"Form 20 template not found: {template_path}")

    doc = fitz.open(template_path)
    try:
        for i in range(doc.page_count):
            page = doc[i]
            positions = field_positions if i == 0 else {}
            _overlay_on_page(page, data, positions, fontsize, footer_y)
        doc.save(str(out_path))
    finally:
        doc.close()


def generate_form20_pdfs(
    subfolder: str,
    customer: dict[str, Any],
    vehicle: dict[str, Any],
    vehicle_id: int | None = None,
    dealer_id: int | None = None,
    uploads_dir: Path | None = None,
) -> list[str]:
    """
    Generate Form 20.pdf and Gate Pass.pdf by overlaying data on templates.
    Form 20: templates/word/FORM 20 Template.docx, or Official FORM-20, or Form 20 Front/back PDFs.
    Gate Pass: templates/word/Gate Pass Template.docx.
    Saves to Uploaded scans/subfolder.
    """
    import re

    uploads_path = Path(uploads_dir or UPLOADS_DIR).resolve()
    safe_sub = re.sub(r"[^\w\-]", "_", (subfolder or "").strip()) or "default"
    subfolder_path = uploads_path / safe_sub
    subfolder_path.mkdir(parents=True, exist_ok=True)

    data = _build_form20_data(customer, vehicle, vehicle_id, dealer_id)

    # Template paths: prefer Word, then PDF overlay
    uploads_path = Path(uploads_dir or UPLOADS_DIR).resolve()
    project_root = uploads_path.parent

    docx_template = Path(FORM20_TEMPLATE_DOCX).resolve()

    single_template = Path(FORM20_TEMPLATE_SINGLE).resolve()
    front_template = Path(FORM20_TEMPLATE_FRONT).resolve()
    back_template = Path(FORM20_TEMPLATE_BACK).resolve()
    if not single_template.exists():
        fallback = project_root / "Raw Scans" / "Official FORM-20 english.pdf"
        if fallback.exists():
            single_template = fallback.resolve()
    if not front_template.exists():
        fallback = project_root / "Raw Scans" / "Form 20 Front.pdf"
        if fallback.exists():
            front_template = fallback.resolve()
    if not back_template.exists():
        fallback = project_root / "Raw Scans" / "Form 20 back.pdf"
        if fallback.exists():
            back_template = fallback.resolve()

    use_docx = docx_template.exists()
    use_single = single_template.exists()
    use_separate = front_template.exists() and back_template.exists()

    logger.info(
        "form20: docx=%s exists=%s | single=%s | separate=%s",
        docx_template,
        use_docx,
        use_single,
        use_separate,
    )

    def _add_gate_pass(saved: list[str]) -> list[str]:
        """Append Gate Pass.pdf if template exists and generation succeeds."""
        gate_pass_template = Path(GATE_PASS_TEMPLATE_DOCX).resolve()
        if gate_pass_template.exists():
            gate_pass_data = _build_gate_pass_data(customer, vehicle, vehicle_id, dealer_id)
            saved = saved + _generate_gate_pass_via_docx(gate_pass_template, subfolder_path, gate_pass_data)
        return saved

    # Prefer Word template - do not fall back on failure so user sees the actual error
    if use_docx:
        saved = _generate_form20_via_docx(docx_template, subfolder_path, data)
        return _add_gate_pass(saved)

    if not use_single and not use_separate:
        logger.warning("form20: No PDF templates found at %s or %s/%s. Using HTML fallback.", single_template, front_template, back_template)
        saved = _generate_form20_via_html(subfolder_path, data)
        return _add_gate_pass(saved)

    try:
        import fitz  # noqa: F401
    except ImportError as e:
        logger.warning("form20: pymupdf not installed (%s). Using HTML fallback. Run: pip install pymupdf", e)
        saved = _generate_form20_via_html(subfolder_path, data)
        return _add_gate_pass(saved)

    form20_out = subfolder_path / "Form 20.pdf"

    if use_single:
        # Official FORM-20: overlay on all pages (front, back, page 3 if present)
        try:
            _overlay_all_pages(
                single_template,
                form20_out,
                data,
                FORM20_FIELD_POSITIONS,
                fontsize=10,
                footer_y=820,
            )
            logger.info("form20: saved %s (from Official FORM-20, all pages)", form20_out)
        except Exception as e:
            logger.warning("form20: PDF generation failed: %s", e)
            raise
    else:
        # Separate front and back templates: overlay each, merge with optional page 3
        front_tmp = subfolder_path / "_form20_pdf_front.pdf"
        back_tmp = subfolder_path / "_form20_pdf_back.pdf"
        _overlay_all_pages(front_template, front_tmp, data, FORM20_FIELD_POSITIONS, fontsize=10, footer_y=820)
        _overlay_all_pages(back_template, back_tmp, data, {}, fontsize=10, footer_y=820)

        page3_template = project_root / "Raw Scans" / "Form 20 page 3.pdf"
        if not page3_template.exists():
            page3_template = project_root / "Raw Scans" / "Form 20 Page 3.pdf"
        to_merge = [front_tmp, back_tmp]
        if page3_template.exists():
            page3_tmp = subfolder_path / "_form20_pdf_page3.pdf"
            _overlay_all_pages(page3_template, page3_tmp, data, {}, fontsize=10, footer_y=820)
            to_merge.append(page3_tmp)

        merged = fitz.open()
        for p in to_merge:
            merged.insert_pdf(fitz.open(p), from_page=0, to_page=-1)
            p.unlink(missing_ok=True)
        merged.save(str(form20_out))
        merged.close()
        logger.info("form20: saved %s (from separate templates, %d pages)", form20_out, len(to_merge))

    saved = ["Form 20.pdf"]
    return _add_gate_pass(saved)


def _generate_form20_via_html(
    subfolder_path: Path,
    data: dict[str, str],
) -> list[str]:
    """Fallback: generate Form 20 using HTML templates when PDF templates are missing."""
    front_html = (TEMPLATES_DIR / "form20_front.html").read_text(encoding="utf-8")
    back_html = (TEMPLATES_DIR / "form20_back.html").read_text(encoding="utf-8")
    page3_html = (TEMPLATES_DIR / "form20_page3.html").read_text(encoding="utf-8")
    for key, val in data.items():
        front_html = front_html.replace("{{" + key + "}}", val or "\u2014")
        back_html = back_html.replace("{{" + key + "}}", val or "\u2014")
        page3_html = page3_html.replace("{{" + key + "}}", val or "\u2014")
    front_html = re.sub(r"\{\{[^}]+\}\}", "\u2014", front_html)
    back_html = re.sub(r"\{\{[^}]+\}\}", "\u2014", back_html)
    page3_html = re.sub(r"\{\{[^}]+\}\}", "\u2014", page3_html)

    from xhtml2pdf import pisa

    temp_pdfs = [
        subfolder_path / "_form20_html_front.pdf",
        subfolder_path / "_form20_html_back.pdf",
        subfolder_path / "_form20_html_page3.pdf",
    ]
    for html, pdf_path in [(front_html, temp_pdfs[0]), (back_html, temp_pdfs[1]), (page3_html, temp_pdfs[2])]:
        with open(pdf_path, "w+b") as dest:
            status = pisa.CreatePDF(html, dest=dest, encoding="utf-8")
        if status.err:
            for p in temp_pdfs:
                p.unlink(missing_ok=True)
            raise RuntimeError(f"xhtml2pdf error: {status.err}")

    # Merge into one Form 20.pdf
    import fitz
    out_path = subfolder_path / "Form 20.pdf"
    merged = fitz.open()
    for p in temp_pdfs:
        merged.insert_pdf(fitz.open(p), from_page=0, to_page=-1)
        p.unlink(missing_ok=True)
    merged.save(str(out_path))
    merged.close()

    logger.info("form20: saved %s (HTML fallback, 3 pages)", out_path)
    return ["Form 20.pdf"]
